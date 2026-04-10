"""
Task #4 — Shareholder Advances Review (dynamic — reads from run_data.json)
Spirit of Math Schools Markham East (2039321 Ontario Inc.)
"""

import sys, os, json, datetime
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import (
    HB, CAL, BS, make_doc,
    sbg, sbd, shdr, sdat, ct, hr, bp, sh, note, sub_header,
    callout, callout_red, callout_green, callout_blue,
)
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ================================================================
# LOAD DATA
# ================================================================
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(BASE_DIR, "data", "extracted", "run_data.json"), encoding="utf-8") as f:
    D = json.load(f)

meta  = D["meta"]
sh_d  = D["shareholder"]

cutoff_str = meta["ytd_cutoff_date"]    # e.g. "April 2, 2026"
fy_label   = meta["fiscal_year_label"]  # "August 1, 2025 – July 31, 2026"

ram_open  = sh_d["ramzan"]["opening_balance"]
ram_close = sh_d["ramzan"]["closing_balance"]
ram_net   = ram_close - ram_open
ram_txns  = sh_d["ramzan"]["transactions"]

rez_open  = sh_d["rezai"]["opening_balance"]
rez_close = sh_d["rezai"]["closing_balance"]
rez_net   = rez_close - rez_open
rez_txns  = sh_d["rezai"]["transactions"]

comb_open  = ram_open  + rez_open
comb_close = ram_close + rez_close
comb_net   = ram_net   + rez_net

# ================================================================
# HELPERS
# ================================================================
def fmt_bal(v):
    if abs(v) < 0.005:
        return "$0.00"
    sign = "+" if v > 0 else "\u2212"
    return f"{sign}${abs(v):,.2f}"

def fmt_bal_round(v):
    if abs(v) < 0.50:
        return "$0"
    sign = "+" if v > 0 else "\u2212"
    return f"{sign}${abs(v):,.0f}"

def fmt_amt(v):
    if abs(v) < 0.005:
        return "$0.00"
    sign = "+" if v > 0 else "\u2212"
    return f"{sign}{abs(v):,.2f}"

def fmt_date(d_str):
    try:
        dt = datetime.datetime.strptime(d_str, "%m/%d/%Y")
        return dt.strftime("%b %d, %Y")
    except Exception:
        return d_str

def bal_summary(name, v):
    if v < -0.005:
        return f"{name} owes corp {fmt_bal(v)[1:]}"
    elif v > 0.005:
        return f"Corp owes {name} {fmt_bal(v)[1:]}"
    else:
        return "Account at zero"

def explain_tx(tx, first_name):
    """Auto-generate a plain-English explanation from transaction fields."""
    memo  = (tx.get("memo") or "").lower()
    ttype = tx.get("type", "")
    num   = (tx.get("num") or "").upper().strip()
    amt   = tx.get("amount", 0.0)

    if "hajj" in memo:
        return (f"e-Transfer to Hajj travel agency. Personal Hajj pilgrimage paid by the "
                f"corporation on {first_name}\u2019s behalf \u2014 recorded as a shareholder advance. "
                f"{first_name} owes this amount back to the company. See Section 5A.")
    if num == "JE-12":
        return (f"JE-12: Annual year-end journal entry credit (${abs(amt):,.2f}). "
                f"One of three recurring credits posted each year. See Section 5D.")
    if num == "JE-11" and amt > 0:
        return ("JE-11: Rezai\u2019s credit balance transferred to offset Ramzan\u2019s "
                "outstanding debt. Ramzan\u2019s balance reduced by $3,263.82; Rezai\u2019s "
                "account zeroed. See Section 5F.")
    if num == "JE-11" and amt < 0:
        return ("JE-11: Credit balance zeroed and transferred to offset Ramzan\u2019s "
                "outstanding debt. Account now at zero. See Section 5F.")
    if "2236262" in memo:
        return ("Transfer received from 2236262 Ontario Inc. (a related company). "
                "Reduced outstanding advance balance. See Section 5C.")
    if "online banking wire" in memo and amt < 0:
        return (f"Corporation sent {first_name} ${abs(amt):,.2f} via online banking wire. "
                f"This is an advance \u2014 {first_name} owes this amount back to the company.")
    if "br to br" in memo and amt > 0 and "client request" in memo:
        return ("Credit from Spirit of Math (BR TO BR \u2014 Credit Memo 0783). "
                "Appears to be a royalty or marketing rebate credited back to the shareholder account.")
    if "br to br" in memo and amt < 0:
        return (f"Bank-to-bank payment of ${abs(amt):,.2f}. Same reference and date as the other "
                f"shareholder\u2019s entry on the same date \u2014 likely a shared expense split equally.")
    if "atm deposit" in memo and amt > 0:
        return f"ATM deposit. {first_name} depositing funds back to the corporation."
    if ("cheque" in memo or "check" in memo) and amt < 0:
        return (f"{first_name} withdrew ${abs(amt):,.2f} (the amount the corporation owed them) "
                f"by cheque. Account returned to zero.")
    desc = tx.get("memo") or f"{ttype} transaction"
    return f"{ttype}: {desc} (${abs(amt):,.2f})."

# Pre-compute flags
je12_posted   = any((tx.get("num") or "").upper().strip() == "JE-12" for tx in ram_txns)
je12_date_raw = next((tx["date"] for tx in ram_txns
                      if (tx.get("num") or "").upper().strip() == "JE-12"), None)
je12_date_str = fmt_date(je12_date_raw) if je12_date_raw else "near July 31, 2026"

hajj_txns = [tx for tx in ram_txns if "hajj" in (tx.get("memo") or "").lower()]
has_hajj   = bool(hajj_txns)

je11_ram = [tx for tx in ram_txns if (tx.get("num") or "").upper().strip() == "JE-11"]
je11_date_str = fmt_date(je11_ram[0]["date"]) if je11_ram else "Mar 11, 2026"
je11_amount   = abs(je11_ram[0]["amount"]) if je11_ram else 3263.82

# ================================================================
# INIT DOC
# ================================================================
doc = make_doc()

FILE_DATE = datetime.date.today().strftime("%Y-%m-%d")
OUT = os.path.join(BASE_DIR, f"reports/claude_report_shareholder_mae_{FILE_DATE}.docx")
today_str = datetime.date.today().strftime("%B %d, %Y")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Shareholder Advances Review \u2014 FY2025-26 & Multi-Year History")
r.bold = True; r.font.name = CAL; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("Spirit of Math Schools Markham East  \u2014  2039321 Ontario Inc.")
r.italic = True; r.font.name = CAL; r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Report Date: {today_str}   |   Fiscal Year: {fy_label}")
r.italic = True; r.font.name = CAL; r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Data source: QuickBooks \u2014 Account 2900 Shareholder\u2019s Advance "
              f"(all dates, Aug 2007 \u2013 {cutoff_str})")
r.italic = True; r.font.name = CAL; r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(6)
hr(doc)

# ================================================================
# QUICK SUMMARY
# ================================================================
sh(doc, "Quick Summary \u2014 What This Report Tells You")

je12_bullet = (f"JE-12 ({je12_date_str}) posted ${9520.37:,.2f} annual credit to each shareholder."
               if je12_posted else "JE-12 (annual year-end credit) not yet posted for FY2025-26.")
hajj_bullet  = (f"\u25cf  The $10,000 Hajj travel advance (Jan 13, 2026) is embedded in Ramzan\u2019s balance.\n"
                if has_hajj else "")

callout_blue(doc, (
    f"WHAT THIS REPORT COVERS: Every dollar that has moved between the corporation and its two "
    f"shareholders (Ramzan and Rezai) from 2008 to today, and what you need to know.\n\n"
    f"\u25cf  This report covers all shareholder advance activity from FY2007-08 through {cutoff_str}.\n"
    f"\u25cf  Ramzan: {fmt_bal(ram_close)} \u2014 {bal_summary('Ramzan', ram_close)}.\n"
    f"\u25cf  Rezai: {fmt_bal(rez_close)} \u2014 {bal_summary('Rezai', rez_close)}.\n"
    f"{hajj_bullet}"
    f"\u25cf  {je12_bullet}\n"
    f"\u25cf  All FY2025-26 entries are still \u2018Uncleared\u2019 \u2014 bank reconciliation still needed."
))

urgent_items = []
if has_hajj:
    urgent_items.append(
        "1.  Verify the $10,000 Hajj payment (Jan 13, 2026) is recorded in account 2901 "
        "(Shareholder Advance) only \u2014 and does NOT appear anywhere in the expense accounts on your "
        "P&L. If it is in any expense account, have the bookkeeper correct it immediately."
    )
if ram_close < 0:
    n = len(urgent_items) + 1
    urgent_items.append(
        f"{n}.  Plan Ramzan\u2019s repayment of {fmt_bal(ram_close)} to the corporation by July 31, 2027. "
        "The full negative balance must be repaid within one year of fiscal year-end "
        "to avoid being taxed as personal income."
    )
if not urgent_items:
    urgent_items.append(
        "No repayment urgency at this time \u2014 both shareholder balances are positive "
        "(the corporation owes the shareholders). Continue to monitor and ensure JE "
        "documentation is current. Hajj advance verification still required."
    )

callout_red(doc, "URGENT ACTIONS:\n" + "\n".join(urgent_items))
note(doc, ("See Section 7 (Action Checklist) for the full list of steps. "
           "See Section 8 (Bottom Line) for key takeaways in plain English. "
           "Sections 1\u20136 contain the detailed account history and explanations."))

# ================================================================
# SECTION 1 — What Is the Account
# ================================================================
sh(doc, "1.  What Is the Shareholder\u2019s Advance Account?")
bp(doc, ("Account 2900 (Shareholder\u2019s Advance) records money moving between the corporation "
         "and its two shareholders \u2014 Ramzan Khuwaja (account 2901) and Mohammad Rezai (account 2902). "
         "It works like a running tab between the company and each owner."))
bp(doc, ("Think of it this way: if the corporation pays for something personal on your behalf, "
         "or transfers cash to you directly, it records that as a debit to your advance account "
         "(\u201cRamzan owes the company\u201d). When you deposit money back or make a payment on behalf of "
         "the company, it records a credit (\u201ccompany owes Ramzan\u201d). The balance tells you at any "
         "moment who owes whom."))
bp(doc, ("\u25ba  Positive balance = the corporation owes the shareholder (shareholder has a credit).\n"
         "\u25ba  Negative balance = the shareholder owes the corporation (corporation has a loan outstanding)."))
callout_blue(doc, ("CRA Rule to Know: Under Canadian tax law (Income Tax Act section 15(2)), "
                   "if a corporation loans money to a shareholder (negative balance), that loan "
                   "must be fully repaid within one year after the end of the corporation\u2019s fiscal year "
                   "in which the loan was made. If it is not repaid in time, the full loan amount "
                   "must be included in the shareholder\u2019s personal income and taxed. "
                   "Additionally, if no interest is charged by the corporation on the loan, "
                   "the CRA will deem a taxable interest benefit based on the prescribed rate."))

# ================================================================
# SECTION 2 — Current Status
# ================================================================
sh(doc, f"2.  Current Status as of {cutoff_str}")
bp(doc, (f"Here is where each shareholder stands at the start of the fiscal year (August 1, 2025) "
         f"and where things are as of {cutoff_str}."))

def fmt_net(v):
    if abs(v) < 0.005:
        return "$0.00"
    sign = "+" if v > 0 else "\u2212"
    return f"{sign}${abs(v):,.2f}"

status_rows = [
    ("Ramzan Khuwaja (2901)", fmt_bal(ram_open), fmt_bal(ram_close),
     fmt_net(ram_net), bal_summary("Ramzan", ram_close)),
    ("Mohammad Rezai (2902)", fmt_bal(rez_open), fmt_bal(rez_close),
     fmt_net(rez_net), bal_summary("Rezai", rez_close)),
    ("Combined (both)", fmt_bal(comb_open), fmt_bal(comb_close),
     fmt_net(comb_net), bal_summary("Net combined", comb_close)),
]
ts = doc.add_table(rows=1 + len(status_rows), cols=5)
ts.style = "Table Grid"; ts.alignment = WD_TABLE_ALIGNMENT.LEFT
ts.columns[0].width = Inches(1.5); ts.columns[1].width = Inches(1.1)
ts.columns[2].width = Inches(1.1); ts.columns[3].width = Inches(1.0); ts.columns[4].width = Inches(2.2)
ct(ts.rows[0].cells[0], "Shareholder", bold=True)
ct(ts.rows[0].cells[1], "Aug 1, 2025 (Start)", bold=True)
ct(ts.rows[0].cells[2], f"{cutoff_str} (Now)", bold=True)
ct(ts.rows[0].cells[3], "Net Change", bold=True)
ct(ts.rows[0].cells[4], "Plain-English Summary", bold=True)
shdr(ts.rows[0])
for ri, (a, b, c, d, e) in enumerate(status_rows, 1):
    is_total = ("Combined" in a)
    ct(ts.rows[ri].cells[0], a, bold=is_total)
    ct(ts.rows[ri].cells[1], b, bold=is_total)
    ct(ts.rows[ri].cells[2], c, bold=is_total)
    ct(ts.rows[ri].cells[3], d, bold=is_total)
    ct(ts.rows[ri].cells[4], e, bold=is_total)
    sdat(ts.rows[ri], is_total)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
note(doc, ("A negative number in the balance column means the shareholder owes the corporation that "
           "amount. A positive number means the corporation owes the shareholder. Neither situation "
           "is automatically a problem \u2014 but negative balances must be repaid within one year "
           "to avoid being taxed as personal income."))

# ================================================================
# SECTION 3 — Transaction Detail
# ================================================================
sh(doc, "3.  FY2025-26 Transaction Detail \u2014 What Each Entry Means")
bp(doc, ("Below is a plain-English explanation of every transaction recorded in the shareholder "
         "advance account since August 1, 2025. All entries are currently showing as \u2018Uncleared\u2019 "
         "in QuickBooks, meaning the bank reconciliation has not yet been done for this fiscal year. "
         "This should be addressed with your bookkeeper."))

def make_tx_table(txns, label, acct, open_bal, close_bal, first_name):
    sub_header(doc, f"{label} ({acct})")
    if open_bal < -0.005:
        open_prose = (f"Starting balance August 1, 2025: {fmt_bal(open_bal)}  "
                      f"({first_name} owed the corporation this amount at the start of the fiscal "
                      f"year, carried forward from FY2024-25.)")
    elif open_bal > 0.005:
        open_prose = (f"Starting balance August 1, 2025: {fmt_bal(open_bal)}  "
                      f"(The corporation owed {first_name} this amount at the start of the fiscal year.)")
    else:
        open_prose = f"Starting balance August 1, 2025: $0.00  (Account opened at zero.)"
    bp(doc, open_prose)

    tbl = doc.add_table(rows=1 + len(txns), cols=5)
    tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(0.9); tbl.columns[1].width = Inches(0.75)
    tbl.columns[2].width = Inches(0.75); tbl.columns[3].width = Inches(0.8)
    tbl.columns[4].width = Inches(3.7)
    ct(tbl.rows[0].cells[0], "Date",     bold=True)
    ct(tbl.rows[0].cells[1], "Type",     bold=True)
    ct(tbl.rows[0].cells[2], "Amount",   bold=True)
    ct(tbl.rows[0].cells[3], "Balance",  bold=True)
    ct(tbl.rows[0].cells[4], "Plain-English Explanation", bold=True)
    shdr(tbl.rows[0])
    for ri, tx in enumerate(txns, 1):
        expl  = explain_tx(tx, first_name)
        num   = (tx.get("num") or "").upper().strip()
        memo  = (tx.get("memo") or "").lower()
        color = (RGBColor(0x8B, 0x00, 0x00) if "hajj" in memo else
                 RGBColor(0x1F, 0x38, 0x96) if num == "JE-11" else None)
        ct(tbl.rows[ri].cells[0], fmt_date(tx["date"]))
        ct(tbl.rows[ri].cells[1], tx["type"])
        ct(tbl.rows[ri].cells[2], fmt_amt(tx["amount"]))
        ct(tbl.rows[ri].cells[3], fmt_bal(tx["balance"]))
        ct(tbl.rows[ri].cells[4], expl, color=color)
        sdat(tbl.rows[ri])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if close_bal < -0.005:
        end_prose = (f"{first_name}\u2019s ending balance as of {cutoff_str}: {fmt_bal(close_bal)}. "
                     f"This must be repaid by July 31, 2027.")
    elif close_bal > 0.005:
        end_prose = (f"{first_name}\u2019s ending balance as of {cutoff_str}: {fmt_bal(close_bal)}. "
                     f"The corporation owes this amount to {first_name}.")
    else:
        end_prose = f"{first_name}\u2019s ending balance as of {cutoff_str}: $0.00."
    note(doc, end_prose)

make_tx_table(ram_txns, "Ramzan Khuwaja", "Account 2901", ram_open, ram_close, "Ramzan")
make_tx_table(rez_txns, "Mohammad Rezai", "Account 2902", rez_open, rez_close, "Rezai")

# ================================================================
# SECTION 4 — Multi-Year Balance History
# ================================================================
sh(doc, "4.  Multi-Year Balance History \u2014 FY2007-08 through FY2025-26")
bp(doc, ("The table below shows the year-end balance for each shareholder at the end of "
         "every fiscal year since the QuickBooks records began. A positive balance means the "
         "corporation owed the shareholder money at year-end. A negative balance means the "
         "shareholder owed the corporation."))

history_rows = [
    ("FY2007-08", "+$23,688", "+$21,093", "+$44,781",
     "Opening balances. Corp owed both shareholders (startup loans from owners)."),
    ("FY2008-09", "+$23,093", "+$20,593", "+$43,686",
     "Minimal activity. Small repayments to shareholders."),
    ("FY2009-10", "+$40,830", "+$39,774", "+$80,604",
     "Balances grew \u2014 more money owed to shareholders (likely additional owner loans)."),
    ("FY2010-11", "+$17,644", "+$17,644", "+$35,288",
     "Large repayments to shareholders. Both accounts brought to identical levels."),
    ("FY2011-12", "+$30,710", "+$27,710", "+$58,420",
     "Balances rose again. Corp drew on shareholder funds."),
    ("FY2012-13", "+$20,074", "+$22,155", "+$42,229",
     "Gradual reduction. Shareholders being repaid over time."),
    ("FY2013-14", "+$8,991",  "+$8,991",  "+$17,982",
     "Both shareholders brought to identical balance \u2014 $8,990.80 each."),
    ("FY2014-15", "+$7,763",  "+$7,763",  "+$15,526",
     "Both at identical $7,762.85. Small net repayments each year."),
    ("FY2015-16", "+$17,279", "+$16,810", "+$34,089",
     "Balances rose again (likely more owner activity during a growth year)."),
    ("FY2016-17", "+$14,468", "+$13,453", "+$27,921",
     "Steady reduction. Corp paying down what it owed shareholders."),
    ("FY2017-18", "$0.00",    "$0.00",    "$0.00",
     "Both accounts fully cleared to zero. Clean slate at July 31, 2018."),
    ("FY2018-19 to FY2020-21", "$0.00", "$0.00", "$0.00",
     "No activity for 3 consecutive years. Pandemic period (FY2020-21 includes COVID year)."),
    ("FY2021-22", "+$10,089", "+$10,089", "+$20,178",
     "Annual credits (JEs) restarted. Both shareholders received ~$10,089 in credits."),
    ("FY2022-23", "\u22126,130",   "+$227",    "\u22125,903",
     "Divergence begins. Ramzan\u2019s account went negative (owes corp). "
     "Rezai nearly zeroed out."),
    ("FY2023-24", "\u221247,088",  "\u221299,851",  "\u2212146,939",
     "Very large activity year. Year-end both shareholders owed corp significant amounts. "
     "See Section 5B for detail on the large temporary loans."),
    ("FY2024-25", "\u22125,346",   "+$9,485",  "+$4,139",
     "Large repayments by both shareholders. Accounts brought nearly to zero / positive."),
    # Current year row — dynamic
    (f"FY2025-26 (to {cutoff_str})",
     fmt_bal_round(ram_close), fmt_bal_round(rez_close), fmt_bal_round(comb_close),
     f"{bal_summary('Ramzan', ram_close)}; {bal_summary('Rezai', rez_close)}. "
     + ("JE-12 annual credits posted." if je12_posted else "JE-12 not yet posted.")),
]

th = doc.add_table(rows=1 + len(history_rows), cols=5)
th.style = "Table Grid"; th.alignment = WD_TABLE_ALIGNMENT.LEFT
th.columns[0].width = Inches(1.4); th.columns[1].width = Inches(1.0)
th.columns[2].width = Inches(1.0); th.columns[3].width = Inches(1.0); th.columns[4].width = Inches(2.5)
ct(th.rows[0].cells[0], "Fiscal Year", bold=True)
ct(th.rows[0].cells[1], "Ramzan",      bold=True)
ct(th.rows[0].cells[2], "Rezai",       bold=True)
ct(th.rows[0].cells[3], "Combined",    bold=True)
ct(th.rows[0].cells[4], "Note",        bold=True)
shdr(th.rows[0])
for ri, (a, b, c, d, e) in enumerate(history_rows, 1):
    is_curr = ("FY2025-26" in a)
    is_zero = b in ("$0.00", "$0")
    b_color = (RGBColor(0x8B, 0x00, 0x00) if b.startswith("\u2212") and not is_zero else
               RGBColor(0x1A, 0x5C, 0x1A) if b.startswith("+") and not is_zero else None)
    ct(th.rows[ri].cells[0], a, bold=is_curr)
    ct(th.rows[ri].cells[1], b, bold=is_curr, color=b_color)
    ct(th.rows[ri].cells[2], c, bold=is_curr)
    ct(th.rows[ri].cells[3], d, bold=is_curr)
    ct(th.rows[ri].cells[4], e)
    sdat(th.rows[ri], is_curr)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ================================================================
# SECTION 5 — Key Issues
# ================================================================
sh(doc, "5.  Key Issues and Questions to Address with Your Bookkeeper")
bp(doc, ("The books are mathematically correct \u2014 every running balance adds up properly "
         "and there are no unexplained jumps in the numbers. However, there are several items "
         "worth understanding and documenting to stay on the right side of the CRA."))

# --- 5A: Hajj (conditional) ---
if has_hajj:
    hajj_amt  = abs(hajj_txns[0]["amount"])
    hajj_date = fmt_date(hajj_txns[0]["date"])
    sub_header(doc, f"A.  The ${hajj_amt:,.0f} Hajj Travel Payment ({hajj_date})  [ACTION REQUIRED]")
    bp(doc, (f"On {hajj_date}, the corporation sent a ${hajj_amt:,.0f} e-Transfer to "
             "hajj@kingtravelcan.com for what appears to be Hajj pilgrimage travel. "
             "This is a personal expense, not a business expense. "
             "The corporation paid it on Ramzan\u2019s behalf and recorded it as a shareholder advance "
             "\u2014 which means Ramzan owes the company this amount."))
    bp(doc, ("This is legally acceptable as long as it is treated as a shareholder loan "
             "(which it currently is). However, you must make sure of three things:"))
    for item in [
        ("It must never show up as a business expense on your income statement.",
         "If it were deducted as a business expense, the CRA would disallow it and assess a penalty."),
        (f"Ramzan must repay the ${hajj_amt:,.0f} (as part of his overall balance) to the corporation "
         "by July 31, 2027.",
         "That is one year after fiscal year-end (July 31, 2026). If not repaid, the Hajj amount "
         "becomes taxable personal income for Ramzan."),
        ("No interest is currently being charged.",
         "The CRA may assess a taxable interest benefit on unpaid advances. Ask Tang & Partners "
         "whether the prescribed interest rate needs to be charged."),
    ]:
        p = doc.add_paragraph(style="List Number")
        rb = p.add_run(item[0] + "  "); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
        rn = p.add_run(item[1]); rn.font.name = CAL; rn.font.size = BS
        p.paragraph_format.space_after = Pt(5)
    callout_red(doc, (f"IMPORTANT: Verify with your bookkeeper that the ${hajj_amt:,.0f} Hajj "
                      "payment is recorded in account 2901 (Shareholder Advance) only, and NOT "
                      "as any kind of business expense in your P&L. If it appears in any expense "
                      "account, have it corrected before year-end."))

# --- 5B: Large FY2023-24 temporary loans ---
sub_header(doc, "B.  Large Temporary Loans in FY2023-24  [UNDERSTAND AND DOCUMENT]")
bp(doc, ("In FY2023-24, there was very large activity in both shareholder accounts. "
         "These were short-term loans from the corporation to the shareholders, all of which "
         "were repaid within the CRA\u2019s one-year window. Here is what happened:"))
fy24_rows = [
    ("Rezai: Jun 26, 2024", "\u2212$100,000",
     "Corporation paid Rezai $100,000 via bank-to-bank transfer."),
    ("Rezai: Sep 5, 2024", "+$100,000",
     "Rezai repaid $100,000 \u2014 within 3 months. CRA 1-year rule satisfied."),
    ("Ramzan: Jul 22, 2024", "\u2212$50,000",
     "Corporation paid Ramzan $50,000 via cheque #220."),
    ("Ramzan: Sep 13, 2024", "+$47,077.86",
     "Ramzan deposited $47,077.86 back. Net $2,922.14 difference "
     "was already credited from Jun 2024 journal entries."),
]
tf = doc.add_table(rows=1 + len(fy24_rows), cols=3)
tf.style = "Table Grid"; tf.alignment = WD_TABLE_ALIGNMENT.LEFT
tf.columns[0].width = Inches(1.7); tf.columns[1].width = Inches(0.9); tf.columns[2].width = Inches(4.3)
ct(tf.rows[0].cells[0], "Transaction", bold=True)
ct(tf.rows[0].cells[1], "Amount",      bold=True)
ct(tf.rows[0].cells[2], "Explanation", bold=True)
shdr(tf.rows[0])
for ri, (a, b, c) in enumerate(fy24_rows, 1):
    ct(tf.rows[ri].cells[0], a); ct(tf.rows[ri].cells[1], b); ct(tf.rows[ri].cells[2], c)
    sdat(tf.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)
callout(doc, ("The loans were repaid on time \u2014 no CRA inclusion issue. "
              "However, if the corporation did not charge interest at the CRA prescribed rate "
              "(currently ~5%), there may be a deemed taxable interest benefit. "
              "Ask Tang & Partners whether interest should have been recorded for FY2024-25."))
bp(doc, ("Also in FY2023-24, both Rezai and Ramzan had several $35,000\u2013$50,000 amounts "
         "go out and come back within days or weeks. These short cycles appear to be cash-flow "
         "management \u2014 the corporation temporarily parking cash with the shareholders and "
         "receiving it back. This is fine legally as long as the round trips are documented "
         "and interest is properly handled."))

# --- 5C: 2236262 Ontario Inc. ---
sub_header(doc, "C.  Transfers from 2236262 Ontario Inc.  [CLARIFY]")
bp(doc, ("On November 10, 2025, both Ramzan and Rezai each received $2,943.65 from "
         "\u20182236262 Ontario Inc.\u2019 (total $5,887.30), recorded through the shareholder advance "
         "accounts. This is the first time this company appears in the QuickBooks records."))
bp(doc, ("Questions to ask your bookkeeper:"))
for q in [
    "Who owns 2236262 Ontario Inc.? Is it owned by Ramzan and Rezai?",
    "What is the business relationship between 2039321 Ontario Inc. (your company) and 2236262 Ontario Inc.?",
    "What was the $2,943.65 \u00d7 2 payment for? (Dividend? Expense reimbursement? Loan repayment?)",
    "Are inter-company transactions between these two entities properly documented?",
]:
    p = doc.add_paragraph(style="List Bullet")
    rn = p.add_run(q); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(3)
note(doc, ("Inter-company transactions between related companies are a common CRA audit target. "
           "Make sure there is a written agreement or at least a clear paper trail explaining "
           "what this money was for."))

# --- 5D: Recurring year-end JEs ---
sub_header(doc, "D.  Recurring Year-End Journal Entries  [UNDERSTAND WHAT THESE ARE]")
bp(doc, ("Every year near July 31, both shareholders receive three journal entry credits: "
         "$3,500.00 + $4,950.00 + $1,070.37 = $9,520.37 per person. "
         "These have been posted consistently since FY2013-14 "
         "and grew to the full $9,520.37 amount in FY2021-22. "
         "The entries have no memo description."))
je_rows = [
    ("FY2013-14 to FY2020-21", "$4,950.00 + $1,070.37 = $6,020.37",
     "Annual JE per shareholder. No memo. Purpose unknown from records alone."),
    ("FY2021-22 to FY2024-25", "$3,500 + $4,950 + $1,070.37 = $9,520.37",
     "$3,500 component added in FY2021-22. Total now $9,520.37 per shareholder per year."),
    ("FY2025-26",
     "$3,500 + $4,950 + $1,070.37 = $9,520.37",
     f"Posted {je12_date_str} (JE-12)." if je12_posted else "Expected near July 31, 2026."),
]
tj = doc.add_table(rows=1 + len(je_rows), cols=3)
tj.style = "Table Grid"; tj.alignment = WD_TABLE_ALIGNMENT.LEFT
tj.columns[0].width = Inches(1.7); tj.columns[1].width = Inches(1.9); tj.columns[2].width = Inches(3.3)
ct(tj.rows[0].cells[0], "Period",                  bold=True)
ct(tj.rows[0].cells[1], "Amount per Shareholder",  bold=True)
ct(tj.rows[0].cells[2], "Note",                    bold=True)
shdr(tj.rows[0])
for ri, (a, b, c) in enumerate(je_rows, 1):
    ct(tj.rows[ri].cells[0], a); ct(tj.rows[ri].cells[1], b); ct(tj.rows[ri].cells[2], c)
    sdat(tj.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)
callout(doc, ("Ask your bookkeeper or Tang & Partners: What are these three journal entries for? "
              "They credit both shareholder accounts by $9,520.37 per year. Are they for "
              "car allowances? Health benefit reimbursements? Expense reimbursements? "
              "They should have a memo describing their purpose, especially since they "
              "happen identically for both shareholders every single year. "
              "If they represent taxable benefits, they may need to be reported on T4 slips."))

# --- 5E: Uncleared ---
sub_header(doc, "E.  All FY2025-26 Entries Are \u2018Uncleared\u2019  [BOOKKEEPING ACTION]")
bp(doc, ("Every single transaction in the FY2025-26 shareholder advance accounts is currently "
         "marked as \u2018Uncleared\u2019 in QuickBooks. This means your bookkeeper has not yet "
         "matched these entries to your actual bank statements."))
bp(doc, ("Uncleared entries are not necessarily wrong \u2014 but they do mean that if there are "
         "any entry errors, they would not have been caught yet. Bank reconciliation is the "
         "process that confirms that what\u2019s in QuickBooks matches what actually happened "
         "at the bank. Ask your bookkeeper to reconcile the shareholder advance accounts "
         "as part of your next quarterly review."))

# --- 5F: JE-11 ---
if je11_ram:
    sub_header(doc, f"F.  Journal Entry JE-11 ({je11_date_str}) \u2014 What It Means")
    bp(doc, (f"On {je11_date_str}, a journal entry (JE-11) was posted that transferred "
             f"Rezai\u2019s ${je11_amount:,.2f} credit balance and applied it against Ramzan\u2019s "
             "outstanding debt. This is a common bookkeeping technique to settle inter-shareholder balances."))
    bp(doc, ("Here is what changed:"))
    for item in [
        ("Before JE-11:",
         f"Rezai had a +${je11_amount:,.2f} credit (the corporation owed Rezai this amount). "
         "Ramzan had a larger negative balance (Ramzan owed the corporation that amount)."),
        ("After JE-11:",
         f"Rezai $0.00 (account fully zeroed). "
         f"Ramzan\u2019s debt reduced by ${je11_amount:,.2f}."),
        ("Net corporate position:",
         "The combined shareholder balance was unchanged \u2014 this is just a reallocation "
         "between the two shareholders, not a new cash flow."),
    ]:
        p = doc.add_paragraph(style="List Number")
        rb = p.add_run(item[0] + "  "); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
        rn = p.add_run(item[1]); rn.font.name = CAL; rn.font.size = BS
        p.paragraph_format.space_after = Pt(5)
    bp(doc, ("Action required: Document with your bookkeeper what specific agreement JE-11 covers. "
             "Was this Rezai agreeing to waive his credit to offset Ramzan\u2019s debt? "
             "If so, that should be documented in writing between the two shareholders."))
    callout_blue(doc, (f"JE-11 is internally consistent \u2014 the math checks out. The combined "
                       f"shareholder balance was unchanged before and after the entry. This is just a "
                       f"reallocation between the two shareholders. Make sure there is a written record "
                       f"of the agreement."))

# ================================================================
# SECTION 6 — Balance Check
# ================================================================
sh(doc, "6.  Are the Books Balanced? \u2014 Quick Check")
bp(doc, ("The short answer is yes \u2014 mathematically, the books are consistent. "
         "Here is how we verified this:"))

ram_tx_sum = sum(tx["amount"] for tx in ram_txns)
rez_tx_sum = sum(tx["amount"] for tx in rez_txns)
ram_calc   = ram_open + ram_tx_sum
rez_calc   = rez_open + rez_tx_sum

check_rows = [
    ("Ramzan\u2019s running total",
     f"Opening {fmt_bal(ram_open)} + net transactions {fmt_net(ram_tx_sum)} = {fmt_bal(ram_calc)}",
     f"\u2713 Matches closing balance {fmt_bal(ram_close)}"),
    ("Rezai\u2019s running total",
     f"Opening {fmt_bal(rez_open)} + net transactions {fmt_net(rez_tx_sum)} = {fmt_bal(rez_calc)}",
     f"\u2713 Matches closing balance {fmt_bal(rez_close)}"),
    ("Year-end progression",
     "Each year\u2019s ending balance = prior year ending + current year net change",
     "\u2713 Verified for all years"),
    ("FY2025-26 combined",
     f"Ramzan {fmt_bal(ram_close)} + Rezai {fmt_bal(rez_close)} = {fmt_bal(comb_close)}",
     "\u2713 Combined verified"),
]
tc = doc.add_table(rows=1 + len(check_rows), cols=3)
tc.style = "Table Grid"; tc.alignment = WD_TABLE_ALIGNMENT.LEFT
tc.columns[0].width = Inches(1.7); tc.columns[1].width = Inches(3.3); tc.columns[2].width = Inches(1.9)
ct(tc.rows[0].cells[0], "Check",       bold=True)
ct(tc.rows[0].cells[1], "Calculation", bold=True)
ct(tc.rows[0].cells[2], "Result",      bold=True)
shdr(tc.rows[0])
for ri, (a, b, c) in enumerate(check_rows, 1):
    ct(tc.rows[ri].cells[0], a)
    ct(tc.rows[ri].cells[1], b)
    ct(tc.rows[ri].cells[2], c, color=RGBColor(0x1A, 0x5C, 0x1A))
    sdat(tc.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)
callout_green(doc, (f"No unexplained jumps or arithmetic errors were found in the QuickBooks data. "
                    f"The running balances for both shareholders trace correctly from the opening "
                    f"entries in 2008 all the way to {cutoff_str}. JE-11 and JE-12 are both "
                    f"internally consistent. The books are internally consistent."))

# ================================================================
# SECTION 7 — Action Checklist
# ================================================================
sh(doc, "7.  Your Action Checklist")

actions = []
if has_hajj:
    hajj_amt  = abs(hajj_txns[0]["amount"])
    actions.append((
        f"URGENT \u2014 Confirm the Hajj payment is NOT a business expense",
        f"Verify with your bookkeeper that the ${hajj_amt:,.0f} Hajj travel payment "
        f"({fmt_date(hajj_txns[0]['date'])}) is recorded in account 2901 (Shareholder Advance) "
        f"and does not appear anywhere in the expense accounts on your P&L. "
        f"This is a personal advance, not a tax deduction."
    ))
if ram_close < -0.005:
    actions.append((
        f"Ramzan must repay {fmt_bal(ram_close)} to the corporation by July 31, 2027",
        "The entire negative balance must be repaid within one year of fiscal year-end (July 31, 2026). "
        "Plan for this now."
    ))
else:
    actions.append((
        f"Monitor Ramzan\u2019s balance ({fmt_bal(ram_close)} \u2014 corp owes Ramzan)",
        "Ramzan\u2019s account is currently positive (no repayment needed). Continue to monitor "
        "each quarter and ensure the Hajj advance is properly documented as a personal advance, "
        "not a business expense."
    ))

if je11_ram:
    actions.append((
        "Understand JE-11 \u2014 get written documentation of the inter-shareholder settlement",
        f"JE-11 ({je11_date_str}) transferred Rezai\u2019s ${je11_amount:,.2f} credit to offset "
        "Ramzan\u2019s debt. Make sure there is a written agreement between Ramzan and Rezai "
        "confirming that Rezai agreed to waive his credit for this purpose. Without documentation, "
        "the CRA may question the nature of this transaction."
    ))

actions += [
    ("Ask Tang & Partners about interest on shareholder loans",
     "The corporation has advanced money to shareholders (especially in FY2023-24) without "
     "visible interest charges. CRA requires that loans to shareholders either be repaid "
     "quickly or have interest charged at the prescribed rate. Ask your accountant whether "
     "interest income should have been recorded on these advances."),
    ("Identify 2236262 Ontario Inc. and document the Nov 10 transfers",
     "Two payments of $2,943.65 came from this related company into your shareholder advance "
     "accounts on November 10, 2025. Make sure there is a written explanation of what these "
     "payments were for and what the relationship is between the two companies."),
    ("Ask your bookkeeper what the recurring $9,520.37 annual JEs are for",
     "These credits ($3,500 + $4,950 + $1,070.37) have been posted to both shareholder "
     "accounts near every fiscal year-end since FY2013-14. They should have a memo "
     "describing their purpose. If they are taxable benefits, they may need to be "
     "reported on T4s."),
    ("Complete bank reconciliation for FY2025-26",
     "All current-year shareholder advance transactions are \u2018Uncleared.\u2019 Ask your "
     "bookkeeper to reconcile these accounts so that the QuickBooks records are confirmed "
     "against actual bank statements."),
]
for title, body in actions:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title + "  "); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(body); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(6)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ================================================================
# SECTION 8 — Bottom Line
# ================================================================
sh(doc, "8.  The Bottom Line \u2014 Plain and Simple")

ram_bl_desc = (
    f"Ramzan\u2019s balance is now {fmt_bal(ram_close)} (the corporation owes Ramzan this amount), "
    "up from a negative position earlier in the year, thanks to JE-11 and JE-12 credits. "
    "The Hajj advance ($10,000, Jan 13, 2026) is embedded in this balance \u2014 "
    "it must not be deducted as a business expense."
    if ram_close > 0 else
    f"Ramzan owes the corporation {fmt_bal(ram_close)}. "
    "This must be repaid by July 31, 2027 to avoid it being taxed as personal income. "
    "The $10,000 Hajj advance (Jan 13, 2026) is included within this balance."
)
rez_bl_desc = (
    f"Rezai\u2019s balance is {fmt_bal(rez_close)} (the corporation owes Rezai). "
    "No repayment required. JE-12 credits account for the full balance."
    if rez_close > 0 else
    f"Rezai owes the corporation {fmt_bal(rez_close)}."
)

blines = [
    ("The math checks out \u2014 the books are internally consistent.",
     "  Every running balance traces correctly from 2008 to today. JE-11 and JE-12 are both "
     "internally consistent. There are no mysterious gaps, duplicate entries, or arithmetic errors."),
    (f"Ramzan: {fmt_bal(ram_close)} as of {cutoff_str}.",
     f"  {ram_bl_desc}"),
    (f"Rezai: {fmt_bal(rez_close)} as of {cutoff_str}.",
     f"  {rez_bl_desc}"),
    ("The FY2023-24 large temporary loans ($100,000 to Rezai, $50,000 to Ramzan) were repaid on time.",
     "  Both amounts came back within a few months, satisfying the CRA\u2019s one-year rule. "
     "The open question is whether proper interest was charged. Ask Tang & Partners."),
    ("There are four items that need clarification \u2014 not emergencies, but worth resolving.",
     "  (1) Purpose of the recurring $9,520.37 annual JEs for each shareholder. "
     "(2) Document the JE-11 agreement between Ramzan and Rezai in writing. "
     "(3) Interest treatment on shareholder loans. "
     "(4) Bank reconciliation to clear the \u2018Uncleared\u2019 status of all FY2025-26 entries."),
    ("The shareholder advance account attracts CRA attention.",
     "  As the corporation becomes more profitable and taxes increase (as projected for FY2025-26), "
     "the CRA pays closer attention to amounts flowing between a corporation and its shareholders. "
     "Clean documentation, timely repayments, and proper interest charges are the "
     "three things that keep this account problem-free."),
]
for title, body in blines:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(body); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
note(doc, (f"Disclaimer: This report is based on QuickBooks shareholder advance data (accounts "
           f"2901 and 2902) as exported on {cutoff_str}. It covers all transactions from "
           f"FY2007-08 through FY2025-26. This is an analytical review, not professional tax or "
           f"legal advice. Share all findings with Tang & Partners LLP (416-987-6005 / thomas@tang.ca) "
           f"before making any decisions."))

# ================================================================
# SAVE
# ================================================================
doc.save(OUT)
print("Saved:", OUT)

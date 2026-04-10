"""
Task #1 — Marketing & Advertising Spend Analysis
Spirit of Math Schools Markham East (2039321 Ontario Inc.)
All figures read dynamically from data/extracted/run_data.json.
"""

import sys, os, json, datetime
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import (
    HB, CAL, BS, make_doc,
    sbg, sbd, shdr, sdat, ct, hr, bp, sh, note,
    callout, callout_red_bullets, callout_blue_bullets,
)

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ------------------------------------------------------------------ #
# Load run data
# ------------------------------------------------------------------ #
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(BASE_DIR, "data", "extracted", "run_data.json"), encoding="utf-8") as f:
    D = json.load(f)

# ------------------------------------------------------------------ #
# Date / period calculations
# ------------------------------------------------------------------ #
TODAY       = datetime.date.today()
REPORT_DATE = TODAY.strftime("%B %d, %Y")
FILE_DATE   = TODAY.strftime("%Y-%m-%d")

# Parse ytd_cutoff_date safely (handles "April 2, 2026" or "Apr. 2 2026")
raw_date = D["meta"]["ytd_cutoff_date"]
for fmt in ("%B %d, %Y", "%B %d %Y", "%b. %d, %Y", "%b. %d %Y", "%b %d, %Y", "%b %d %Y"):
    try:
        cutoff_dt = datetime.datetime.strptime(raw_date, fmt).date()
        break
    except ValueError:
        pass

# Format cutoff without leading zero (Windows-safe)
cutoff_month = cutoff_dt.strftime("%B")
cutoff_day   = str(cutoff_dt.day)
cutoff_year  = str(cutoff_dt.year)
cutoff_str   = f"{cutoff_month} {cutoff_day}, {cutoff_year}"   # e.g. "April 2, 2026"

FY_START = datetime.date(2025, 8, 1)
FY_END   = datetime.date(2026, 7, 31)

# Months elapsed (fractional) from Aug 1 to cutoff
elapsed_days     = (cutoff_dt - FY_START).days
months_elapsed   = elapsed_days / 30.44
months_remaining = (FY_END - cutoff_dt).days / 30.44

# ------------------------------------------------------------------ #
# Revenue figures
# ------------------------------------------------------------------ #
rev              = D["revenue"]
ytd_cur          = rev["ytd_tuition_current"]
ytd_py           = rev["ytd_tuition_prior_year"]
full_year_prior  = rev["full_year_prior"]          # 3,020,723 (audited)
ytd_ratio_pct    = rev["ytd_to_annual_ratio"] * 100
projected_rev    = rev["projected_full_year"]
yoy_pct          = rev["yoy_growth_pct"]

# ------------------------------------------------------------------ #
# Marketing figures
# ------------------------------------------------------------------ #
mkt              = D["marketing"]
total_ytd_cur    = mkt["total_ytd_current"]
total_ytd_py     = mkt["total_ytd_prior"]
ob_cons          = mkt["obligation_conservative"]
ob_proj          = mkt["obligation_projected"]
gap_cons         = mkt["gap_conservative"]
gap_proj         = mkt["gap_projected"]
monthly_cons     = gap_cons / months_remaining if months_remaining > 0 else 0
monthly_proj     = gap_proj / months_remaining if months_remaining > 0 else 0

# ------------------------------------------------------------------ #
# Benchmark constants (3-year historical — fixed until new P&L arrives)
# ------------------------------------------------------------------ #
bmark_mkt_avg = D["benchmarks_3yr"]["marketing_annual_avg"]
bmark_ftc_avg = D["benchmarks_3yr"]["ftc_annual_avg"]
bmark_3yr_rev = D["benchmarks_3yr"]["tuition_annual_avg"]

# Per-account 3-year annual averages (fixed from Aug 2022–Jul 2025 P&L)
BMARK_ACCT = {
    "6200":   7462,
    "62010":  19199,
    "6201.1": 24058,
    "6201":   6159,
    "6202":   5108,
    "6205":   17688,
    "6208":   2289,
    "6209":   2764,
}

# ------------------------------------------------------------------ #
# Per-account marketing dict (from JSON)
# ------------------------------------------------------------------ #
accounts = mkt["accounts"]

def acct_cur(key):
    for k, v in accounts.items():
        if k.startswith(key.split()[0]):
            return v["current"]
    return 0.0

def acct_py(key):
    for k, v in accounts.items():
        if k.startswith(key.split()[0]):
            return v["prior"]
    return 0.0

def fmt_dollar(v):
    """Format a dollar value. Negative = prepend minus sign."""
    if v < 0:
        return f"-${abs(v):,.0f}"
    return f"${v:,.0f}"

def fmt_dollar_yr(v):
    return f"${v:,.0f}/yr"

# ------------------------------------------------------------------ #
# Dynamic status for Section 6 benchmark table
# ------------------------------------------------------------------ #
def bmark_status(acct_code, cur_val):
    avg = BMARK_ACCT.get(acct_code, 0)
    if acct_code == "6201.1":
        if cur_val == 0:
            return f"\u26a0 MISSING \u2014 $0 this year vs ${avg:,.0f} avg"
        return f"${cur_val:,.0f} booked (avg ${avg:,.0f}/yr)"
    if acct_code == "62010":
        if cur_val == 0:
            return "Not started this year"
        return f"Already ${cur_val:,.0f} \u2014 {'above' if cur_val > avg else 'below'} 3-yr average"
    if acct_code == "6209":
        if cur_val == 0:
            return "Not yet started this year"
        return f"${cur_val:,.0f} booked this year"
    if cur_val == 0:
        return "No spend yet this year"
    if cur_val >= avg:
        return f"On track (${cur_val:,.0f} already, avg ${avg:,.0f}/yr)"
    # behind
    pct = (cur_val / avg * 100) if avg else 0
    return f"Behind \u2014 ${cur_val:,.0f} with {months_remaining:.1f} months left"

# ================================================================
# BUILD DOCUMENT
# ================================================================
doc = make_doc()

# ------------------------------------------------------------------ #
# TITLE
# ------------------------------------------------------------------ #
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Marketing & Advertising Spend Analysis")
r.bold = True; r.font.name = CAL; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("Spirit of Math Schools Markham East  \u2014  2039321 Ontario Inc.")
r.italic = True; r.font.name = CAL; r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Report Date: {REPORT_DATE}   |   Fiscal Year: August 1, 2025 \u2013 July 31, 2026")
r.italic = True; r.font.name = CAL; r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"YTD data through: {cutoff_str}")
r.italic = True; r.font.name = CAL; r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(6)
hr(doc)

# ------------------------------------------------------------------ #
# QUICK SUMMARY
# ------------------------------------------------------------------ #
sh(doc, "Quick Summary")

ftc_cur = 0.0
for k, v in accounts.items():
    if "6201.1" in k or "FTC" in k:
        ftc_cur = v["current"]
        ftc_py  = v["prior"]
        break
else:
    ftc_py = 0.0

ftc_missing = (ftc_cur == 0)

blue_bullets = [
    f"This report covers FY2025-26 marketing spend from August 1 to {cutoff_str} "
    f"({months_elapsed:.1f} months in).",
    f"You have spent {fmt_dollar(total_ytd_cur)} on marketing so far. "
    f"Your full-year obligation is {fmt_dollar(ob_cons)}\u2013{fmt_dollar(ob_proj)}.",
    f"You need to spend roughly {fmt_dollar(gap_cons)}\u2013{fmt_dollar(gap_proj)} more before July 31, 2026.",
    f"Revenue is growing: {fmt_dollar(ytd_cur)} YTD vs {fmt_dollar(ytd_py)} same period last year "
    f"(+{yoy_pct:.1f}%).",
]
if ftc_missing:
    blue_bullets.append(
        f"FTC charges (account 6201.1) are still $0 this year \u2014 "
        f"were {fmt_dollar(ftc_py)} at this point last year. Still unexplained."
    )

mkt_wages_cur = 0.0
for k, v in accounts.items():
    if "62010" in k or "Marketing Wages" in k:
        mkt_wages_cur = v["current"]
        break
if mkt_wages_cur > 0:
    blue_bullets.append(
        f"Marketing Wages (62010) of {fmt_dollar(mkt_wages_cur)} are new this year \u2014 "
        "this is the single largest marketing item."
    )

callout_blue_bullets(doc, blue_bullets)

red_bullets = []
if ftc_missing:
    red_bullets.append(
        f"URGENT: FTC charges are $0. Last year you paid {fmt_dollar(ftc_py)} to Spirit of Math\u2019s "
        "marketing fund by this date. Confirm with head office whether this is still owed \u2014 "
        "booking it closes a significant portion of the conservative gap."
    )
red_bullets.append(
    f"You need to spend ~{fmt_dollar(monthly_cons)}\u2013{fmt_dollar(monthly_proj)} per month on marketing "
    f"for the remaining {months_remaining:.1f} months to meet your obligation."
)
callout_red_bullets(doc, red_bullets)

note(doc, "See the Action Checklist (Section 8) and Bottom Line (Section 9) at the end of this report.")

# ------------------------------------------------------------------ #
# SECTION 1 — The Obligation
# ------------------------------------------------------------------ #
sh(doc, "1.  The Obligation")
bp(doc, ("The franchise agreement requires you to spend a minimum of 3% of your gross revenue "
         "on local marketing each fiscal year. This is not optional \u2014 it is a contract requirement. "
         "Every dollar you spend on marketing also reduces your taxable income, so there is a "
         "tax benefit built into this obligation."))

# ------------------------------------------------------------------ #
# SECTION 2 — Revenue Projection
# ------------------------------------------------------------------ #
sh(doc, "2.  Step 1 \u2014 Estimate Full-Year Gross Revenue")
bp(doc, (f"We use the prior year to estimate how much revenue you will earn this full year. "
         f"In FY2024-25, you had earned {ytd_ratio_pct:.1f}% of your full-year revenue by {cutoff_str}. "
         f"Applying that same ratio to this year\u2019s YTD gives us the full-year projection."))

s2 = [
    ("FY2024-25 full-year tuition (confirmed from audited financials)",
     fmt_dollar(full_year_prior), False),
    (f"FY2024-25 YTD tuition (Aug 1 \u2013 {cutoff_str})",
     fmt_dollar(ytd_py), False),
    ("YTD-to-annual ratio",
     f"{ytd_ratio_pct:.1f}%", False),
    (f"Current YTD tuition (Aug 1 \u2013 {cutoff_str})",
     fmt_dollar(ytd_cur), False),
    ("Projected FY2025-26 full-year revenue",
     f"~{fmt_dollar(projected_rev)}", True),
]
t = doc.add_table(rows=1 + len(s2), cols=2)
t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
t.columns[0].width = Inches(4.0); t.columns[1].width = Inches(2.0)
ct(t.rows[0].cells[0], "Item", bold=True); ct(t.rows[0].cells[1], "Amount", bold=True)
shdr(t.rows[0])
for i, (a, b, bl) in enumerate(s2, 1):
    ct(t.rows[i].cells[0], a, bold=bl); ct(t.rows[i].cells[1], b, bold=bl)
    sdat(t.rows[i], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

diff_rev = ytd_cur - ytd_py
callout(doc, (f"Revenue is growing strongly. You earned {fmt_dollar(ytd_cur)} in the first "
              f"{months_elapsed:.1f} months \u2014 that is {fmt_dollar(diff_rev)} more than the same period "
              f"last year (+{yoy_pct:.1f}%). This is great for the business, but it also means your "
              "3% marketing obligation is larger this year than last."))

# ------------------------------------------------------------------ #
# SECTION 3 — 3% Obligation
# ------------------------------------------------------------------ #
sh(doc, "3.  Step 2 \u2014 Required Minimum Spend (3% Obligation)")
s3 = [
    ("Conservative (background doc estimate)",
     fmt_dollar(ob_cons / 0.03), fmt_dollar(ob_cons), False),
    ("Projected (YTD extrapolation)",
     f"~{fmt_dollar(projected_rev)}", fmt_dollar(ob_proj), False),
    ("3-year historical average revenue",
     fmt_dollar(bmark_3yr_rev), fmt_dollar(bmark_mkt_avg), False),
]
t3 = doc.add_table(rows=1 + len(s3), cols=3)
t3.style = "Table Grid"; t3.alignment = WD_TABLE_ALIGNMENT.LEFT
t3.columns[0].width = Inches(2.8); t3.columns[1].width = Inches(1.7); t3.columns[2].width = Inches(1.5)
ct(t3.rows[0].cells[0], "Scenario", bold=True)
ct(t3.rows[0].cells[1], "Revenue Basis", bold=True)
ct(t3.rows[0].cells[2], "3% Obligation", bold=True)
shdr(t3.rows[0])
for i, (a, b, c, bl) in enumerate(s3, 1):
    ct(t3.rows[i].cells[0], a, bold=bl); ct(t3.rows[i].cells[1], b, bold=bl)
    ct(t3.rows[i].cells[2], c, bold=bl); sdat(t3.rows[i], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

bp(doc, (f"Recommendation: Budget to meet the higher figure (~{fmt_dollar(ob_proj)}) to avoid falling short. "
         f"This year\u2019s obligation is about {fmt_dollar(ob_proj - bmark_mkt_avg)} above your 3-year average "
         "because your revenue is growing."), bold=True)

# ------------------------------------------------------------------ #
# SECTION 4 — YTD Spend
# ------------------------------------------------------------------ #
sh(doc, f"4.  Step 3 \u2014 Marketing Spend Year-to-Date (Aug 1 \u2013 {cutoff_str})")
bp(doc, ("The table below shows what you spent in each marketing category this year vs the same "
         "period last year. These figures come directly from QuickBooks."))

# Build rows dynamically from accounts dict
s4_rows = []
for acct_name, vals in accounts.items():
    s4_rows.append((acct_name, vals["current"], vals["prior"]))
# Add total row
s4_rows.append(("Total YTD Marketing Spent", total_ytd_cur, total_ytd_py))

t4 = doc.add_table(rows=1 + len(s4_rows), cols=3)
t4.style = "Table Grid"; t4.alignment = WD_TABLE_ALIGNMENT.LEFT
t4.columns[0].width = Inches(2.8); t4.columns[1].width = Inches(1.5); t4.columns[2].width = Inches(1.5)
ct(t4.rows[0].cells[0], "Account", bold=True)
ct(t4.rows[0].cells[1], f"FY2025-26 YTD\n(Aug 1\u2013{cutoff_str})", bold=True)
ct(t4.rows[0].cells[2], f"FY2024-25 YTD\n(Aug 1\u2013{cutoff_month} {cutoff_day}, 2025)", bold=True)
shdr(t4.rows[0])
for i, (acct, cur, py) in enumerate(s4_rows, 1):
    bl = ("Total" in acct)
    ct(t4.rows[i].cells[0], acct, bold=bl)
    ct(t4.rows[i].cells[1], fmt_dollar(cur), bold=bl)
    ct(t4.rows[i].cells[2], fmt_dollar(py), bold=bl)
    sdat(t4.rows[i], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Find notable accounts for prose
online_cur = online_py = 0.0
for k, v in accounts.items():
    if "6208" in k:
        online_cur = v["current"]; online_py = v["prior"]

py_neg_6200 = 0.0
for k, v in accounts.items():
    if k.startswith("6200 ") and v["prior"] < 0:
        py_neg_6200 = v["prior"]

if py_neg_6200 < 0:
    note(doc, (f"Prior year {fmt_dollar(py_neg_6200)} in account 6200 reflects a credit or refund on a "
               "prior advertising booking. "
               f"Prior year total of {fmt_dollar(total_ytd_py)} includes {fmt_dollar(ftc_py)} in FTC charges "
               "that are absent this year." if ftc_missing else
               f"Prior year total of {fmt_dollar(total_ytd_py)}."))

if mkt_wages_cur > 0 and online_cur > 0 and online_py > 0:
    online_pct = (online_cur / online_py - 1) * 100 if online_py != 0 else 0
    py_own = total_ytd_py - ftc_py
    own_pct = (total_ytd_cur / py_own - 1) * 100 if py_own != 0 else 0
    callout(doc, (f"The biggest item this year is Marketing Wages at {fmt_dollar(mkt_wages_cur)} \u2014 "
                  f"this account was $0 last year. "
                  f"Excluding FTC, your own marketing spend of {fmt_dollar(total_ytd_cur)} is well ahead of "
                  f"last year\u2019s own spending ({fmt_dollar(total_ytd_py)} \u2212 {fmt_dollar(ftc_py)} "
                  f"= {fmt_dollar(py_own)}) \u2014 a {own_pct:.0f}% increase. "
                  f"Online advertising is the standout: up from {fmt_dollar(online_py)} to "
                  f"{fmt_dollar(online_cur)} (+{online_pct:.0f}%)."))

# ------------------------------------------------------------------ #
# SECTION 5 — Gap
# ------------------------------------------------------------------ #
sh(doc, "5.  Step 4 \u2014 How Much More Must Be Spent")
s5 = [
    ("Required minimum",
     fmt_dollar(ob_cons), fmt_dollar(ob_proj), False),
    (f"Already spent (Aug 1 \u2013 {cutoff_str})",
     fmt_dollar(total_ytd_cur), fmt_dollar(total_ytd_cur), False),
    ("Still required",
     fmt_dollar(gap_cons), fmt_dollar(gap_proj), True),
    (f"Months remaining ({cutoff_month} {cutoff_day} \u2013 Jul 31)",
     f"~{months_remaining:.1f} months", f"~{months_remaining:.1f} months", False),
    ("Monthly run rate needed",
     f"~{fmt_dollar(monthly_cons)}", f"~{fmt_dollar(monthly_proj)}", True),
]
t5 = doc.add_table(rows=1 + len(s5), cols=3)
t5.style = "Table Grid"; t5.alignment = WD_TABLE_ALIGNMENT.LEFT
t5.columns[0].width = Inches(3.0); t5.columns[1].width = Inches(1.5); t5.columns[2].width = Inches(1.5)
ct(t5.rows[0].cells[0], "", bold=True)
ct(t5.rows[0].cells[1], "Conservative", bold=True)
ct(t5.rows[0].cells[2], "Projected", bold=True)
shdr(t5.rows[0])
for i, (a, b, c, bl) in enumerate(s5, 1):
    ct(t5.rows[i].cells[0], a, bold=bl)
    ct(t5.rows[i].cells[1], b, bold=bl)
    ct(t5.rows[i].cells[2], c, bold=bl)
    sdat(t5.rows[i], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ------------------------------------------------------------------ #
# SECTION 6 — 3-Year Benchmark
# ------------------------------------------------------------------ #
sh(doc, "6.  Three-Year Marketing Benchmark (Aug 2022 \u2013 Jul 2025)")
bp(doc, ("This compares this year\u2019s spending pattern to the annual average over the past "
         "three full fiscal years."))

# Look up current-year value for each benchmark account
def cur_for_code(code):
    for k, v in accounts.items():
        acct_num = k.split()[0].rstrip(".")
        if acct_num == code or k.startswith(code + " ") or code in k:
            return v["current"]
    return 0.0

s6 = [
    ("6200  Marketing/Adv/Promo (misc)",
     fmt_dollar_yr(BMARK_ACCT["6200"]),
     bmark_status("6200", cur_for_code("6200"))),
    ("62010  Marketing Wages",
     fmt_dollar_yr(BMARK_ACCT["62010"]),
     bmark_status("62010", mkt_wages_cur)),
    ("6201.1  FTC charges",
     fmt_dollar_yr(BMARK_ACCT["6201.1"]),
     bmark_status("6201.1", ftc_cur)),
    ("6201  Marketing Material",
     fmt_dollar_yr(BMARK_ACCT["6201"]),
     bmark_status("6201", cur_for_code("6201 "))),
    ("6202  Space Rental",
     fmt_dollar_yr(BMARK_ACCT["6202"]),
     bmark_status("6202", cur_for_code("6202"))),
    ("6205  Public Sign",
     fmt_dollar_yr(BMARK_ACCT["6205"]),
     bmark_status("6205", cur_for_code("6205"))),
    ("6208  Online",
     fmt_dollar_yr(BMARK_ACCT["6208"]),
     bmark_status("6208", online_cur)),
    ("6209  Event Sponsorship",
     fmt_dollar_yr(BMARK_ACCT["6209"]),
     bmark_status("6209", cur_for_code("6209"))),
    ("All 6200-series combined",
     fmt_dollar_yr(int(bmark_mkt_avg)),
     f"This year needs {fmt_dollar(ob_cons)}\u2013{fmt_dollar(ob_proj)}"),
]
t6 = doc.add_table(rows=1 + len(s6), cols=3)
t6.style = "Table Grid"; t6.alignment = WD_TABLE_ALIGNMENT.LEFT
t6.columns[0].width = Inches(2.3); t6.columns[1].width = Inches(1.3); t6.columns[2].width = Inches(2.4)
ct(t6.rows[0].cells[0], "Account", bold=True)
ct(t6.rows[0].cells[1], "3-Yr Avg/Year", bold=True)
ct(t6.rows[0].cells[2], "Status This Year", bold=True)
shdr(t6.rows[0])
for i, (a, b, c) in enumerate(s6, 1):
    bl = ("All 6200" in a)
    ct(t6.rows[i].cells[0], a, bold=bl)
    ct(t6.rows[i].cells[1], b, bold=bl)
    ct(t6.rows[i].cells[2], c, bold=bl)
    sdat(t6.rows[i], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout(doc, (f"The 3-year average annual spend is {fmt_dollar(int(bmark_mkt_avg))}. "
              f"This year you need {fmt_dollar(ob_cons)}\u2013{fmt_dollar(ob_proj)}. "
              + (f"Marketing Wages are already above the historical average. " if mkt_wages_cur > BMARK_ACCT["62010"] else "")
              + (f"Online spending is {online_cur / BMARK_ACCT['6208']:.0f}x the average. " if BMARK_ACCT["6208"] > 0 else "")
              + ("The missing FTC and the underspend on public signage are the two gaps to close."
                 if ftc_missing else "Monitor FTC and public signage to stay on track.")))

# ------------------------------------------------------------------ #
# SECTION 7 — Key Observations
# ------------------------------------------------------------------ #
sh(doc, "7.  Key Observations")

obs = []

if ftc_missing:
    obs.append((
        "FTC charges are still missing. ",
        f"Account 6201.1 shows $0 for this reporting period. By this point last year, "
        f"you had paid {fmt_dollar(ftc_py)} to Spirit of Math\u2019s central marketing fund. "
        f"The 3-year average is {fmt_dollar(int(bmark_ftc_avg))}/year. If FTC is owed but not yet "
        "booked, recording it would significantly close the conservative gap. Confirm with head office."
    ))
else:
    obs.append((
        "FTC charges have been booked. ",
        f"Account 6201.1 shows {fmt_dollar(ftc_cur)} this year (vs {fmt_dollar(ftc_py)} last year, "
        f"3-yr avg {fmt_dollar(int(bmark_ftc_avg))}/yr). Monitor for any remaining balance."
    ))

if mkt_wages_cur > 0:
    obs.append((
        "Marketing Wages are a new major item. ",
        f"Account 62010 shows {fmt_dollar(mkt_wages_cur)} this year vs $0 last year. "
        f"This is the single biggest marketing category this year. "
        + (f"It already exceeds the 3-year average of {fmt_dollar(BMARK_ACCT['62010'])}/year. "
           if mkt_wages_cur > BMARK_ACCT["62010"] else "")
        + "Make sure these wages are for genuine marketing activities and are well-documented "
          "in case CRA asks."
    ))

if online_py > 0:
    online_pct2 = (online_cur / online_py - 1) * 100
    online_x    = online_cur / BMARK_ACCT["6208"] if BMARK_ACCT["6208"] > 0 else 0
    obs.append((
        "Online spending is a standout. ",
        f"Account 6208 (Online) went from {fmt_dollar(online_py)} last year to "
        f"{fmt_dollar(online_cur)} this year (+{online_pct2:.0f}%). "
        f"You have already invested {online_x:.0f}x the historical annual average in just "
        f"{months_elapsed:.0f} months. This is a smart shift toward digital marketing."
    ))

sign_cur = cur_for_code("6205")
if sign_cur < BMARK_ACCT["6205"]:
    py_sign = 0.0
    for k, v in accounts.items():
        if "6205" in k:
            py_sign = v["prior"]
    obs.append((
        "Public signage is behind last year. ",
        f"Account 6205 (Public Sign) is at {fmt_dollar(sign_cur)} vs "
        f"{fmt_dollar(py_sign)} last year and a 3-year average of "
        f"{fmt_dollar(BMARK_ACCT['6205'])}. With {months_remaining:.1f} months left, "
        "you can still catch up, but it needs attention."
    ))

event_cur = cur_for_code("6209")
if event_cur == 0:
    obs.append((
        "Event sponsorship has not started. ",
        f"Account 6209 averages {fmt_dollar(BMARK_ACCT['6209'])}/year historically. "
        "A local event sponsorship before July 31 would count toward the obligation and build "
        "community visibility."
    ))

obs.append((
    "Revenue growth increases the obligation. ",
    f"Revenue is projected at ~{fmt_dollar(projected_rev)} vs {fmt_dollar(full_year_prior)} last year "
    f"(+{yoy_pct:.1f}%). Your 3% obligation is roughly {fmt_dollar(ob_proj - ob_cons)} above the "
    "conservative estimate as a result."
))

for bpt, npt in obs:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(bpt); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(npt); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(5)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ------------------------------------------------------------------ #
# SECTION 8 — Action Checklist
# ------------------------------------------------------------------ #
sh(doc, "8.  Action Checklist")

actions = []
if ftc_missing:
    actions.append((
        "Confirm FTC status with Spirit of Math head office. ",
        "Find out if FTC invoices are still being issued. If yes, get the amount and have your "
        "bookkeeper record it in account 6201.1. This single item could close a significant portion "
        "of your gap."
    ))
if mkt_wages_cur > 0:
    actions.append((
        "Verify Marketing Wages are properly documented. ",
        f"Keep a contract or written record showing the {fmt_dollar(mkt_wages_cur)} in wages is for "
        "marketing work. CRA could ask about this in an audit."
    ))
actions.append((
    f"Plan {fmt_dollar(monthly_cons)}\u2013{fmt_dollar(monthly_proj)}/month in marketing for the next "
    f"{months_remaining:.1f} months. ",
    "This is the monthly spend needed to meet your obligation. Balance online ads "
    "with signage, events, and direct mail."
))
if sign_cur < BMARK_ACCT["6205"]:
    actions.append((
        "Boost public signage spending. ",
        f"At {fmt_dollar(sign_cur)}, signage is behind the {fmt_dollar(BMARK_ACCT['6205'])} annual average. "
        "Plan at least one more public sign campaign before July 31."
    ))
if event_cur == 0:
    actions.append((
        "Consider an event sponsorship. ",
        f"A local community event sponsorship counts toward the 3% obligation and builds "
        f"school visibility. Average annual spend historically is {fmt_dollar(BMARK_ACCT['6209'])}."
    ))

for bpt, npt in actions:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(bpt); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(npt); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(5)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ------------------------------------------------------------------ #
# SECTION 9 — Bottom Line
# ------------------------------------------------------------------ #
sh(doc, "9.  Bottom Line")

bottom = [
    (f"You need {fmt_dollar(gap_cons)}\u2013{fmt_dollar(gap_proj)} more in marketing spend before July 31. ",
     f"That works out to roughly {fmt_dollar(monthly_cons)}\u2013{fmt_dollar(monthly_proj)} per month "
     f"for the next {months_remaining:.1f} months. You are ahead of where you were last year at this point."),
]
if ftc_missing:
    gap_after_ftc_cons = max(0, gap_cons - bmark_ftc_avg)
    gap_after_ftc_proj = max(0, gap_proj - bmark_ftc_avg)
    bottom.append((
        "FTC is the biggest unknown. ",
        f"If FTC charges of ~{fmt_dollar(int(bmark_ftc_avg))} are still owed and get booked, your remaining "
        f"gap drops to roughly {fmt_dollar(gap_after_ftc_cons)}\u2013{fmt_dollar(gap_after_ftc_proj)} "
        "\u2014 very manageable. Resolve this first before planning other spending."
    ))
if mkt_wages_cur > 0:
    bottom.append((
        "Marketing Wages are now your top category. ",
        f"The {fmt_dollar(mkt_wages_cur)} in wages is the single largest marketing item this year. "
        "Make sure it continues and is well-supported by documentation."
    ))
bottom.append((
    "Online is your best-performing channel. ",
    "You have made a strong shift to digital. Balance it with traditional channels "
    "(signage, events) to cover all parts of your community."
))

for bpt, npt in bottom:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(bpt); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(npt); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(5)

note(doc, (f"Disclaimer: YTD data is from QuickBooks as of {cutoff_str}. "
           f"Full-year revenue projection uses the FY2024-25 YTD-to-annual ratio ({ytd_ratio_pct:.1f}%) "
           "derived from audited financial statements. "
           "3-year benchmarks are from the aggregate P&L (Aug 2022\u2013Jul 2025). "
           "All figures in Canadian dollars. Consult Tang & Partners before acting on any tax-related items."))

# ------------------------------------------------------------------ #
# SAVE
# ------------------------------------------------------------------ #
OUT = os.path.join(BASE_DIR, f"reports/claude_report_marketing_mae_{FILE_DATE}.docx")
doc.save(OUT)
print("Saved:", OUT)

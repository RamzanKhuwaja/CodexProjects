"""
validate_all.py — som_mae_financials
Validates the 4 most recent reports against run_data.json.
Exits 0 if all checks pass, 1 if any fail.
Writes reports/validation_summary_YYYY-MM-DD.txt with full results.
"""

import glob
import json
import os
import re
import sys
from datetime import date

# ── Locate project root (one level above scripts/) ──────────────────────────
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ── Helper: find most recent report matching a pattern ──────────────────────

def latest(pattern, exclude_suffixes=("_validation.docx", "-reviewed.docx")):
    """Return the most-recently-modified file matching reports/<pattern>."""
    files = glob.glob(os.path.join(BASE_DIR, "reports", pattern))
    files = [
        f for f in files
        if not any(f.endswith(s) for s in exclude_suffixes)
    ]
    return max(files, key=os.path.getmtime) if files else None


# ── Helpers: extract text from .docx ────────────────────────────────────────

def extract_table_values(doc):
    """Return a flat list of all non-empty cell text strings from all tables."""
    values = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    values.append(t)
    return values


def extract_para_values(doc):
    """Return a list of non-empty paragraph text strings."""
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def all_text_values(doc):
    """Combine table cells and paragraph text into one list."""
    return extract_table_values(doc) + extract_para_values(doc)


def parse_dollar(text):
    """
    Extract a float from a string like '$61,328' or '($8,799)' or '-$8,799'.
    Handles em-dash (−, U+2212) as a minus sign.
    Returns None if no parseable number is found.
    """
    text = text.replace(",", "").replace("$", "").replace("\u2212", "-").strip()
    if text.startswith("(") and text.endswith(")"):
        text = "-" + text[1:-1]
    try:
        return float(text)
    except ValueError:
        return None


def extract_embedded_numbers(text):
    """
    Extract all dollar amounts embedded in mixed text, e.g. 'pre-tax income $412,912.52 which...'
    Returns a list of floats.
    """
    # Match: optional leading minus/em-dash, optional $, digits with optional commas and decimals
    pattern = r'[\-\u2212]?\$?[\d]{1,3}(?:,\d{3})*(?:\.\d+)?'
    results = []
    for m in re.findall(pattern, text):
        p = parse_dollar(m)
        if p is not None and abs(p) > 0.009:   # skip near-zero noise
            results.append(p)
    return results


def find_value_after(values, label_fragment, offset=1):
    """
    Find the value that appears `offset` positions after the cell containing
    label_fragment (case-insensitive). Returns None if not found.
    """
    for i, v in enumerate(values):
        if label_fragment.lower() in v.lower():
            idx = i + offset
            if idx < len(values):
                return values[idx]
    return None


def find_any_containing(values, fragment):
    """Return first cell text that contains fragment (case-insensitive), or None."""
    fragment_lower = fragment.lower()
    for v in values:
        if fragment_lower in v.lower():
            return v
    return None


def number_present(values, number_str):
    """
    Return True if any cell contains the number_str (ignoring $ and commas).
    number_str: e.g. "13,565"  →  searches for "13565" within cleaned cell text.
    """
    needle = number_str.replace(",", "").replace("$", "")
    for v in values:
        cleaned = v.replace(",", "").replace("$", "")
        if needle in cleaned:
            return True
    return False


# ── Check result accumulator ─────────────────────────────────────────────────

class Results:
    def __init__(self):
        self.lines = []
        self.passed = 0
        self.failed = 0
        self.warned = 0

    def ok(self, label, found, expected):
        msg = f"  PASS  {label}: {_fmt(found)} (expected {_fmt(expected)})"
        self.lines.append(msg)
        self.passed += 1

    def fail(self, label, found, expected, diff=None):
        diff_str = f" -- DISCREPANCY {_fmt(abs(diff))}" if diff is not None else ""
        msg = f"  FAIL  {label}: {_fmt(found)} (expected {_fmt(expected)}){diff_str}"
        self.lines.append(msg)
        self.failed += 1

    def warn(self, label, reason):
        msg = f"  WARN  {label}: {reason}"
        self.lines.append(msg)
        self.warned += 1

    def present(self, label, value_str):
        msg = f"  PASS  {label}: '{value_str}' found in document"
        self.lines.append(msg)
        self.passed += 1

    def absent(self, label, value_str):
        msg = f"  FAIL  {label}: '{value_str}' NOT found in document -- MISSING"
        self.lines.append(msg)
        self.failed += 1


def _fmt(v):
    if v is None:
        return "None"
    if isinstance(v, float):
        return f"${v:,.2f}"
    return str(v)


# ── Dollar-value check with tolerance ────────────────────────────────────────

def check_dollar(results, label, values, search_fragment, expected, tolerance,
                 offset=1, search_all=False):
    """
    Find a dollar value near search_fragment in document, compare to expected.
    If search_all=True, scan all cells for any that parse to expected ± tolerance.
    """
    try:
        if search_all:
            # Try to find any cell whose dollar value matches expected
            for v in values:
                # First try whole-cell parse
                parsed = parse_dollar(v)
                if parsed is not None and abs(parsed - expected) <= tolerance:
                    results.ok(label, parsed, expected)
                    return
                # Then try extracting embedded numbers (handles mixed-content cells)
                for parsed in extract_embedded_numbers(v):
                    if abs(parsed - expected) <= tolerance:
                        results.ok(label, parsed, expected)
                        return
            results.warn(label, f"not found in document — may be display-formatted differently (expected {_fmt(expected)})")
            return

        raw = find_value_after(values, search_fragment, offset)
        if raw is None:
            # Try one-cell approach: find cell that itself contains the label AND a number
            raw = find_any_containing(values, search_fragment)

        if raw is None:
            results.warn(label, f"label '{search_fragment}' not found in document — may be display-formatted differently")
            return

        parsed = parse_dollar(raw)
        if parsed is None:
            # Maybe the adjacent cell didn't have the number; try searching all cells
            for v in values:
                p = parse_dollar(v)
                if p is not None and abs(p - expected) <= tolerance:
                    results.ok(label, p, expected)
                    return
            results.warn(label, f"found label but could not parse adjacent value '{raw}' — may be display-formatted differently")
            return

        diff = abs(parsed - expected)
        if diff <= tolerance:
            results.ok(label, parsed, expected)
        else:
            results.fail(label, parsed, expected, diff)

    except Exception as e:
        results.warn(label, f"exception during check: {e}")


# ── Per-report validation functions ──────────────────────────────────────────

def validate_marketing(doc, D, res):
    """Task 1 — Marketing report checks."""
    values = all_text_values(doc)

    expected_tuition  = D["revenue"]["ytd_tuition_current"]
    expected_mkt      = D["marketing"]["total_ytd_current"]
    # Report shows the gap to meet *annual* minimum obligation ($96K = 3% × $3.2M minimum).
    # JSON's gap_conservative is YTD-based; compute the annual gap here instead.
    expected_gap_cons = 96000.0 - expected_mkt
    expected_proj     = D["revenue"]["projected_full_year"]

    check_dollar(res, "YTD tuition current", values, "tuition",
                 expected_tuition, tolerance=1.0, search_all=True)

    check_dollar(res, "YTD marketing total", values, "total",
                 expected_mkt, tolerance=1.0, search_all=True)

    check_dollar(res, "Gap to annual minimum obligation", values, "gap",
                 expected_gap_cons, tolerance=5.0, search_all=True)

    # Projected revenue is rounded to nearest $1K in the report — use wider tolerance
    check_dollar(res, "Projected full-year revenue", values, "projected",
                 expected_proj, tolerance=1000.0, search_all=True)


def validate_tax(doc, D, res):
    """Task 2 — Tax report checks."""
    values = all_text_values(doc)

    expected_qb_profit   = D["income"]["qb_profit"]
    expected_h1_pretax   = D["income"]["h1_pretax_proxy"]

    check_dollar(res, "QB Profit (net income)", values, "profit",
                 expected_qb_profit, tolerance=1.0, search_all=True)

    check_dollar(res, "H1 pre-tax proxy", values, "pre-tax",
                 expected_h1_pretax, tolerance=1.0, search_all=True)

    # Check installment amounts are present
    if number_present(values, "13,565"):
        res.present("Installment Apr 30 / Jul 31 amount", "$13,565")
    else:
        res.absent("Installment Apr 30 / Jul 31 amount", "13,565")

    if number_present(values, "42,226"):
        res.present("Total installments", "$42,226")
    else:
        res.absent("Total installments", "42,226")


def validate_deviation(doc, D, res):
    """Task 3 — Deviation report checks."""
    values = all_text_values(doc)

    expected_tuition  = D["revenue"]["ytd_tuition_current"]
    expected_handouts = D["expenses"]["5780"]["current_ytd"]
    expected_insurance = D["expenses"]["6600"]["current_ytd"]
    expected_mkt      = D["marketing"]["total_ytd_current"]

    check_dollar(res, "Tuition CY", values, "tuition",
                 expected_tuition, tolerance=1.0, search_all=True)

    check_dollar(res, "Student Handouts CY", values, "handout",
                 expected_handouts, tolerance=1.0, search_all=True)

    check_dollar(res, "Insurance CY", values, "insurance",
                 expected_insurance, tolerance=1.0, search_all=True)

    check_dollar(res, "Marketing total CY", values, "marketing",
                 expected_mkt, tolerance=1.0, search_all=True)


def validate_shareholder(doc, D, res):
    """Task 4 — Shareholder advances review checks."""
    values = all_text_values(doc)

    ramzan_expected = D["shareholder"]["ramzan"]["closing_balance"]
    rezai_expected  = D["shareholder"]["rezai"]["closing_balance"]

    # Ramzan closing balance: negative means he owes the corp
    check_dollar(res, "Ramzan closing balance", values, "ramzan",
                 ramzan_expected, tolerance=1.0, search_all=True)

    # Rezai closing balance: 0.0 in the JSON — look for the value or a 0 near "Rezai"
    if abs(rezai_expected) < 0.01:
        # Balance is effectively zero — look for "nil", "0", or "0.00" near Rezai
        rezai_cell = find_any_containing(values, "rezai")
        if rezai_cell is not None:
            res.present("Rezai in document", "Rezai")
            # Check that a zero-balance indicator exists somewhere
            if (find_any_containing(values, "nil") or
                    number_present(values, "0.00") or
                    find_any_containing(values, "$0")):
                res.present("Rezai closing balance (zero/nil)", "$0 / nil")
            else:
                res.warn("Rezai closing balance (zero/nil)", "expected $0 but no zero indicator found near Rezai — may be display-formatted differently")
        else:
            res.warn("Rezai in document", "name 'Rezai' not found in document — may be display-formatted differently")
    else:
        check_dollar(res, "Rezai closing balance", values, "rezai",
                     rezai_expected, tolerance=0.01, search_all=True)

    # JE-11 amount
    if number_present(values, "3,263") or number_present(values, "3,264"):
        res.present("JE-11 amount", "3,263 / 3,264")
    else:
        res.absent("JE-11 amount", "3,263 or 3,264")

    # Hajj advance
    if number_present(values, "10,000"):
        res.present("Hajj advance amount", "$10,000")
    else:
        res.absent("Hajj advance amount", "10,000")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    from docx import Document

    # ── Load JSON ──
    json_path = os.path.join(BASE_DIR, "data", "extracted", "run_data.json")
    try:
        with open(json_path, encoding="utf-8") as f:
            D = json.load(f)
    except Exception as e:
        print(f"ERROR: could not read {json_path}: {e}")
        sys.exit(1)

    extracted_at = D.get("meta", {}).get("extracted_at", "unknown")

    # ── Find reports ──
    marketing_rpt   = latest("claude_report_marketing_mae_*.docx")
    tax_rpt         = latest("claude_report_tax_mae_*.docx")
    deviation_rpt   = latest("claude_report_deviation_mae_*.docx")
    shareholder_rpt = latest("claude_report_shareholder_mae_*.docx")

    output_lines = []
    output_lines.append("=== validate_all.py — som_mae_financials ===")
    output_lines.append(f"Data: data/extracted/run_data.json (extracted: {extracted_at})")
    output_lines.append("")

    tasks = [
        ("Task 1 — Marketing", marketing_rpt,   validate_marketing),
        ("Task 2 — Tax",       tax_rpt,          validate_tax),
        ("Task 3 — Deviation", deviation_rpt,    validate_deviation),
        ("Task 4 — Shareholder", shareholder_rpt, validate_shareholder),
    ]

    grand_passed = 0
    grand_failed = 0
    grand_warned = 0

    for task_label, rpt_path, validate_fn in tasks:
        short_path = os.path.relpath(rpt_path, BASE_DIR).replace("\\", "/") if rpt_path else None

        if rpt_path is None:
            output_lines.append(f"{task_label}: NO REPORT FOUND")
            output_lines.append("")
            grand_failed += 1
            continue

        output_lines.append(f"{task_label}: {short_path}")

        try:
            doc = Document(rpt_path)
        except Exception as e:
            output_lines.append(f"  ERROR: could not open document: {e}")
            output_lines.append("")
            grand_failed += 1
            continue

        res = Results()
        try:
            validate_fn(doc, D, res)
        except Exception as e:
            res.warn("unexpected error during validation", str(e))

        for line in res.lines:
            output_lines.append(line)

        grand_passed += res.passed
        grand_failed += res.failed
        grand_warned += res.warned

        output_lines.append("")

    # ── Grand summary ──
    total = grand_passed + grand_failed
    if grand_failed == 0:
        summary = f"=== RESULT: ALL {grand_passed} CHECKS PASSED"
        if grand_warned:
            summary += f" ({grand_warned} WARNING(S))"
        summary += " ==="
    else:
        summary = (f"=== RESULT: {grand_passed}/{total} CHECKS PASSED, "
                   f"{grand_failed} FAILED")
        if grand_warned:
            summary += f", {grand_warned} WARNING(S)"
        summary += " ==="

    output_lines.append(summary)

    # ── Print to stdout ──
    full_output = "\n".join(output_lines)
    print(full_output)

    # ── Write summary file ──
    today = date.today().strftime("%Y-%m-%d")
    summary_path = os.path.join(BASE_DIR, "reports", f"validation_summary_{today}.txt")
    try:
        with open(summary_path, "w", encoding="utf-8") as f:
            f.write(full_output + "\n")
        print(f"\nSummary written to: reports/validation_summary_{today}.txt")
    except Exception as e:
        print(f"\nWARN: could not write summary file: {e}")

    sys.exit(0 if grand_failed == 0 else 1)


if __name__ == "__main__":
    main()

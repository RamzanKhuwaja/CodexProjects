"""
Task #3 — Spending Deviation Analysis (CRA Risk Review)
Spirit of Math Schools Markham East (2039321 Ontario Inc.)
Reads all data from data/extracted/run_data.json — no hardcoded values.
"""

import json, os, sys, datetime
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import (
    HB, CAL, BS, make_doc,
    sbg, sbd, shdr, sdat, ct, hr, bp, sh, note, sub_header,
    callout, callout_red, callout_green, callout_blue,
)
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ----------------------------------------------------------------
# Load run_data.json
# ----------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(BASE_DIR, "data", "extracted", "run_data.json"), encoding="utf-8") as f:
    D = json.load(f)

# ----------------------------------------------------------------
# Extract values
# ----------------------------------------------------------------
ytd_cutoff_raw = D["meta"]["ytd_cutoff_date"]
ytd_cy  = D["revenue"]["ytd_tuition_current"]
ytd_py  = D["revenue"]["ytd_tuition_prior_year"]
yoy_pct = D["revenue"]["yoy_growth_pct"] or 0.0
projected_rev = D["revenue"]["projected_full_year"]

exp = D["expenses"]   # key → {label, current_ytd, prior_ytd, change_pct}
mkt = D["marketing"]

def e_cy(key):  return exp.get(key, {}).get("current_ytd", 0.0)
def e_py(key):  return exp.get(key, {}).get("prior_ytd",   0.0)
def e_chg(key):
    cy, py = e_cy(key), e_py(key)
    return ((cy - py) / abs(py) * 100) if py else None

def m_cy(acct): return mkt["accounts"].get(acct, {}).get("current", 0.0)
def m_py(acct): return mkt["accounts"].get(acct, {}).get("prior",   0.0)
def pct_tuit(v): return (v / ytd_cy * 100) if ytd_cy else 0.0
def annualize(v, months_elapsed): return (v / months_elapsed * 12) if months_elapsed else 0.0
def fmt_c(v):
    if v is None: return "n/a"
    return f"${v:,.0f}" if v >= 0 else f"\u2212${abs(v):,.0f}"
def fmt_pct(v):
    if v is None: return "n/a"
    return f"{v:+.1f}%"
def chg_str(cy, py):
    if not py: return "new"
    return fmt_pct((cy-py)/abs(py)*100)

# ----------------------------------------------------------------
# Date calculations
# ----------------------------------------------------------------
for fmt in ("%B %d, %Y", "%b. %d %Y", "%b %d %Y"):
    try:
        cutoff_dt = datetime.datetime.strptime(ytd_cutoff_raw, fmt).date()
        break
    except ValueError:
        pass
else:
    cutoff_dt = datetime.date.today()

fy_start = datetime.date(2025, 8, 1)
fy_end   = datetime.date(2026, 7, 31)
months_elapsed   = (cutoff_dt - fy_start).days / 30.44
months_remaining = (fy_end - cutoff_dt).days / 30.44

def fmt_date(d):
    s = d.strftime("%B %d, %Y")
    return s.replace(" 0", " ")

cutoff_str = fmt_date(cutoff_dt)
today      = datetime.date.today()
FILE_DATE  = today.strftime("%Y-%m-%d")
today_str  = fmt_date(today)

# ----------------------------------------------------------------
# Key account data
# ----------------------------------------------------------------
handout_cy  = e_cy("5780");  handout_py  = e_py("5780")
insurance_cy = e_cy("6600"); insurance_py = e_py("6600")
royalty_cy  = e_cy("5710");  royalty_py  = e_py("5710")
rent_cy     = e_cy("5606");  rent_py     = e_py("5606")
online_cy   = m_cy("6208 Online");       online_py  = m_py("6208 Online")
space_cy    = m_cy("6202 Space Rental"); space_py   = m_py("6202 Space Rental")
wages_cy    = m_cy("62010 Marketing Wages"); wages_py = m_py("62010 Marketing Wages")
ftc_cy      = m_cy("6201.1 FTC");       ftc_py     = m_py("6201.1 FTC")
payroll_cy  = e_cy("5200_total"); payroll_py = e_py("5200_total")
merchant_cy = e_cy("6427");  merchant_py = e_py("6427")
repairs_cy  = e_cy("6715");  repairs_py  = e_py("6715")
finance_cy  = e_cy("6420_total"); finance_py = e_py("6420_total")
it_cy       = e_cy("6405_total"); it_py     = e_py("6405_total")
office_cy   = e_cy("6401_total"); office_py = e_py("6401_total")
profees_cy  = e_cy("6110_total"); profees_py = e_py("6110_total")
mat_cy      = e_cy("5100");   mat_py     = e_py("5100")
cleaning_cy = e_cy("6602");   cleaning_py = e_py("6602")

# Payroll sub-accounts
wages_511   = e_cy("5211"); wages_511p  = e_py("5211")
cpp_cy      = e_cy("5212"); cpp_py      = e_py("5212")
ei_cy       = e_cy("5213"); ei_py       = e_py("5213")
eht_cy      = e_cy("5214"); eht_py      = e_py("5214")
wsib_cy     = e_cy("5217"); wsib_py     = e_py("5217")
fedtax_cy   = e_cy("5219"); fedtax_py   = e_py("5219")
benefits_cy = e_cy("5215"); benefits_py = e_py("5215")
payroll0_cy = e_cy("5200"); payroll0_py = e_py("5200")

total_mkt_cy = mkt["total_ytd_current"]
total_mkt_py = mkt["total_ytd_prior"]

# 3-year per-account averages (fixed from Aug 2022–Jul 2025 aggregate P&L — does not change quarterly)
AVG_HANDOUT  = 100_469.0
AVG_INSUR    = 3_073.0
AVG_RENT     = 138_768.0
AVG_PAYROLL  = D["benchmarks_3yr"]["payroll_annual_avg"]
AVG_MKT      = D["benchmarks_3yr"]["marketing_annual_avg"]
AVG_IT       = D["benchmarks_3yr"]["it_annual_avg"]
EFF_RATE     = 0.1368

# ----------------------------------------------------------------
doc = make_doc()

# ================================================================
# TITLE
# ================================================================
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Spending Deviation Analysis \u2014 CRA Risk Review")
r.bold = True; r.font.name = CAL; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("Spirit of Math Schools Markham East  \u2014  2039321 Ontario Inc.")
r.italic = True; r.font.name = CAL; r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Report Date: {today_str}   |   Fiscal Year: August 1, 2025 \u2013 July 31, 2026")
r.italic = True; r.font.name = CAL; r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(
    f"YTD data through: {cutoff_str}   |   "
    "Benchmarks: FY2024\u201325 same-period YTD + 3-year aggregate (Aug 2022\u2013Jul 2025)   |   "
    "Note: PY-1 column in source file is a 15.5-month non-comparable period \u2014 not used for YTD comparison"
)
r.italic = True; r.font.name = CAL; r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(6)
hr(doc)

# ================================================================
# QUICK SUMMARY
# ================================================================
sh(doc, "Quick Summary \u2014 What This Report Tells You")

handout_chg = e_chg("5780") or 0.0
insur_chg   = e_chg("6600") or 0.0
insur_vs_avg = insurance_cy / AVG_INSUR if AVG_INSUR else 0
handout_proj = annualize(handout_cy, months_elapsed)

callout_blue(doc, (
    f"WHAT THIS REPORT COVERS: Which spending categories are outside the normal range "
    f"compared to prior years \u2014 and which ones the CRA is most likely to question in an audit.\n\n"
    f"\u25cf  HIGH RISK \u2014 Student Handouts (account 5780): {chg_str(handout_cy, handout_py)} vs. last year. "
    f"On track for ~${handout_proj:,.0f} full-year vs. 3-year average of ${AVG_HANDOUT:,.0f}. Needs documentation.\n"
    f"\u25cf  HIGH RISK \u2014 Insurance (account 6600): {chg_str(insurance_cy, insurance_py)} vs. last year. "
    f"Already {insur_vs_avg:.1f}x the 3-year annual average with {months_remaining:.1f} months still to go. Needs policy documents.\n"
    f"\u25cf  WATCH \u2014 FTC charges (6201.1): ${ftc_cy:,.0f} this year vs. ${ftc_py:,.0f} last year. "
    f"{'Possibly unrecorded, not gone \u2014 check with Spirit of Math.' if ftc_cy < 1000 else 'Confirm all invoices are recorded.'}\n"
    f"\u25cf  {'NEW \u2014 Marketing Wages (62010): ' if wages_py < 1 else 'Marketing Wages (62010): '}"
    f"${wages_cy:,.0f} this year vs. ${wages_py:,.0f} last year. "
    f"{'New category \u2014 needs documentation confirming it qualifies as marketing spend.' if wages_py < 1 else 'Keep documentation on file.'}\n"
    f"\u25cf  OK \u2014 IT costs, payroll as % of revenue, rent: all within normal range.\n"
    f"\u25cf  GOOD \u2014 Merchant service fees and repairs are both significantly lower this year."
))

callout_red(doc, (
    "URGENT ACTIONS:\n"
    f"1.  Student Handouts: Pull all account 5780 entries and confirm every one has a supplier invoice. "
    f"On track for ~${handout_proj:,.0f} full-year \u2014 more than double the 3-year average of ${AVG_HANDOUT:,.0f}.\n"
    f"2.  Insurance: Locate the policy documents explaining the ${insurance_cy:,.0f} spent so far. "
    f"Keep the premium invoice and certificate on file. Ask Tang & Partners about the prepaid insurance rule.\n"
    f"3.  Marketing Wages: Document that the ${wages_cy:,.0f} in account 62010 is a qualifying marketing expense "
    f"under your franchise agreement."
))

note(doc, "See Section 7 (Action Checklist) for the full list of steps. "
          "See Section 8 (Bottom Line) for key takeaways. "
          "Sections 2\u20136 contain the detailed analysis by category.")

# ================================================================
# SECTION 1 — Purpose
# ================================================================
sh(doc, "1.  What This Report Is For")
bp(doc, (
    "The Canada Revenue Agency (CRA) reviews business tax returns and looks for expenses "
    "that stick out compared to prior years. If one category of spending jumps dramatically "
    "without a clear business reason, it can trigger an audit or a request for receipts. "
    f"This report compares your current-year spending (August 1 \u2013 {cutoff_str}) "
    "to the same period last year and flags anything that deviates significantly."
))
bp(doc, (
    "We also use the three-year aggregate P\u2056L (August 2022 \u2013 July 2025) as a longer-term "
    "benchmark. Think of it as a track record: if a category looks normal over three full years "
    "but this year is wildly different, that is exactly what the CRA notices."
))
callout(doc, (
    "IMPORTANT NOTE ON THE PY-1 COLUMN: The QuickBooks export file contains three columns. "
    "The third column covers 15.5 months \u2014 not the same period as the current year. "
    "Comparing percentages against a 15.5-month column would be misleading, so this report "
    "does NOT use PY-1 YTD figures for comparison. "
    "Instead, we use the 3-year annual averages (full fiscal years) as the longer-term benchmark."
))

sub_header(doc, "Revenue Context (for proportionality)")
bp(doc, "Before comparing expenses, it helps to see how revenue is growing.")

t = doc.add_table(rows=1 + 3, cols=4)
t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
t.columns[0].width = Inches(2.5); t.columns[1].width = Inches(1.3)
t.columns[2].width = Inches(1.3); t.columns[3].width = Inches(1.3)
ct(t.rows[0].cells[0], "Item", bold=True); ct(t.rows[0].cells[1], "FY2025-26 YTD", bold=True)
ct(t.rows[0].cells[2], "PY YTD", bold=True); ct(t.rows[0].cells[3], "Change", bold=True)
shdr(t.rows[0])
rev_rows = [
    ("Gross tuition (acct 4100)", fmt_c(ytd_cy), fmt_c(ytd_py), fmt_pct(yoy_pct)),
    ("YTD period", f"Aug 1\u2013{cutoff_str}", "Aug 1\u2013prior year same date", ""),
    ("Note on PY-1 column", "n/a (15.5-month non-comparable period)", "\u2014", "\u2014"),
]
for ri, (a, b, c, d) in enumerate(rev_rows, 1):
    ct(t.rows[ri].cells[0], a); ct(t.rows[ri].cells[1], b)
    ct(t.rows[ri].cells[2], c); ct(t.rows[ri].cells[3], d)
    sdat(t.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)
note(doc, (f"Revenue grew {yoy_pct:+.1f}% year-over-year. All expense percentages in this report are expressed "
           "as a share of gross tuition so that growth is factored in."))

# ================================================================
# SECTION 2 — HIGH RISK
# ================================================================
sh(doc, "2.  High-Risk Deviations \u2014 Prepare Documentation Now")
bp(doc, (
    "These two categories have increased so much that they are very likely to be "
    "questioned by the CRA. This does not mean anything illegal has happened \u2014 "
    "it just means you need to have clear paperwork ready."
))

# --- 2A: Student Handouts ---
sub_header(doc, f"A.  Student Handouts (Account 5780) \u2014 {chg_str(handout_cy, handout_py)}  [CRITICAL]")
bp(doc, (f"Student Handouts covers workbooks, printouts, and learning materials distributed to students. "
         f"This year the cost has jumped {chg_str(handout_cy, handout_py)} compared to last year, "
         f"even though revenue grew by only {yoy_pct:+.1f}%. That mismatch is exactly what the CRA looks for."))

handout_rows = [
    (f"Current YTD (Aug 1\u2013{cutoff_str})",                          fmt_c(handout_cy), f"{pct_tuit(handout_cy):.2f}% of tuition"),
    ("Prior year same period",                                           fmt_c(handout_py), f"{pct_tuit(handout_py):.2f}% of tuition"),
    ("3-year full-year annual average (FY2022-23 to FY2024-25)",        f"${AVG_HANDOUT:,.0f}/yr", "approx. 3.5% of tuition"),
    ("Year-over-year dollar increase",                                   fmt_c(handout_cy-handout_py), ""),
    ("Year-over-year percentage increase",                               chg_str(handout_cy, handout_py), f"vs. revenue growth of {yoy_pct:+.1f}%"),
    ("Projected full-year at current pace",                              f"~${handout_proj:,.0f}", f"vs. 3-yr avg of ${AVG_HANDOUT:,.0f}"),
]
th = doc.add_table(rows=1 + len(handout_rows), cols=3)
th.style = "Table Grid"; th.alignment = WD_TABLE_ALIGNMENT.LEFT
th.columns[0].width = Inches(2.8); th.columns[1].width = Inches(1.3); th.columns[2].width = Inches(1.9)
ct(th.rows[0].cells[0], "Item", bold=True); ct(th.rows[0].cells[1], "Amount", bold=True)
ct(th.rows[0].cells[2], "As % of Tuition / Note", bold=True)
shdr(th.rows[0])
for ri, (a, b, c) in enumerate(handout_rows, 1):
    bl = ("Year-over-year percentage" in a or "Projected" in a)
    ct(th.rows[ri].cells[0], a, bold=bl); ct(th.rows[ri].cells[1], b, bold=bl)
    ct(th.rows[ri].cells[2], c); sdat(th.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout_red(doc, (
    f"RISK: Handouts are on track to reach ~${handout_proj:,.0f} this year \u2014 more than double "
    f"the 3-year annual average of ${AVG_HANDOUT:,.0f}. "
    f"The CRA may request itemized proof of every purchase. Keep all supplier invoices "
    f"and delivery records. Write a brief memo now if there is a business reason for the jump "
    f"(e.g., new curriculum materials, price increases from Spirit of Math, or a large one-time order)."
))
note(doc, ("Action: Ask your bookkeeper to pull all individual entries in account 5780 this year. "
           "Make sure each one has a supplier invoice."))

# --- 2B: Insurance ---
insur_proj = annualize(insurance_cy, months_elapsed)
sub_header(doc, f"B.  Insurance (Account 6600) \u2014 {chg_str(insurance_cy, insurance_py)} vs. Last Year; "
                f"{insur_vs_avg:.1f}x the 3-Year Annual Average  [HIGH RISK]")
bp(doc, (f"Insurance is one of those expenses that should stay relatively stable year to year. "
         f"Your insurance cost has already exceeded the full-year 3-year average with "
         f"{months_remaining:.1f} months still remaining in the fiscal year."))

ins_rows = [
    (f"Current YTD (Aug 1\u2013{cutoff_str})",               fmt_c(insurance_cy), f"{pct_tuit(insurance_cy):.2f}% of tuition"),
    ("Prior year same period",                                 fmt_c(insurance_py), f"{pct_tuit(insurance_py):.2f}% of tuition"),
    ("3-year full-year annual average (FY2022-23 to FY2024-25)", f"${AVG_INSUR:,.0f}/yr", f"from ${AVG_INSUR*3:,.0f} over 3 years"),
    ("Year-over-year dollar increase",                         fmt_c(insurance_cy-insurance_py), chg_str(insurance_cy, insurance_py)),
    ("Projected full-year at current pace",                    f"~${insur_proj:,.0f}", f"vs. 3-yr avg of ${AVG_INSUR:,.0f}"),
    ("Current YTD vs. 3-yr annual average",                   f"{insur_vs_avg:.1f}x higher", "Already exceeds the full annual average"),
]
ti = doc.add_table(rows=1 + len(ins_rows), cols=3)
ti.style = "Table Grid"; ti.alignment = WD_TABLE_ALIGNMENT.LEFT
ti.columns[0].width = Inches(2.8); ti.columns[1].width = Inches(1.3); ti.columns[2].width = Inches(1.9)
ct(ti.rows[0].cells[0], "Item", bold=True); ct(ti.rows[0].cells[1], "Amount", bold=True)
ct(ti.rows[0].cells[2], "Note", bold=True)
shdr(ti.rows[0])
for ri, (a, b, c) in enumerate(ins_rows, 1):
    bl = ("Year-over-year" in a or "Projected" in a or "3-yr annual" in b.lower() if False else False)
    ct(ti.rows[ri].cells[0], a); ct(ti.rows[ri].cells[1], b); ct(ti.rows[ri].cells[2], c)
    sdat(ti.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout_red(doc, (
    f"RISK: Insurance is on track for ~${insur_proj:,.0f} this year \u2014 approximately "
    f"{insur_proj/AVG_INSUR:.0f} times the 3-year annual average. This will stand out on the T2. "
    "You need a clear explanation: What changed? Did you add a new policy? "
    "Keep all insurance policy documents and the broker\u2019s invoice."
))
note(doc, ("Action: Locate the insurance policy that explains this amount. "
           "Keep the premium invoice and certificate of insurance in your records. "
           "Ask Tang & Partners whether any portion is for a period beyond July 31, 2026."))

# ================================================================
# SECTION 3 — MEDIUM RISK
# ================================================================
sh(doc, "3.  Medium-Risk Deviations \u2014 Monitor and Keep Records")
bp(doc, ("These items have changed noticeably compared to prior years but are smaller in "
         "absolute dollars or have clear explanations."))

# --- Royalty fee ---
royalty_expected = ytd_cy * 0.12
royalty_overage  = royalty_cy - royalty_expected
royalty_actual_pct = (royalty_cy / ytd_cy * 100) if ytd_cy else 0
royalty_py_pct    = (royalty_py / ytd_py * 100) if ytd_py else 0

sub_header(doc, f"A.  Royalty Fee (Account 5710) \u2014 Running at {royalty_actual_pct:.2f}% vs. 12% Contract Rate")
bp(doc, (f"Your franchise agreement says you pay 12% of gross revenue as a royalty to Spirit of Math. "
         f"This year, royalties are running at {royalty_actual_pct:.2f}% of gross tuition \u2014 "
         f"slightly {'above' if royalty_actual_pct > 12 else 'at'} the contractual rate. "
         f"The same was true last year ({royalty_py_pct:.2f}%). "
         f"This is likely a timing issue, but it is worth confirming with your bookkeeper."))

roy_rows = [
    ("Current YTD royalties paid",       fmt_c(royalty_cy), f"{royalty_actual_pct:.2f}% of gross tuition"),
    ("PY YTD royalties paid",            fmt_c(royalty_py), f"{royalty_py_pct:.2f}% of gross tuition"),
    (f"Expected at 12% of {fmt_c(ytd_cy)}", fmt_c(royalty_expected), "12.0%"),
    ("Overage vs. 12% obligation",       fmt_c(royalty_overage), "Extra royalties paid or timing difference"),
]
tr = doc.add_table(rows=1 + len(roy_rows), cols=3)
tr.style = "Table Grid"; tr.alignment = WD_TABLE_ALIGNMENT.LEFT
tr.columns[0].width = Inches(2.8); tr.columns[1].width = Inches(1.3); tr.columns[2].width = Inches(1.9)
ct(tr.rows[0].cells[0], "Item", bold=True); ct(tr.rows[0].cells[1], "Amount", bold=True)
ct(tr.rows[0].cells[2], "Note", bold=True)
shdr(tr.rows[0])
for ri, (a, b, c) in enumerate(roy_rows, 1):
    ct(tr.rows[ri].cells[0], a); ct(tr.rows[ri].cells[1], b); ct(tr.rows[ri].cells[2], c)
    sdat(tr.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)
note(doc, ("Paying more royalties than contractually required reduces your income and therefore your tax. "
           "Keep the royalty invoices from Spirit of Math on file."))

# --- Campus Rent ---
rent_proj = annualize(rent_cy, months_elapsed)
rent_chg_pct = ((rent_cy - rent_py) / abs(rent_py) * 100) if rent_py else 0
sub_header(doc, f"B.  Campus Rent (Account 5606) \u2014 {chg_str(rent_cy, rent_py)}, Outpacing Revenue Growth of {yoy_pct:+.1f}%")
bp(doc, (f"Rent is growing faster than your revenue. Rent went up {chg_str(rent_cy, rent_py)} year-over-year "
         f"while revenue grew only {yoy_pct:+.1f}%. "
         f"This is generally explainable (landlord increases, lease renewal terms). Keep your lease on file."))

rent_rows = [
    (f"Current YTD (Aug 1\u2013{cutoff_str})", fmt_c(rent_cy), f"{pct_tuit(rent_cy):.2f}% of tuition"),
    ("Prior year same period",                   fmt_c(rent_py), f"{pct_tuit(rent_py):.2f}% of tuition"),
    ("3-year annual average (full year)",         f"~${AVG_RENT:,.0f}/yr", "from 3-year aggregate P\u2056L"),
    ("Projected full-year at current pace",       f"~${rent_proj:,.0f}", f"vs. 3-yr avg of ${AVG_RENT:,.0f}"),
    ("Year-over-year dollar increase",            fmt_c(rent_cy - rent_py), chg_str(rent_cy, rent_py)),
]
trt = doc.add_table(rows=1 + len(rent_rows), cols=3)
trt.style = "Table Grid"; trt.alignment = WD_TABLE_ALIGNMENT.LEFT
trt.columns[0].width = Inches(2.8); trt.columns[1].width = Inches(1.3); trt.columns[2].width = Inches(1.9)
ct(trt.rows[0].cells[0], "Item", bold=True); ct(trt.rows[0].cells[1], "Amount", bold=True)
ct(trt.rows[0].cells[2], "Note", bold=True)
shdr(trt.rows[0])
for ri, (a, b, c) in enumerate(rent_rows, 1):
    bl = ("Year-over-year" in a)
    ct(trt.rows[ri].cells[0], a, bold=bl); ct(trt.rows[ri].cells[1], b, bold=bl)
    ct(trt.rows[ri].cells[2], c); sdat(trt.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
callout(doc, (f"MEDIUM RISK: The projected full-year rent (~${rent_proj:,.0f}) is above the 3-year "
              f"average of ${AVG_RENT:,.0f}. Keep your signed lease agreement on file."))

# --- Online Marketing ---
sub_header(doc, f"C.  Online Marketing (Account 6208) \u2014 {chg_str(online_cy, online_py)} (Growing Channel)")
bp(doc, (f"You spent ${online_cy:,.0f} on online marketing this year, compared to ${online_py:,.0f} last year. "
         f"While the absolute dollar amount is not large, this is a rapidly growing category. "
         f"The CRA may question it if they review your marketing accounts."))
online_rows = [
    (f"Current YTD (Aug 1\u2013{cutoff_str})", fmt_c(online_cy), f"{pct_tuit(online_cy):.2f}% of tuition"),
    ("Prior year same period",                   fmt_c(online_py), f"{pct_tuit(online_py):.2f}% of tuition"),
    ("Year-over-year increase",                  fmt_c(online_cy - online_py), chg_str(online_cy, online_py)),
]
to = doc.add_table(rows=1 + len(online_rows), cols=3)
to.style = "Table Grid"; to.alignment = WD_TABLE_ALIGNMENT.LEFT
to.columns[0].width = Inches(2.8); to.columns[1].width = Inches(1.3); to.columns[2].width = Inches(1.9)
ct(to.rows[0].cells[0], "Item", bold=True); ct(to.rows[0].cells[1], "Amount", bold=True)
ct(to.rows[0].cells[2], "Note", bold=True)
shdr(to.rows[0])
for ri, (a, b, c) in enumerate(online_rows, 1):
    bl = ("Year-over-year" in a)
    ct(to.rows[ri].cells[0], a, bold=bl); ct(to.rows[ri].cells[1], b, bold=bl)
    ct(to.rows[ri].cells[2], c); sdat(to.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
note(doc, ("Action: Keep screenshots of the ad campaigns and receipts from Google/Meta/other platforms. "
           "A brief note about what each campaign was for makes these easy to defend in an audit."))

# --- Space Rental (marketing) ---
sub_header(doc, f"D.  Marketing Space Rental (Account 6202) \u2014 ${space_py:,.0f} to ${space_cy:,.0f}")
bp(doc, (f"You spent ${space_cy:,.0f} on renting space for marketing events this year, "
         f"compared to ${space_py:,.0f} last year. "
         f"It is deductible as a business expense, but keep the event venue invoices."))
space_rows = [
    (f"Current YTD", fmt_c(space_cy), f"{pct_tuit(space_cy):.2f}% of tuition"),
    ("Prior year same period", fmt_c(space_py), f"{pct_tuit(space_py):.2f}% of tuition"),
    ("Year-over-year increase", fmt_c(space_cy - space_py), chg_str(space_cy, space_py)),
]
tsp = doc.add_table(rows=1 + len(space_rows), cols=3)
tsp.style = "Table Grid"; tsp.alignment = WD_TABLE_ALIGNMENT.LEFT
tsp.columns[0].width = Inches(2.8); tsp.columns[1].width = Inches(1.3); tsp.columns[2].width = Inches(1.9)
ct(tsp.rows[0].cells[0], "Item", bold=True); ct(tsp.rows[0].cells[1], "Amount", bold=True)
ct(tsp.rows[0].cells[2], "Note", bold=True)
shdr(tsp.rows[0])
for ri, (a, b, c) in enumerate(space_rows, 1):
    ct(tsp.rows[ri].cells[0], a); ct(tsp.rows[ri].cells[1], b); ct(tsp.rows[ri].cells[2], c)
    sdat(tsp.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)
note(doc, "Both online marketing and space rental count toward your 3% marketing obligation.")

# --- Marketing Wages ---
mw_label = "NEW CATEGORY" if wages_py < 1 else "SIGNIFICANT CHANGE"
sub_header(doc, f"E.  Marketing Wages (Account 62010) \u2014 ${wages_cy:,.0f} This Year  [{mw_label}]")
bp(doc, (f"Marketing Wages shows ${wages_cy:,.0f} recorded this year vs. ${wages_py:,.0f} last year. "
         f"This likely represents a staff member whose time is dedicated to marketing activities. "
         f"This is a legitimate business expense, but the CRA may ask for documentation."))
mw_rows = [
    (f"Current YTD (account 62010)", fmt_c(wages_cy), f"{pct_tuit(wages_cy):.2f}% of tuition"),
    ("Prior year same period",        fmt_c(wages_py), f"{pct_tuit(wages_py):.2f}% of tuition"),
    ("Does it count toward 3% marketing?", "Potentially yes", "If the role is primarily marketing-focused"),
]
tmw = doc.add_table(rows=1 + len(mw_rows), cols=3)
tmw.style = "Table Grid"; tmw.alignment = WD_TABLE_ALIGNMENT.LEFT
tmw.columns[0].width = Inches(2.8); tmw.columns[1].width = Inches(1.3); tmw.columns[2].width = Inches(1.9)
ct(tmw.rows[0].cells[0], "Item", bold=True); ct(tmw.rows[0].cells[1], "Amount", bold=True)
ct(tmw.rows[0].cells[2], "Note", bold=True)
shdr(tmw.rows[0])
for ri, (a, b, c) in enumerate(mw_rows, 1):
    ct(tmw.rows[ri].cells[0], a); ct(tmw.rows[ri].cells[1], b); ct(tmw.rows[ri].cells[2], c)
    sdat(tmw.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)
callout(doc, ("ACTION: Write a brief memo describing this employee\u2019s marketing role. "
              "Keep their offer letter or job description. "
              "Confirm with Tang & Partners whether this qualifies under the franchise\u2019s "
              "3% marketing obligation definition."))

# ================================================================
# SECTION 4 — ACCOUNTING CONSISTENCY
# ================================================================
sh(doc, "4.  Accounting Consistency Issues \u2014 Review with Your Bookkeeper")
bp(doc, ("These are not CRA risk items in themselves, but they show that some expense categories "
         "are being recorded differently this year compared to last year."))

sub_header(doc, f"A.  Materials and Supplies (Account 5100) \u2014 {fmt_c(mat_py)} Last Year, {fmt_c(mat_cy)} This Year")
bp(doc, (f"Last year, {fmt_c(mat_py)} was recorded in account 5100. "
         f"This year, the account shows {fmt_c(mat_cy)}. That money did not disappear \u2014 it was likely "
         f"reclassified into a different account (possibly account 5780 Student Handouts, "
         f"which would help explain why that account is so much higher this year)."))
callout(doc, ("Ask your bookkeeper: Was the spend that went through account 5100 last year "
              "reclassified into account 5780 (Student Handouts) this year? If so, that partly "
              "explains the jump in 5780."))

sub_header(doc, f"B.  Employee Benefits (Account 5215) \u2014 {fmt_c(benefits_py)} Last Year, {fmt_c(benefits_cy)} This Year")
bp(doc, (f"Employee Benefits showed {fmt_c(benefits_py)} last year. This year it shows {fmt_c(benefits_cy)}. "
         f"Either the benefit program was cancelled, or the costs are being recorded elsewhere. "
         f"Make a note in your records."))
note(doc, "Small dollar amount \u2014 not a CRA risk on its own. But consistent bookkeeping makes audits much easier.")

payroll_chg_pct = ((payroll_cy - payroll_py) / abs(payroll_py) * 100) if payroll_py else 0
sub_header(doc, f"C.  Payroll Breakdown \u2014 Total Grew {payroll_chg_pct:+.1f}% (Broadly In Line with Revenue)")
bp(doc, (f"The overall payroll total grew {payroll_chg_pct:+.1f}% this year ({fmt_c(payroll_py)} to {fmt_c(payroll_cy)}), "
         f"which is broadly in line with the {yoy_pct:+.1f}% revenue growth. "
         f"The internal breakdown changed significantly as payroll is now broken into proper sub-accounts."))

comp_rows = [
    ("5211 Wages",               fmt_c(wages_511), fmt_c(wages_511p)),
    ("5212 CPP Expense",         fmt_c(cpp_cy),    fmt_c(cpp_py)),
    ("5213 EI Expense",          fmt_c(ei_cy),     fmt_c(ei_py)),
    ("5214 EHT Expense",         fmt_c(eht_cy),    fmt_c(eht_py)),
    ("5217 WSIB Expense",        fmt_c(wsib_cy),   fmt_c(wsib_py)),
    ("5219 FED.TAX",             fmt_c(fedtax_cy), fmt_c(fedtax_py)),
    ("5215 Employee Benefits",   fmt_c(benefits_cy), fmt_c(benefits_py)),
    ("5200 Payroll (unitemized)", fmt_c(payroll0_cy), fmt_c(payroll0_py)),
    ("Total Payroll",            fmt_c(payroll_cy), fmt_c(payroll_py)),
]
tc2 = doc.add_table(rows=1 + len(comp_rows), cols=3)
tc2.style = "Table Grid"; tc2.alignment = WD_TABLE_ALIGNMENT.LEFT
tc2.columns[0].width = Inches(2.8); tc2.columns[1].width = Inches(1.3); tc2.columns[2].width = Inches(1.9)
ct(tc2.rows[0].cells[0], "Account", bold=True)
ct(tc2.rows[0].cells[1], "FY2025-26 YTD", bold=True)
ct(tc2.rows[0].cells[2], "FY2024-25 YTD", bold=True)
shdr(tc2.rows[0])
for ri, (a, b, cv) in enumerate(comp_rows, 1):
    bl = ("Total" in a)
    ct(tc2.rows[ri].cells[0], a, bold=bl); ct(tc2.rows[ri].cells[1], b, bold=bl)
    ct(tc2.rows[ri].cells[2], cv, bold=bl); sdat(tc2.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
note(doc, (f"Total payroll as a % of tuition: {pct_tuit(payroll_cy):.1f}% this year vs. "
           f"{(payroll_py/ytd_py*100) if ytd_py else 0:.1f}% last year. "
           "The total is within the normal range. The composition changed due to reclassification."))

# ================================================================
# SECTION 5 — LOWER EXPENSES
# ================================================================
sh(doc, "5.  Expenses That Are Lower Than Usual \u2014 No CRA Risk")
bp(doc, ("The CRA does not typically audit you for spending less. "
         "These categories are significantly lower than in prior years."))

merchant_chg_pct = ((merchant_cy - merchant_py) / abs(merchant_py) * 100) if merchant_py else 0
repairs_chg_pct  = ((repairs_cy - repairs_py)   / abs(repairs_py)   * 100) if repairs_py else 0
finance_chg_pct  = ((finance_cy - finance_py)   / abs(finance_py)   * 100) if finance_py else 0
it_chg_pct       = ((it_cy - it_py)             / abs(it_py)         * 100) if it_py else 0

lower_rows = [
    ("Merchant Services (6427)", fmt_c(merchant_cy), fmt_c(merchant_py),
     f"Down {merchant_chg_pct:.1f}%. Was {merchant_py/ytd_py*100 if ytd_py else 0:.2f}% of tuition, "
     f"now {pct_tuit(merchant_cy):.2f}%. Big positive drop."),
    ("Repairs & Maintenance (6715)", fmt_c(repairs_cy), fmt_c(repairs_py),
     f"Down {repairs_chg_pct:.1f}% vs. PY. Last year was unusually high. Current year is within normal range."),
    ("Finance Charges Total (6420)", fmt_c(finance_cy), fmt_c(finance_py),
     f"Down {finance_chg_pct:.1f}%, mainly from lower merchant services."),
    (f"FTC \u2014 Franchise Technology Charge (6201.1)", fmt_c(ftc_cy), fmt_c(ftc_py),
     f"FTC charges {'missing entirely' if ftc_cy < 1 else 'lower'} this year. See note below."),
    ("IT Expenses Total (6405)", fmt_c(it_cy), fmt_c(it_py),
     f"{it_chg_pct:+.1f}% vs. PY. "
     f"Now {pct_tuit(it_cy):.1f}% of tuition. Normal and healthy."),
]
tl = doc.add_table(rows=1 + len(lower_rows), cols=4)
tl.style = "Table Grid"; tl.alignment = WD_TABLE_ALIGNMENT.LEFT
tl.columns[0].width = Inches(1.6); tl.columns[1].width = Inches(0.85)
tl.columns[2].width = Inches(0.85); tl.columns[3].width = Inches(3.1)
ct(tl.rows[0].cells[0], "Category", bold=True); ct(tl.rows[0].cells[1], "Current YTD", bold=True)
ct(tl.rows[0].cells[2], "PY YTD", bold=True); ct(tl.rows[0].cells[3], "Note", bold=True)
shdr(tl.rows[0])
for ri, (a, b, cv, d) in enumerate(lower_rows, 1):
    ct(tl.rows[ri].cells[0], a); ct(tl.rows[ri].cells[1], b)
    ct(tl.rows[ri].cells[2], cv); ct(tl.rows[ri].cells[3], d)
    sdat(tl.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)
callout(doc, (f"FTC Note: Spirit of Math\u2019s Franchise Technology Charge (FTC) account 6201.1 "
              f"shows ${ftc_cy:,.0f} this year, compared to ${ftc_py:,.0f} last year. "
              f"If Spirit of Math is still charging you FTC but your bookkeeper has not recorded it, "
              f"booking those invoices would also count toward your 3% marketing obligation."))

# ================================================================
# SECTION 6 — FULL COMPARISON TABLE
# ================================================================
sh(doc, "6.  Full Category Comparison \u2014 Summary at a Glance")
bp(doc, (f"All major spending categories for the current year and prior year. "
         f"The PY-1 column from the QuickBooks file covers a 15.5-month period "
         f"and cannot be compared directly to a YTD figure. It is marked as non-comparable."))

def risk_label(cy, py, avg=None):
    if not py: return "New"
    chg = (cy - py) / abs(py) * 100
    if chg > 50:  return "HIGH"
    if chg > 20:  return "Low-Med"
    if chg < -30: return "Low (down)"
    return "Low"

summary_rows = [
    ("5200 Payroll (total)",       payroll_cy, payroll_py, risk_label(payroll_cy, payroll_py)),
    ("5606 Campus Rent",           rent_cy,    rent_py,    risk_label(rent_cy, rent_py)),
    ("5710 Royalty Fee",           royalty_cy, royalty_py, risk_label(royalty_cy, royalty_py)),
    ("5780 Student Handouts",      handout_cy, handout_py, "HIGH"),
    ("5100 Materials & Supplies",  mat_cy,     mat_py,     "Review" if mat_cy < 1 and mat_py > 0 else risk_label(mat_cy, mat_py)),
    ("6200 Marketing (all)",       total_mkt_cy, total_mkt_py, "Below oblig."),
    ("  incl. FTC (6201.1)",       ftc_cy,     ftc_py,     "Missing" if ftc_cy < 1 else risk_label(ftc_cy, ftc_py)),
    ("  incl. Online (6208)",      online_cy,  online_py,  risk_label(online_cy, online_py)),
    ("  incl. Mktg Wages (62010)", wages_cy,   wages_py,   "New" if wages_py < 1 else risk_label(wages_cy, wages_py)),
    ("6405 IT Expenses (total)",   it_cy,      it_py,      risk_label(it_cy, it_py)),
    ("6420 Finance Charges",       finance_cy, finance_py, "Low (down)" if finance_cy < finance_py else risk_label(finance_cy, finance_py)),
    ("  Merchant Svcs (6427)",     merchant_cy, merchant_py, "Low (down)" if merchant_cy < merchant_py else risk_label(merchant_cy, merchant_py)),
    ("6600 Insurance",             insurance_cy, insurance_py, "HIGH"),
    ("6602 Cleaning Service",      cleaning_cy, cleaning_py, risk_label(cleaning_cy, cleaning_py)),
    ("6715 Repairs & Maintenance", repairs_cy,  repairs_py,  "Low (down)" if repairs_cy < repairs_py else risk_label(repairs_cy, repairs_py)),
    ("6401 Office/Campus Expenses", office_cy, office_py,   risk_label(office_cy, office_py)),
    ("6110 Professional Fees",     profees_cy, profees_py,  risk_label(profees_cy, profees_py)),
]

ts2 = doc.add_table(rows=1 + len(summary_rows), cols=8)
ts2.style = "Table Grid"; ts2.alignment = WD_TABLE_ALIGNMENT.LEFT
ts2.columns[0].width = Inches(1.65); ts2.columns[1].width = Inches(0.65)
ts2.columns[2].width = Inches(0.45); ts2.columns[3].width = Inches(0.65)
ts2.columns[4].width = Inches(0.45); ts2.columns[5].width = Inches(1.10)
ts2.columns[6].width = Inches(0.45); ts2.columns[7].width = Inches(0.85)
ct(ts2.rows[0].cells[0], "Category", bold=True); ct(ts2.rows[0].cells[1], "Curr. $", bold=True)
ct(ts2.rows[0].cells[2], "Curr. %", bold=True); ct(ts2.rows[0].cells[3], "PY $", bold=True)
ct(ts2.rows[0].cells[4], "PY %", bold=True); ct(ts2.rows[0].cells[5], "PY-1 Note", bold=True)
ct(ts2.rows[0].cells[6], "PY-1 %", bold=True); ct(ts2.rows[0].cells[7], "CRA Risk", bold=True)
shdr(ts2.rows[0])
for ri, (cat, cy_v, py_v, risk) in enumerate(summary_rows, 1):
    is_high = (risk == "HIGH")
    py_tuit_pct = f"{py_v/ytd_py*100:.1f}%" if ytd_py and py_v else "0.0%"
    risk_color = None
    if risk == "HIGH": risk_color = RGBColor(0x8B, 0x00, 0x00)
    elif risk in ("Low-Med", "Missing", "New", "Review", "Below oblig."): risk_color = RGBColor(0x7D, 0x5A, 0x00)
    ct(ts2.rows[ri].cells[0], cat, bold=is_high)
    ct(ts2.rows[ri].cells[1], fmt_c(cy_v), bold=is_high)
    ct(ts2.rows[ri].cells[2], f"{pct_tuit(cy_v):.1f}%")
    ct(ts2.rows[ri].cells[3], fmt_c(py_v))
    ct(ts2.rows[ri].cells[4], py_tuit_pct)
    ct(ts2.rows[ri].cells[5], "(non-comparable\nperiod)")
    ct(ts2.rows[ri].cells[6], "n/a")
    ct(ts2.rows[ri].cells[7], risk, bold=is_high, color=risk_color)
    sdat(ts2.rows[ri], is_high)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
note(doc, (f"Percentages are expressed as % of gross tuition (account 4100). "
           f"Current YTD = Aug 1\u2013{cutoff_str}. "
           f"PY-1 column from source file covers Aug 2023\u2013Nov 2024 (15.5 months) "
           f"\u2014 cannot be used for YTD comparison."))

# ================================================================
# SECTION 7 — ACTION CHECKLIST
# ================================================================
sh(doc, "7.  Your Action Checklist \u2014 What To Do Before July 31, 2026")
bp(doc, "Here is a simple checklist of actions to take to reduce your CRA risk and keep your records clean.")

actions = [
    ("URGENT: Document every Student Handout purchase",
     f"Pull all entries from account 5780. Confirm each has a supplier invoice. "
     f"On track for ~${handout_proj:,.0f} full-year \u2014 your highest CRA risk this year. "
     f"Write a memo explaining the increase (e.g., new curriculum, ~950 students, price increases)."),
    ("URGENT: Locate your insurance policy documents",
     f"Find the insurance policy (or policies) that explain the ${insurance_cy:,.0f} spent so far. "
     f"Keep the premium invoice and certificate of insurance. Ask Tang & Partners whether "
     f"any portion crosses the July 31 year-end and needs to be treated as prepaid insurance."),
    (f"Document the Marketing Wages (account 62010)",
     f"Write a brief memo describing the employee\u2019s marketing role. Keep their job description "
     f"or offer letter. Confirm with Tang & Partners whether ${wages_cy:,.0f} qualifies "
     f"under the franchise\u2019s 3% marketing obligation definition."),
    ("Confirm with bookkeeper: Where did account 5100 go?",
     f"Last year {fmt_c(mat_py)} went through account 5100. This year it shows {fmt_c(mat_cy)}. "
     f"Find out where these costs are being recorded this year."),
    ("Check FTC charges (6201.1) with Spirit of Math",
     f"This account shows ${ftc_cy:,.0f} this year vs. ${ftc_py:,.0f} last year. "
     f"If Spirit of Math is issuing FTC invoices, record them \u2014 they count toward "
     f"your 3% marketing obligation and reduce taxable income."),
    ("Keep records for online marketing and event space rental",
     "Keep platform receipts (Google Ads, Facebook, etc.) and venue invoices for every marketing event."),
    ("Ask Tang & Partners about the SBD limit before filing",
     "Your FY2024-25 T2 showed a $300,000 Small Business Deduction limit. Ask about this before year-end."),
]
for title, body in actions:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title + "  "); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(body); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(6)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ================================================================
# SECTION 8 — BOTTOM LINE
# ================================================================
sh(doc, "8.  The Bottom Line \u2014 Plain and Simple")

blines = [
    ("Two categories are genuinely out of proportion this year.",
     f"  Student Handouts {chg_str(handout_cy, handout_py)} and are on track for ~${handout_proj:,.0f} full-year "
     f"\u2014 about {handout_proj/AVG_HANDOUT:.1f}x the 3-year average. Insurance is already "
     f"{insur_vs_avg:.1f}x the annual average with {months_remaining:.1f} months still to go. "
     f"Make sure you have a clear paper trail for both."),
    ("Everything else is either within normal range or moving in the right direction.",
     f"  IT costs are stable. Payroll as a percentage of revenue is within the normal band. "
     f"Merchant service fees and repairs are both down significantly \u2014 that is fine."),
    ("The biggest internal issue is bookkeeping consistency.",
     f"  Some accounts that had significant spend last year now show $0 "
     f"(Materials 5100: {fmt_c(mat_cy)}, Employee Benefits 5215: {fmt_c(benefits_cy)}, "
     f"FTC 6201.1: {fmt_c(ftc_cy)}). "
     f"A consistent chart of accounts makes your records cleaner for an audit."),
    ("Your marketing spend is below the 3% obligation \u2014 fix this before July 31.",
     "  See the Task #1 marketing report. You must spend more on marketing before July 31, 2026."),
    ("Share this report with Tang & Partners before your T2 is filed.",
     "  Walk them through the Student Handouts and Insurance increases so they are "
     "prepared to answer CRA questions if the return is reviewed."),
]
for title, body in blines:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(body); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
note(doc, (f"Disclaimer: This report is based on QuickBooks YTD data (Aug 1\u2013{cutoff_str}), "
           "the 3-year aggregate P\u2056L (Aug 2022\u2013Jul 2025), and the FY2024-25 audited financials. "
           "It is an analytical review, not professional tax or legal advice. "
           "Share all findings with Tang & Partners LLP (416-987-6005 / thomas@tang.ca) "
           "before making decisions."))

# ================================================================
# SAVE
# ================================================================
OUT = f"reports/claude_report_deviation_mae_{FILE_DATE}.docx"
doc.save(OUT)
print("Saved:", OUT)

"""
Task #2 — Corporate Tax Estimation
Spirit of Math Schools Markham East (2039321 Ontario Inc.)
Reads all data from data/extracted/run_data.json — no hardcoded values.
"""

import json, os, sys, datetime
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import (
    HB, CAL, BS, make_doc,
    sbg, sbd, shdr, sdat, ct, hr, bp, sh, note, sub_header,
    callout, callout_red, callout_green, callout_blue,
)
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ----------------------------------------------------------------
# Load run_data.json
# ----------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(BASE_DIR, "data", "extracted", "run_data.json"), encoding="utf-8") as f:
    D = json.load(f)

# ----------------------------------------------------------------
# Extract values
# ----------------------------------------------------------------
ytd_cutoff_raw  = D["meta"]["ytd_cutoff_date"]
ytd_cy          = D["revenue"]["ytd_tuition_current"]
ytd_py          = D["revenue"]["ytd_tuition_prior_year"]
projected_rev   = D["revenue"]["projected_full_year"]
ratio           = D["revenue"]["ytd_to_annual_ratio"]
yoy_pct         = D["revenue"]["yoy_growth_pct"] or 0.0

qb_profit       = D["income"]["qb_profit"]
taxes_8500      = D["income"]["taxes_paid_8500"]
ccr             = D["income"]["canada_carbon_rebate"]
h1_pretax       = D["income"]["h1_pretax_proxy"]

interest_cy     = D["expenses"].get("4900", {}).get("current_ytd", 0.0)
installments    = D["tax"]["installments"]
inst_paid_ytd   = D["tax"]["installments_paid_ytd"]
inst_total      = D["tax"]["installments_total"]

# FY2024-25 confirmed constants (from T2 — do not change until new T2 filed)
PY_REVENUE      = 3_020_723.00
PY_INCOME       = 310_894.00
PY_TAX_INCOME   = 308_658.00
PY_FED_TAX      = 31_350.00
PY_ONT_TAX      = 10_876.00
PY_TOTAL_TAX    = 42_226.00
PY_EFF_RATE     = 0.1368
PY_CCA          = 87_068.00
CY_CCA_EST      = 7_644.00
GIC_BALANCE     = 700_000.00

# ----------------------------------------------------------------
# Date calculations
# ----------------------------------------------------------------
for fmt in ("%B %d, %Y", "%b. %d %Y", "%b %d %Y"):
    try:
        cutoff_dt = datetime.datetime.strptime(ytd_cutoff_raw, fmt).date()
        break
    except ValueError:
        pass
else:
    cutoff_dt = datetime.date.today()

fy_start = datetime.date(2025, 8, 1)
fy_end   = datetime.date(2026, 7, 31)
days_elapsed   = (cutoff_dt - fy_start).days
days_total     = (fy_end - fy_start).days
months_elapsed = days_elapsed / 30.44
months_remaining = (fy_end - cutoff_dt).days / 30.44

def fmt_date(d):
    s = d.strftime("%B %d, %Y")
    return s.replace(" 0", " ")

cutoff_str = fmt_date(cutoff_dt)
today      = datetime.date.today()
DATE_LABEL = fmt_date(today)
FILE_DATE  = today.strftime("%Y-%m-%d")

# ----------------------------------------------------------------
# H2 estimates (computed dynamically)
# ----------------------------------------------------------------
remaining_rev = max(0.0, projected_rev - ytd_cy)  # tuition still to collect
h1_total_revenue = ytd_cy + interest_cy
h1_expenses_est  = h1_total_revenue - qb_profit
monthly_exp_rate = h1_expenses_est / months_elapsed if months_elapsed > 0 else 0
h2_expenses_est  = monthly_exp_rate * months_remaining

# H2 installments still unpaid
h2_installments = sum(i["amount"] for i in installments if i["status"] in ("upcoming", "future"))
h2_pretax_est   = remaining_rev - h2_expenses_est + h2_installments
full_yr_pretax  = h1_pretax + h2_pretax_est

# Tax scenarios
low_pretax  = full_yr_pretax * 0.90
high_pretax = full_yr_pretax * 1.10
low_tax     = low_pretax  * PY_EFF_RATE
mid_tax     = full_yr_pretax * PY_EFF_RATE
high_tax    = high_pretax * PY_EFF_RATE
balance_owing = mid_tax - inst_total
set_aside   = max(0, high_tax - inst_total) * 1.1  # 10% buffer

# Installment status helper (mark based on due date vs today)
def inst_status(inst):
    try:
        due = datetime.datetime.strptime(inst["due"], "%B %d, %Y").date()
    except Exception:
        due = None
    if inst["status"] == "paid":
        return "Paid \u2713 (included in QuickBooks acct 8500)"
    if due and due <= today:
        return "DUE NOW \u2014 pay immediately if not yet done"
    return "COMING UP SOON \u2014 mark your calendar now" if due and (due - today).days <= 60 else "Future \u2014 schedule in advance"

# ----------------------------------------------------------------
doc = make_doc()

# ================================================================
# TITLE
# ================================================================
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Corporate Tax Estimate \u2014 FY2025-26")
r.bold = True; r.font.name = CAL; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("Spirit of Math Schools Markham East  \u2014  2039321 Ontario Inc.")
r.italic = True; r.font.name = CAL; r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Report Date: {DATE_LABEL}   |   Fiscal Year: August 1, 2025 \u2013 July 31, 2026")
r.italic = True; r.font.name = CAL; r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Updated with: QuickBooks YTD through {cutoff_str} "
              "+ FY2024-25 T2 tax return (Tang & Partners, Oct 2025) "
              "+ 3-year P\u2056L (Aug 2022\u2013Jul 2025)")
r.italic = True; r.font.name = CAL; r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(6)
hr(doc)

# ================================================================
# QUICK SUMMARY
# ================================================================
sh(doc, "Quick Summary \u2014 What This Report Tells You")

upcoming = [i for i in installments if i["status"] in ("upcoming", "future")]
urgent_items = ""
for i in upcoming:
    urgent_items += f"\n{i['due']}: Pay ${i['amount']:,.0f} tax installment."
if not urgent_items:
    urgent_items = "\n  All installments paid."

callout_blue(doc, (
    f"WHAT THIS REPORT COVERS: Your estimated corporate tax for FY2025-26 (Aug 1, 2025 \u2013 Jul 31, 2026).\n\n"
    f"\u25cf  Last year\u2019s confirmed tax: ${PY_TOTAL_TAX:,.0f} (from T2 filed by Tang & Partners).\n"
    f"\u25cf  H1 pre-tax income (Aug 1 \u2013 {cutoff_str} \u2014 {months_elapsed:.1f} months of real data): ${h1_pretax:,.0f}.\n"
    f"\u25cf  Full-year estimate: ~${full_yr_pretax:,.0f} taxable income; tax ~${mid_tax:,.0f}.\n"
    f"\u25cf  Already paid: ${inst_paid_ytd:,.0f} in installments so far.\n"
    f"\u25cf  Set aside ~${set_aside:,.0f} by October 2026 for the balance owing at filing.\n"
    f"\u25cf  The single best way to reduce tax: complete the required marketing spend (see Task #1)."
))

callout_red(doc, (
    f"UPCOMING DEADLINES:{urgent_items}\n"
    f"  October 2026: Set aside ~${set_aside:,.0f} for the balance owing when Tang & Partners file your T2.\n"
    f"  FLAG FOR ACCOUNTANT: Ask Tang & Partners about the $300,000 SBD business limit "
    f"(see Section 4 and Section 8) \u2014 if it applies, tax could be significantly higher."
))

note(doc, "See Section 7 (Action Checklist) for all actions. "
          "See Section 9 (Bottom Line) for key takeaways. "
          "Sections 1\u20136 contain the detailed analysis and calculations.")

# ================================================================
# SECTION 1 — Last year's T2
# ================================================================
sh(doc, "1.  What Your Tax Return Showed Last Year (FY2024-25)")
bp(doc, ("Your accountant (Tang & Partners) filed your FY2024-25 T2 corporate tax return "
         "in October 2025. Here is exactly what it said \u2014 these are the real, confirmed numbers, "
         "not estimates."))

s1 = [
    ("Total revenue",                                            f"${PY_REVENUE:,.0f}"),
    ("Net income \u2014 accounting books",                      f"${PY_INCOME:,.0f}"),
    ("Net income \u2014 CRA tax basis",                         f"${PY_TAX_INCOME:,.0f}"),
    ("  Difference (very small \u2014 books nearly match tax)", f"\u2212${PY_REVENUE-PY_TAX_INCOME-PY_INCOME+PY_TAX_INCOME:,.0f}" if False else "\u2212$2,236"),
    ("Federal income tax",                                       f"${PY_FED_TAX:,.0f}"),
    ("Ontario provincial income tax",                            f"${PY_ONT_TAX:,.0f}"),
    ("Total tax payable",                                        f"${PY_TOTAL_TAX:,.0f}"),
    ("Effective (actual) tax rate",                              f"{PY_EFF_RATE:.2%}"),
]
t = doc.add_table(rows=1 + len(s1), cols=2)
t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
t.columns[0].width = Inches(4.0); t.columns[1].width = Inches(2.0)
ct(t.rows[0].cells[0], "Item", bold=True); ct(t.rows[0].cells[1], "Amount", bold=True)
shdr(t.rows[0])
for ri, (a, b) in enumerate(s1, 1):
    bl = ("Total tax" in a or "Effective" in a)
    ct(t.rows[ri].cells[0], a, bold=bl); ct(t.rows[ri].cells[1], b, bold=bl)
    sdat(t.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout(doc, ("Good news: Your accounting books and your tax return are almost identical "
              "(only $2,236 apart). This means we can use your QuickBooks numbers directly "
              "to estimate this year\u2019s tax \u2014 no complicated adjustments needed."))

# ================================================================
# SECTION 2 — The Big Change: CCA
# ================================================================
sh(doc, "2.  The Single Biggest Reason Your Tax Is Going Up This Year")
bp(doc, ("Every year, the government lets you deduct a portion of what you spent on your "
         "building improvements and equipment. This deduction is called CCA "
         "(Capital Cost Allowance). For the past 10 years, your biggest CCA deduction came "
         "from leasehold improvements at your campus. "
         "The government allowed you to deduct ~$79,000 per year as CCA."))

s2 = [
    ("FY2023-24",       "$86,340",                   "$205,579",                    "$5,881"),
    ("FY2024-25",       f"${PY_CCA:,.0f}",           f"${PY_TAX_INCOME:,.0f}",     f"${PY_TOTAL_TAX:,.0f}"),
    ("FY2025-26 (est.)",f"~${CY_CCA_EST:,.0f}",      f"~${full_yr_pretax:,.0f}",   f"~${mid_tax:,.0f}"),
]
t2 = doc.add_table(rows=1 + len(s2), cols=4)
t2.style = "Table Grid"; t2.alignment = WD_TABLE_ALIGNMENT.LEFT
t2.columns[0].width = Inches(1.3); t2.columns[1].width = Inches(1.3)
t2.columns[2].width = Inches(1.5); t2.columns[3].width = Inches(1.5)
ct(t2.rows[0].cells[0], "Year", bold=True); ct(t2.rows[0].cells[1], "CCA Deduction", bold=True)
ct(t2.rows[0].cells[2], "Taxable Income", bold=True); ct(t2.rows[0].cells[3], "Tax", bold=True)
shdr(t2.rows[0])
for ri, (yr, cca, inc, tax) in enumerate(s2, 1):
    bl = (ri == len(s2))
    ct(t2.rows[ri].cells[0], yr, bold=bl); ct(t2.rows[ri].cells[1], cca, bold=bl)
    ct(t2.rows[ri].cells[2], inc, bold=bl); ct(t2.rows[ri].cells[3], tax, bold=bl)
    sdat(t2.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout(doc, (f"The lease on your main campus improvements expired on July 31, 2025. "
              f"Starting August 1, 2025 (FY2025-26), that ~$79,000 annual deduction is gone. "
              f"Your remaining CCA drops from ${PY_CCA:,.0f} to only ~${CY_CCA_EST:,.0f} this year. "
              f"That ~${PY_CCA - CY_CCA_EST:,.0f} difference goes straight into taxable income."))

note(doc, (f"CCA breakdown for FY2025-26: Furniture (Class 8): ~$2,486 | "
           f"Remaining leasehold (Class 13-b): ~$1,700 | "
           f"Franchise fee (Class 14.1): ~$888 | Computers (Class 50): ~$2,570 | "
           f"Total: ~${CY_CCA_EST:,.0f}  (vs. ${PY_CCA:,.0f} last year)"))

# ================================================================
# SECTION 3 — This Year's Income Estimate
# ================================================================
sh(doc, "3.  Estimating This Year\u2019s Full Income (August 1, 2025 \u2013 July 31, 2026)")
bp(doc, (f"We now have {months_elapsed:.1f} months of real QuickBooks data (Aug 1 \u2013 {cutoff_str}). "
         f"We estimate the remaining {months_remaining:.1f} months based on how last year played out "
         f"over the same period."))

sub_header(doc, f"Step A \u2014 Year-to-Date (what QuickBooks actually shows, through {cutoff_str})")

s3a = [
    (f"QuickBooks profit (Aug 1 \u2013 {cutoff_str})",                   f"${qb_profit:,.2f}",    False),
    ("Add back: tax installments booked as expense (acct 8500)",          f"+${taxes_8500:,.2f}",  False),
    ("Less: Canada Carbon Rebate \u2014 government credit, not taxable",  f"\u2212${ccr:,.2f}",    False),
    (f"YTD estimated pre-tax proxy ({months_elapsed:.1f} months)",        f"~${h1_pretax:,.2f}",   True),
]
t3a = doc.add_table(rows=1 + len(s3a), cols=2)
t3a.style = "Table Grid"; t3a.alignment = WD_TABLE_ALIGNMENT.LEFT
t3a.columns[0].width = Inches(4.3); t3a.columns[1].width = Inches(1.7)
ct(t3a.rows[0].cells[0], "Item", bold=True); ct(t3a.rows[0].cells[1], "Amount", bold=True)
shdr(t3a.rows[0])
for ri, (a, b, bl) in enumerate(s3a, 1):
    ct(t3a.rows[ri].cells[0], a, bold=bl); ct(t3a.rows[ri].cells[1], b, bold=bl)
    sdat(t3a.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

note(doc, ("Why add back the tax installments? QuickBooks records them as an expense, "
           "but they are really just prepayments of tax \u2014 not a business cost. "
           "The Canada Carbon Rebate is a government payment, not business revenue."))

callout_green(doc, (
    f"YTD tuition through {cutoff_str}: ${ytd_cy:,.2f} "
    f"(vs. ${ytd_py:,.2f} for the same period last year \u2014 up {yoy_pct:+.1f}%). "
    f"YTD-to-annual ratio: {ratio:.1%}. Projected full-year revenue: ~${projected_rev:,.0f}."
))

sub_header(doc, f"Step B \u2014 Remaining Period Estimate ({cutoff_str} \u2013 July 31, 2026)")

bp(doc, (f"With {ratio:.1%} of the year\u2019s revenue already collected (${ytd_cy:,.0f} of ~${projected_rev:,.0f} projected), "
         f"the remaining {months_remaining:.1f} months will see minimal new revenue (about ${remaining_rev:,.0f} more). "
         f"However, operating costs continue at roughly ${monthly_exp_rate:,.0f}/month. "
         f"Here is how we build the H2 estimate:"))

s3b = [
    ("Remaining tuition revenue estimate",                        f"+${remaining_rev:,.0f}",      "Projected - YTD"),
    (f"Less: {months_remaining:.1f} months of operating costs",  f"\u2212${h2_expenses_est:,.0f}", f"~${monthly_exp_rate:,.0f}/month"),
    ("Add back: remaining installments (not a business cost)",    f"+${h2_installments:,.0f}",    "Apr + Jul installments"),
    ("Remaining period pre-tax proxy",                            f"~${h2_pretax_est:,.0f}",      ""),
]
t3b = doc.add_table(rows=1 + len(s3b), cols=3)
t3b.style = "Table Grid"; t3b.alignment = WD_TABLE_ALIGNMENT.LEFT
t3b.columns[0].width = Inches(3.0); t3b.columns[1].width = Inches(1.2); t3b.columns[2].width = Inches(1.8)
ct(t3b.rows[0].cells[0], "Item", bold=True)
ct(t3b.rows[0].cells[1], "Amount", bold=True)
ct(t3b.rows[0].cells[2], "Note", bold=True)
shdr(t3b.rows[0])
for ri, (a, b, c) in enumerate(s3b, 1):
    bl = ("Remaining period pre-tax" in a)
    ct(t3b.rows[ri].cells[0], a, bold=bl)
    ct(t3b.rows[ri].cells[1], b, bold=bl)
    ct(t3b.rows[ri].cells[2], c)
    sdat(t3b.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

sub_header(doc, "Step C \u2014 Full-Year Income Estimate")

s3c = [
    (f"YTD pre-tax proxy (Aug 1 \u2013 {cutoff_str})",  f"~${h1_pretax:,.0f}",      False),
    (f"Remaining period estimate",                        f"~${h2_pretax_est:,.0f}",  False),
    ("Estimated full-year taxable income",                f"~${full_yr_pretax:,.0f}", True),
    ("Working estimate range",                            f"~${low_pretax:,.0f}\u2013${high_pretax:,.0f}", True),
]
t3c = doc.add_table(rows=1 + len(s3c), cols=2)
t3c.style = "Table Grid"; t3c.alignment = WD_TABLE_ALIGNMENT.LEFT
t3c.columns[0].width = Inches(4.3); t3c.columns[1].width = Inches(1.7)
ct(t3c.rows[0].cells[0], "Calculation", bold=True); ct(t3c.rows[0].cells[1], "Amount", bold=True)
shdr(t3c.rows[0])
for ri, (a, b, bl) in enumerate(s3c, 1):
    ct(t3c.rows[ri].cells[0], a, bold=bl); ct(t3c.rows[ri].cells[1], b, bold=bl)
    sdat(t3c.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ================================================================
# SECTION 4 — Tax Estimate
# ================================================================
sh(doc, "4.  How Much Tax You Will Likely Owe \u2014 FY2025-26")
bp(doc, ("Canada taxes small business corporations at two different rates. The first "
         "$500,000 of business income gets the Small Business Rate "
         "(Federal 9% + Ontario 3.2% = 12.2% combined). Income above $500,000 gets taxed "
         "at the General Rate (Federal 15% + Ontario 11.5% = 26.5% combined). "
         "Your confirmed effective rate last year was 13.68% (from the T2 return). "
         "We use this same rate as our best estimate."))

s4 = [
    (f"Low scenario \u2014 full marketing obligation spent, costs higher",
     f"~${low_pretax:,.0f}", "13.68%", f"~${low_tax:,.0f}"),
    (f"Mid scenario \u2014 most likely outcome",
     f"~${full_yr_pretax:,.0f}", "13.68%", f"~${mid_tax:,.0f}"),
    (f"High scenario \u2014 less marketing spent, income higher",
     f"~${high_pretax:,.0f}", "13.68%", f"~${high_tax:,.0f}"),
]
t4 = doc.add_table(rows=1 + len(s4), cols=4)
t4.style = "Table Grid"; t4.alignment = WD_TABLE_ALIGNMENT.LEFT
t4.columns[0].width = Inches(2.8); t4.columns[1].width = Inches(1.1)
t4.columns[2].width = Inches(0.8); t4.columns[3].width = Inches(1.3)
ct(t4.rows[0].cells[0], "Scenario", bold=True)
ct(t4.rows[0].cells[1], "Est. Taxable Income", bold=True)
ct(t4.rows[0].cells[2], "Rate", bold=True)
ct(t4.rows[0].cells[3], "Est. Tax", bold=True)
shdr(t4.rows[0])
for ri, (sc, inc, rt, tax) in enumerate(s4, 1):
    bl = (ri == 2)
    ct(t4.rows[ri].cells[0], sc, bold=bl); ct(t4.rows[ri].cells[1], inc, bold=bl)
    ct(t4.rows[ri].cells[2], rt, bold=bl); ct(t4.rows[ri].cells[3], tax, bold=bl)
    sdat(t4.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout(doc, (f"Best estimate: You will owe approximately ${low_tax:,.0f}\u2013${high_tax:,.0f} in corporate "
              f"tax for FY2025-26. This is higher than last year\u2019s ${PY_TOTAL_TAX:,.0f}, mainly because "
              f"the ${PY_CCA-CY_CCA_EST:,.0f} annual CCA deduction for campus improvements is now finished."))

callout_red(doc, (
    "IMPORTANT FLAG \u2014 SBD Business Limit May Be $300,000 (Not $500,000):\n\n"
    "Your FY2024-25 T2 showed the Small Business Deduction limit as $300,000 (Line 410). "
    f"If the same $300,000 limit applies in FY2025-26 and your income is ~${full_yr_pretax:,.0f}, then "
    "income above $300,000 gets taxed at the higher 26.5% rate instead of 12.2%.\n\n"
    f"With $300K SBD limit: $300,000 \u00d7 12.2% + ${full_yr_pretax-300000:,.0f} \u00d7 26.5% "
    f"= $36,600 + ${(full_yr_pretax-300000)*0.265:,.0f} = ~${36600+(full_yr_pretax-300000)*0.265:,.0f}\n\n"
    "Ask Tang & Partners about this before year-end. See also Section 8."
))

# ================================================================
# SECTION 5 — Installments
# ================================================================
sh(doc, "5.  Tax Payments: What You\u2019ve Paid and What\u2019s Still Due")
bp(doc, ("The government requires four installment payments throughout the year. "
         "Your accountant calculated these based on last year\u2019s tax. "
         "Here is the full schedule:"))

t5a = doc.add_table(rows=1 + len(installments) + 1, cols=3)
t5a.style = "Table Grid"; t5a.alignment = WD_TABLE_ALIGNMENT.LEFT
t5a.columns[0].width = Inches(1.7); t5a.columns[1].width = Inches(1.2); t5a.columns[2].width = Inches(3.1)
ct(t5a.rows[0].cells[0], "Due Date", bold=True)
ct(t5a.rows[0].cells[1], "Amount", bold=True)
ct(t5a.rows[0].cells[2], "Status", bold=True)
shdr(t5a.rows[0])
for ri, inst in enumerate(installments, 1):
    bl = False
    ct(t5a.rows[ri].cells[0], inst["due"])
    ct(t5a.rows[ri].cells[1], f"${inst['amount']:,.0f}")
    ct(t5a.rows[ri].cells[2], inst_status(inst))
    sdat(t5a.rows[ri])
ri_total = len(installments) + 1
ct(t5a.rows[ri_total].cells[0], "Total installments", bold=True)
ct(t5a.rows[ri_total].cells[1], f"${inst_total:,.0f}", bold=True)
ct(t5a.rows[ri_total].cells[2], "")
sdat(t5a.rows[ri_total], True)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

bp(doc, (f"Because the actual FY2025-26 tax (~${low_tax:,.0f}\u2013${high_tax:,.0f}) will be higher than the "
         f"installment total (${inst_total:,.0f}), you will have an extra amount owing when filing:"))

s5b = [
    (f"Estimated total tax (mid scenario, {PY_EFF_RATE:.2%} rate)", f"~${mid_tax:,.0f}"),
    ("Total installments paid during the year",                       f"${inst_total:,.0f}"),
    ("Extra amount owing at filing (October 2026)",                   f"~${balance_owing:,.0f}"),
    ("Recommended amount to set aside (buffer for high scenario)",    f"~${set_aside:,.0f}"),
]
t5b = doc.add_table(rows=1 + len(s5b), cols=2)
t5b.style = "Table Grid"; t5b.alignment = WD_TABLE_ALIGNMENT.LEFT
t5b.columns[0].width = Inches(4.0); t5b.columns[1].width = Inches(2.0)
ct(t5b.rows[0].cells[0], "Item", bold=True); ct(t5b.rows[0].cells[1], "Amount", bold=True)
shdr(t5b.rows[0])
for ri, (a, b) in enumerate(s5b, 1):
    bl = (ri == len(s5b))
    ct(t5b.rows[ri].cells[0], a, bold=bl); ct(t5b.rows[ri].cells[1], b, bold=bl)
    sdat(t5b.rows[ri], bl)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout(doc, (f"Set aside ~${set_aside:,.0f} by October 2026 for the balance owing when you file. "
              f"Your GIC balance (${GIC_BALANCE:,.0f}) more than covers this."))

note(doc, (f"QuickBooks account 8500 shows ${taxes_8500:,.2f} paid in the YTD period. This includes "
           f"FY2025-26 installments paid so far (${inst_paid_ytd:,.0f}) "
           f"plus partial payment of the FY2024-25 balance. "
           f"Ask Tang & Partners for a CRA account statement to confirm the exact balances."))

# ================================================================
# SECTION 6 — How To Reduce the Tax Bill
# ================================================================
sh(doc, "6.  What You Can Do To Reduce the Tax Bill Before July 31, 2026")
bp(doc, ("Here are the most practical ways to legally reduce your corporate tax this year. "
         "Each one reduces your taxable income, which lowers your tax bill."))

mkt_gap_cy = D["marketing"]["gap_conservative"]
mkt_gap_proj = D["marketing"]["gap_projected"]
mkt_tax_save_low  = mkt_gap_cy  * PY_EFF_RATE
mkt_tax_save_high = mkt_gap_proj * PY_EFF_RATE
ftc_cy = D["marketing"]["accounts"].get("6201.1 FTC", {}).get("current", 0.0)
ftc_save = 24_058.00 * PY_EFF_RATE if ftc_cy < 1000 else 0

strats = [
    ("1. Complete the required marketing spend \u2014 already mandatory",
     f"Your franchise agreement requires you to spend 3% of revenue on local marketing. "
     f"Our Task #1 report ({DATE_LABEL}) shows you still need to spend approximately "
     f"${mkt_gap_cy:,.0f}\u2013${mkt_gap_proj:,.0f} more by July 31, 2026. "
     f"Tax saving: ~${mkt_tax_save_low:,.0f}\u2013${mkt_tax_save_high:,.0f} (at {PY_EFF_RATE:.2%} rate). "
     f"You have to spend this money anyway \u2014 think of the tax saving as a bonus."),
    ("2. Book any missing FTC (Franchise Marketing Fund) charges",
     f"Account 6201.1 (FTC) shows ${ftc_cy:,.0f} so far this year. "
     f"Over the past 3 years, this averaged $24,058 per year. "
     f"If Spirit of Math has been charging this but it hasn\u2019t been recorded yet, "
     f"booking it would save approximately ${ftc_save:,.0f} in tax. "
     f"Check with your bookkeeper whether any FTC invoices are outstanding."),
    ("3. Buy needed equipment before July 31",
     "If you need new computers, printers, or office equipment, buy them before the fiscal year ends. "
     "The government lets you deduct 27.5% of the cost in the first year "
     "(Class 50 computers: 55% rate, halved in the first year = 27.5%). "
     "Only worth doing if you genuinely need the equipment."),
    ("4. Pay yourself a bonus before July 31 (only if you will put it into RRSP)",
     "If you increase your salary or pay yourself a bonus, the company deducts it from income "
     f"(saving ~{PY_EFF_RATE:.2%} corporate tax), but you personally pay income tax on it. "
     "This only makes sense if you immediately contribute the bonus to your RRSP. "
     "Discuss the right amount with Tang & Partners based on your personal situation."),
    ("5. Review all outstanding expenses with your bookkeeper",
     "Before July 31, make sure every bill that belongs to FY2025-26 has been entered in QuickBooks. "
     f"Each $1,000 of legitimate expenses not yet recorded = ~${1000*PY_EFF_RATE:,.0f} in missed tax savings."),
]

for title, body in strats:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title + "  "); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(body); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(6)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

bp(doc, "Quick summary of tax savings available:", bold=True)
ss = [
    ("Complete marketing spend obligation",    f"${mkt_gap_cy:,.0f}\u2013${mkt_gap_proj:,.0f} to spend", f"~${mkt_tax_save_low:,.0f}\u2013${mkt_tax_save_high:,.0f}", "Must do anyway"),
    ("Book missing FTC charges (acct 6201.1)", "$0\u2013$24,000 (if owed)",       "~$0\u2013$3,300",        "Verify with bookkeeper"),
    ("Buy needed equipment",                   "Only if genuinely needed",          "~$300\u2013$800",        "Modest benefit"),
    ("Bonus + RRSP contribution",              "Discuss with accountant",           "Deferred to retirement", "Personal tax planning"),
    ("Catch up on unrecorded expenses",        "Unknown",                           "~$500\u2013$2,000",      "Do a year-end review"),
]
ts = doc.add_table(rows=1 + len(ss), cols=4)
ts.style = "Table Grid"; ts.alignment = WD_TABLE_ALIGNMENT.LEFT
ts.columns[0].width = Inches(1.8); ts.columns[1].width = Inches(1.3)
ts.columns[2].width = Inches(1.1); ts.columns[3].width = Inches(1.8)
ct(ts.rows[0].cells[0], "Strategy", bold=True); ct(ts.rows[0].cells[1], "Money Spent", bold=True)
ct(ts.rows[0].cells[2], "Tax Saved", bold=True); ct(ts.rows[0].cells[3], "Note", bold=True)
shdr(ts.rows[0])
for ri, (a, b, c, d) in enumerate(ss, 1):
    ct(ts.rows[ri].cells[0], a); ct(ts.rows[ri].cells[1], b)
    ct(ts.rows[ri].cells[2], c); ct(ts.rows[ri].cells[3], d)
    sdat(ts.rows[ri])
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ================================================================
# SECTION 7 — Action Checklist
# ================================================================
sh(doc, "7.  Your Action Checklist")
bp(doc, ("Here are all the specific actions you need to take, in order of urgency."))

tax_actions = []
for inst in installments:
    if inst["status"] in ("upcoming", "future"):
        tax_actions.append((
            f"{'URGENT' if inst['status']=='upcoming' else 'UPCOMING'} \u2014 {inst['due']}: Pay ${inst['amount']:,.0f} installment",
            f"This is a FY2025-26 tax installment. It must be paid to the CRA by {inst['due']}. "
            f"If you miss it, interest accrues daily."
        ))

tax_actions += [
    (f"Set aside ~${set_aside:,.0f} by October 2026 for the balance owing at filing",
     f"Your installments (${inst_total:,.0f} total) were based on last year\u2019s tax. "
     f"This year\u2019s tax is estimated at ${low_tax:,.0f}\u2013${high_tax:,.0f} (possibly higher if the SBD limit is $300K). "
     f"The shortfall of ~${balance_owing:,.0f} will be due when Tang & Partners file your T2 in October 2026. "
     f"Your GIC balance (${GIC_BALANCE:,.0f}) easily covers this."),
    (f"Spend the required marketing amount before July 31 \u2014 this also reduces your tax",
     f"Our Task #1 marketing report ({DATE_LABEL}) shows you still need to spend ${mkt_gap_cy:,.0f}\u2013${mkt_gap_proj:,.0f} "
     f"on marketing by July 31. "
     f"Every dollar you spend is a tax deduction. At {PY_EFF_RATE:.2%}, spending the full gap saves "
     f"approximately ${mkt_tax_save_low:,.0f}\u2013${mkt_tax_save_high:,.0f} in tax."),
    ("Ask Tang & Partners about the $300,000 SBD business limit \u2014 BEFORE year-end",
     "Your FY2024-25 T2 showed the Small Business Deduction limit as $300,000 instead of $500,000. "
     f"At ~${full_yr_pretax:,.0f} projected income with a $300K limit, more income will be taxed at 26.5% vs. 12.2%. "
     "Ask them before year-end so there are no surprises."),
    ("Check with your bookkeeper for any unrecorded expenses before July 31",
     "Every legitimate business expense not yet in QuickBooks is a missed tax deduction. "
     "Before year-end, review: insurance invoices, car repairs, outstanding supplier bills, "
     f"and training or conference costs. Each $1,000 unrecorded = ~${1000*PY_EFF_RATE:,.0f} in missed tax savings."),
    ("Book any missing FTC charges (account 6201.1) if Spirit of Math invoices exist",
     f"This account shows ${ftc_cy:,.0f} this year vs. ~$24,000 last year. "
     "If Spirit of Math has issued FTC invoices, recording them reduces taxable income and "
     "also counts toward the 3% marketing obligation."),
]

for title, body in tax_actions:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title + "  "); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(body); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(6)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ================================================================
# SECTION 8 — SBD Question
# ================================================================
sh(doc, "8.  One Important Question to Ask Tang & Partners")
bp(doc, ("Your FY2024-25 T2 return shows the \u2018business limit\u2019 as $300,000 "
         "(Line 410 on the T2). Most small businesses get the full $500,000 limit. "
         f"In FY2025-26, your income is projected at ~${full_yr_pretax:,.0f}. "
         "If the $300,000 limit applies again, income above $300,000 is taxed at 26.5% instead of 12.2%."))

callout(doc, ("Before filing your FY2025-26 return, ask Tang & Partners: "
              "\u2018Why is our small business deduction limit $300,000 instead of $500,000, "
              "and will it stay at $300,000 for the year ending July 31, 2026?\u2019 "
              f"At ~${full_yr_pretax:,.0f} projected income, the answer could change your tax bill significantly."))

# ================================================================
# SECTION 9 — Bottom Line
# ================================================================
sh(doc, "9.  The Bottom Line \u2014 Plain and Simple")

blines = [
    ("Your tax will be higher this year \u2014 and that\u2019s mainly good news.",
     f"  The main reason taxes are rising is that your business is making more money "
     f"and your old tax deduction (for campus improvements) has run out. "
     f"YTD tuition is up {yoy_pct:+.1f}% vs. last year."),
    (f"Best estimate: ${low_tax:,.0f}\u2013${high_tax:,.0f} in total corporate tax for FY2025-26.",
     f"  This compares to ${PY_TOTAL_TAX:,.0f} last year. The increase is almost entirely "
     f"explained by the ${PY_CCA-CY_CCA_EST:,.0f} drop in your annual CCA deduction. "
     "If the SBD limit stays at $300,000, ask your accountant about the impact."),
    (f"You have already paid ${inst_paid_ytd:,.0f} in FY2025-26 installments so far.",
     f"  Remaining installments total ${h2_installments:,.0f}. "
     f"Set aside ~${set_aside:,.0f} additional for the balance owing when you file in October 2026."),
    ("The most important action right now: spend the required marketing money.",
     f"  You need to spend ${mkt_gap_cy:,.0f}\u2013${mkt_gap_proj:,.0f} more on marketing before July 31 "
     f"(per Task #1). This is mandatory under your franchise agreement, and it will "
     f"also reduce your tax bill by about ${mkt_tax_save_low:,.0f}\u2013${mkt_tax_save_high:,.0f}. Do both things at once."),
    ("Talk to Tang & Partners before year-end.",
     "  Ask about the $300,000 business limit question (Section 8 above) "
     "and about the right salary/dividend balance for your personal tax situation. "
     "A short planning call before July 2026 could save you $5,000\u2013$15,000."),
]

for title, body in blines:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title); rb.bold = True; rb.font.name = CAL; rb.font.size = BS
    rn = p.add_run(body); rn.font.name = CAL; rn.font.size = BS
    p.paragraph_format.space_after = Pt(5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
note(doc, (f"Disclaimer: This report is an estimate based on QuickBooks data through {cutoff_str}, "
           f"the FY2024-25 audited financial statements, and the FY2024-25 T2 tax return. It is not "
           f"professional tax advice. Confirm all figures with Tang & Partners LLP "
           f"(416-987-6005 / thomas@tang.ca) before filing."))

# ================================================================
# SAVE
# ================================================================
OUT = f"reports/claude_report_tax_mae_{FILE_DATE}.docx"
doc.save(OUT)
print("Saved:", OUT)

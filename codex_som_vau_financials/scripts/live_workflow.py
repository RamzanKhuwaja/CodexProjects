"""
live_workflow.py - VAU live Codex workflow helpers

Python extracts evidence and renders reports.
Codex is expected to do the final reasoning live in chat.
"""

from __future__ import annotations

import csv
import glob
import json
import re
from datetime import date, datetime
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from project_context import fiscal_year_bounds, fiscal_year_label, load_historical_context

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None

try:
    from docx import Document as DocxReader
except ImportError:  # pragma: no cover
    DocxReader = None


BASE_DIR = Path(__file__).resolve().parents[1]
CURRENT_DIR = BASE_DIR / "data" / "current"
ARCHIVE_DIR = BASE_DIR / "data" / "archive"
DOCS_DIR = BASE_DIR / "docs"
EXTRACTED_DIR = BASE_DIR / "data" / "extracted"
SOURCE_TEXT_DIR = EXTRACTED_DIR / "source_text"

PACKET_PATH = EXTRACTED_DIR / "live_session_packet.json"
PAYLOAD_TEMPLATE_PATH = EXTRACTED_DIR / "live_report_payload.template.json"

REQUIRED_KEYWORDS = {
    "pl_compare": "Profit and Loss - Compare YTD",
    "pl_3yr": "Profit and Loss - Aug 2022",
    "sh_current": "Shareholder advances - this fiscal year",
    "sh_all": "Shareholder Advances - all dates",
}

MARKETING_OBLIGATION_RATE = 0.03
CORE_REPORT_TOPICS = ["marketing", "tax", "deviation", "shareholder"]


def ensure_dirs() -> None:
    EXTRACTED_DIR.mkdir(parents=True, exist_ok=True)
    SOURCE_TEXT_DIR.mkdir(parents=True, exist_ok=True)


def safe_float(value):
    if value is None:
        return 0.0
    if isinstance(value, str):
        value = value.replace(",", "").replace("$", "").replace("\u2019", "").strip()
        if not value:
            return 0.0
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def change_pct(current, prior):
    if prior == 0.0:
        return None
    return round((current - prior) / abs(prior) * 100, 2)


def row_label(row):
    for cell in row:
        if isinstance(cell, str):
            return cell.strip().rstrip("\xa0").strip()
    return ""


def find_file(keyword, exts=(".xlsx", ".csv")):
    for ext in exts:
        matches = glob.glob(str(CURRENT_DIR / f"*{keyword}*{ext}"))
        if matches:
            return Path(matches[0])
    return None


def load_csv_rows(path: Path):
    rows = []
    with path.open(newline="", encoding="utf-8-sig") as handle:
        for row in csv.reader(handle):
            converted = tuple(cell if cell.strip() else None for cell in row)
            if any(cell is not None for cell in converted):
                rows.append(converted)
    return rows


def load_xlsx_rows(path: Path, sheet_name=None):
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active
    rows = []
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True):
        if any(cell is not None for cell in row):
            rows.append(row)
    return rows


def load_rows(path: Path):
    if path.suffix.lower() == ".csv":
        return load_csv_rows(path)
    return load_xlsx_rows(path)


def extract_docx_text(path: Path):
    if DocxReader is None:
        return ""
    doc = DocxReader(str(path))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines)


def extract_pdf_text(path: Path, page_limit=15):
    if pdfplumber is None:
        return ""
    chunks = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages[:page_limit]:
            text = (page.extract_text() or "").strip()
            if text:
                chunks.append(text)
    return "\n\n".join(chunks)


def write_source_text_cache(path: Path):
    ext = path.suffix.lower()
    if ext == ".txt":
        text = path.read_text(encoding="utf-8", errors="replace")
    elif ext == ".docx":
        text = extract_docx_text(path)
    elif ext == ".pdf":
        text = extract_pdf_text(path)
    else:
        return None

    if not text.strip():
        return None

    cache_name = re.sub(r"[^A-Za-z0-9._-]+", "_", path.stem) + ".txt"
    cache_path = SOURCE_TEXT_DIR / cache_name
    cache_path.write_text(text, encoding="utf-8")
    return cache_path


def discover_sources():
    ensure_dirs()

    required = {}
    missing = []
    for key, keyword in REQUIRED_KEYWORDS.items():
        found = find_file(keyword)
        if found:
            required[key] = found
        else:
            missing.append(keyword)

    if missing:
        raise FileNotFoundError("Missing required QuickBooks files: " + ", ".join(missing))

    optional_current = []
    for path in sorted(CURRENT_DIR.iterdir()):
        if path.is_file() and path not in required.values():
            optional_current.append(path)

    reference_files = []
    for root in [ARCHIVE_DIR, DOCS_DIR]:
        if root.exists():
            for path in sorted(root.iterdir()):
                if path.is_file():
                    reference_files.append(path)

    manifest = {"required": [], "optional_current": [], "reference": []}
    for bucket, files in [
        ("required", required.values()),
        ("optional_current", optional_current),
        ("reference", reference_files),
    ]:
        for path in files:
            cache_path = write_source_text_cache(path)
            manifest[bucket].append(
                {
                    "path": str(path),
                    "name": path.name,
                    "extension": path.suffix.lower(),
                    "cached_text_path": str(cache_path) if cache_path else None,
                }
            )

    return required, manifest


def build_lookup(rows):
    data = {}
    for row in rows:
        label = row_label(row)
        if not label:
            continue
        cy = safe_float(row[1]) if len(row) > 1 else 0.0
        py = safe_float(row[2]) if len(row) > 2 else 0.0
        data[label] = {"cy": cy, "py": py}
    return data


def read_pl_compare(path: Path):
    rows = load_rows(path)
    data = build_lookup(rows)

    cutoff_date = "Unknown"
    for row in rows[:8]:
        label = row_label(row)
        match = re.search(r"(\w+ \d+,?\s*\d{4})-(\w+ \d+,?\s*\d{4})", label)
        if match:
            cutoff_date = match.group(2).strip()
            break

    def get(label):
        return data.get(label, {"cy": 0.0, "py": 0.0})

    accounts = {}
    for label, values in data.items():
        if values["cy"] != 0.0 or values["py"] != 0.0:
            accounts[label] = {
                "current_ytd": round(values["cy"], 2),
                "prior_ytd": round(values["py"], 2),
                "change_pct": change_pct(values["cy"], values["py"]),
                "delta_amount": round(values["cy"] - values["py"], 2),
            }

    return {
        "cutoff_date": cutoff_date,
        "raw_accounts": accounts,
        "tuition_current": get("4100 Tuition")["cy"],
        "tuition_prior": get("4100 Tuition")["py"],
        "income_current": get("Total for Income")["cy"],
        "income_prior": get("Total for Income")["py"],
        "cogs_current": get("Total for Cost of Goods Sold")["cy"],
        "cogs_prior": get("Total for Cost of Goods Sold")["py"],
        "expenses_current": get("Total for Expenses")["cy"],
        "expenses_prior": get("Total for Expenses")["py"],
        "qb_profit_current": get("Profit")["cy"],
        "qb_profit_prior": get("Profit")["py"],
        "taxes_booked": get("6935 Corporate Tax Expense")["cy"],
        "taxes_booked_prior": get("6935 Corporate Tax Expense")["py"],
        "canada_carbon_rebate": get("4110.1 Canada Carbon Rebate")["cy"],
        "canada_carbon_rebate_prior": get("4110.1 Canada Carbon Rebate")["py"],
    }


def read_pl_3yr(path: Path):
    rows = load_rows(path)
    data = {}
    for row in rows:
        label = row_label(row)
        if label:
            data[label] = safe_float(row[1]) if len(row) > 1 else 0.0
    return {
        "tuition_total": round(data.get("4100 Tuition", 0.0), 2),
        "income_total": round(data.get("Total for Income", 0.0), 2),
        "cogs_total": round(data.get("Total for Cost of Goods Sold", 0.0), 2),
        "expenses_total": round(data.get("Total for Expenses", 0.0), 2),
        "marketing_total": round(data.get("Total for 6200 Marketing/Advertising/Promotion", 0.0), 2),
        "ftc_total": round(data.get("6201.2 FTC", 0.0), 2),
        "handouts_total": round(data.get("5780 Student Handouts", 0.0), 2),
        "royalty_total": round(data.get("Total for 5710 Royalty fee", data.get("5710 Royalty fee", 0.0)), 2),
    }


def parse_shareholder_sheet(rows, person_label, acct_key):
    in_section = False
    opening_balance = 0.0
    closing_balance = 0.0
    transactions = []

    for row in rows:
        if row[0] == person_label:
            in_section = True
            continue
        if not in_section:
            continue

        label0 = str(row[0]) if row[0] is not None else ""
        label1 = str(row[1]) if row[1] is not None else ""

        if row[0] is not None and row[0] != person_label and label0.startswith("29"):
            break
        if label0.startswith("Total for 2900") or label0 == "TOTAL" or label1 == "TOTAL":
            break
        if label1 == "Beginning Balance":
            opening_balance = safe_float(row[10])
            continue
        if label0.startswith(f"Total for {person_label}"):
            break

        if label1 == acct_key and row[2] is not None:
            amount = round(safe_float(row[9]), 2)
            balance = round(safe_float(row[10]), 2)
            closing_balance = balance
            transactions.append(
                {
                    "date": str(row[2]) if row[2] is not None else "",
                    "type": str(row[3]) if row[3] is not None else "",
                    "num": str(row[4]) if row[4] is not None else "",
                    "name": str(row[5]) if row[5] is not None else "",
                    "memo": str(row[6]) if row[6] is not None else "",
                    "amount": amount,
                    "balance": balance,
                }
            )

    return {
        "opening_balance": round(opening_balance, 2),
        "closing_balance": round(closing_balance, 2),
        "transactions": transactions,
    }


def parse_parent_shareholder_account(rows):
    in_section = False
    opening_balance = 0.0
    closing_balance = 0.0
    for row in rows:
        label0 = str(row[0]) if row[0] is not None else ""
        label1 = str(row[1]) if row[1] is not None else ""
        if row[0] == "2900 Shareholder's Advance":
            in_section = True
            continue
        if not in_section:
            continue
        if label0.startswith("2901 ") or label0.startswith("2902 "):
            break
        if label1 == "Beginning Balance":
            opening_balance = safe_float(row[10])
            closing_balance = opening_balance
            continue
        if label0.startswith("Total for 2900 Shareholder's Advance"):
            break
        if label1 == "2900 Shareholder's Advance" and row[2] is not None:
            closing_balance = safe_float(row[10])
    return {
        "opening_balance": round(opening_balance, 2),
        "closing_balance": round(closing_balance, 2),
    }


def read_shareholder_file(path: Path):
    rows = load_rows(path)
    parent = parse_parent_shareholder_account(rows)
    ramzan = parse_shareholder_sheet(rows, "2901 Ramzan Khuwaja", "2901 Ramzan Khuwaja")
    farah = parse_shareholder_sheet(rows, "2902 Farah Khuwaja", "2902 Farah Khuwaja")
    return parent, ramzan, farah


def build_account_change_list(raw_accounts, limit=18):
    excluded_prefixes = (
        "Total for Income",
        "Total for Cost of Goods Sold",
        "Gross Profit",
        "Total for Expenses",
        "Profit",
    )
    rows = []
    for label, values in raw_accounts.items():
        if label.startswith(excluded_prefixes):
            continue
        delta = values["delta_amount"]
        current = values["current_ytd"]
        prior = values["prior_ytd"]
        if current == 0.0 and prior == 0.0:
            continue
        rows.append(
            {
                "label": label,
                "current_ytd": current,
                "prior_ytd": prior,
                "delta_amount": delta,
                "change_pct": values["change_pct"],
                "abs_delta": abs(delta),
            }
        )
    rows.sort(key=lambda item: item["abs_delta"], reverse=True)
    return rows[:limit]


def summarize_shareholder_items(ramzan_transactions):
    hajj_total = 0.0
    no_memo_entries = []
    personal_card_items = []

    for tx in ramzan_transactions:
        memo = (tx.get("memo") or "").strip()
        memo_lower = memo.lower()
        if "hajj" in memo_lower:
            hajj_total += abs(tx["amount"])
        if tx.get("num") in {"JE-21", "JE-22", "JE-23", "JE-24"} and not memo:
            no_memo_entries.append(tx)
        if any(token in memo_lower for token in ["walmart", "cineplex", "aritzia", "air can", "uber"]):
            personal_card_items.append(tx)

    return {
        "hajj_total": round(hajj_total, 2),
        "journal_entries_without_memos": no_memo_entries,
        "possible_personal_card_items": personal_card_items,
    }


def build_live_session_packet():
    required, manifest = discover_sources()
    historical = load_historical_context(BASE_DIR)
    prior_full_year_tuition = historical["prior_year"]["review_fs"]["full_year_tuition"]

    pl_compare = read_pl_compare(required["pl_compare"])
    pl_3yr = read_pl_3yr(required["pl_3yr"])
    sh_parent_cy, sh_ramzan_cy, sh_farah_cy = read_shareholder_file(required["sh_current"])
    sh_parent_all, sh_ramzan_all, sh_farah_all = read_shareholder_file(required["sh_all"])

    tuition_current = pl_compare["tuition_current"]
    tuition_prior = pl_compare["tuition_prior"]
    income_current = pl_compare["income_current"]
    income_prior = pl_compare["income_prior"]
    cogs_current = pl_compare["cogs_current"]
    cogs_prior = pl_compare["cogs_prior"]
    expenses_current = pl_compare["expenses_current"]
    expenses_prior = pl_compare["expenses_prior"]
    qb_profit = pl_compare["qb_profit_current"]
    qb_profit_prior = pl_compare["qb_profit_prior"]
    taxes_booked = pl_compare["taxes_booked"]
    carbon_rebate = pl_compare["canada_carbon_rebate"]
    pretax_proxy = qb_profit + taxes_booked - carbon_rebate

    prior_ratio = tuition_prior / prior_full_year_tuition if prior_full_year_tuition else None
    projected_tuition = (tuition_current / prior_ratio) if prior_ratio else None

    marketing_current = pl_compare["raw_accounts"].get(
        "Total for 6200 Marketing/Advertising/Promotion", {}
    ).get("current_ytd", 0.0)
    marketing_prior = pl_compare["raw_accounts"].get(
        "Total for 6200 Marketing/Advertising/Promotion", {}
    ).get("prior_ytd", 0.0)
    marketing_obligation_ytd = round(tuition_current * MARKETING_OBLIGATION_RATE, 2)
    marketing_obligation_projected = (
        round(projected_tuition * MARKETING_OBLIGATION_RATE, 2) if projected_tuition else None
    )

    net_shareholder = round(
        sh_parent_all["closing_balance"] + sh_ramzan_cy["closing_balance"] + sh_farah_cy["closing_balance"], 2
    )
    shareholder_flags = summarize_shareholder_items(sh_ramzan_cy["transactions"])
    cutoff_dt = datetime.strptime(pl_compare["cutoff_date"], "%B %d, %Y")
    fy_start, fy_end = fiscal_year_bounds(cutoff_dt.date())

    packet = {
        "meta": {
            "generated_at": datetime.now().isoformat(),
            "workflow": "VAU live Codex session",
            "company": "Spirit of Math Schools Vaughan",
            "cutoff_date": pl_compare["cutoff_date"],
            "fiscal_year_start": fy_start.isoformat(),
            "fiscal_year_end": fy_end.isoformat(),
            "fiscal_year": fiscal_year_label(fy_start, fy_end),
            "notes": [
                "Python prepared evidence only.",
                "Final reasoning should happen live in the Codex chat session.",
                "Supplemental tax or finance docs may be added to data/current/, data/archive/, or docs/.",
            ],
        },
        "sources": manifest,
        "evidence": {
            "revenue_and_profit": {
                "tuition_current_ytd": round(tuition_current, 2),
                "tuition_prior_ytd": round(tuition_prior, 2),
                "tuition_yoy_pct": change_pct(tuition_current, tuition_prior),
                "income_current_ytd": round(income_current, 2),
                "income_prior_ytd": round(income_prior, 2),
                "income_yoy_pct": change_pct(income_current, income_prior),
                "cogs_current_ytd": round(cogs_current, 2),
                "cogs_prior_ytd": round(cogs_prior, 2),
                "cogs_yoy_pct": change_pct(cogs_current, cogs_prior),
                "expenses_current_ytd": round(expenses_current, 2),
                "expenses_prior_ytd": round(expenses_prior, 2),
                "expenses_yoy_pct": change_pct(expenses_current, expenses_prior),
                "qb_profit_current_ytd": round(qb_profit, 2),
                "qb_profit_prior_ytd": round(qb_profit_prior, 2),
                "qb_profit_yoy_pct": change_pct(qb_profit, qb_profit_prior),
                "profit_margin_current_pct": round((qb_profit / income_current) * 100, 2) if income_current else None,
                "profit_margin_prior_pct": round((qb_profit_prior / income_prior) * 100, 2) if income_prior else None,
                "projected_full_year_tuition_from_prior_ratio": round(projected_tuition, 2) if projected_tuition else None,
            },
            "marketing": {
                "spent_current_ytd": round(marketing_current, 2),
                "spent_prior_ytd": round(marketing_prior, 2),
                "spend_yoy_pct": change_pct(marketing_current, marketing_prior),
                "obligation_ytd": marketing_obligation_ytd,
                "obligation_projected": marketing_obligation_projected,
                "gap_ytd": round(marketing_obligation_ytd - marketing_current, 2),
                "gap_projected": round(marketing_obligation_projected - marketing_current, 2) if marketing_obligation_projected else None,
                "ftc_current": round(pl_compare["raw_accounts"].get("6201.2 FTC", {}).get("current_ytd", 0.0), 2),
            },
            "tax": {
                "qb_profit_current_ytd": round(qb_profit, 2),
                "corporate_tax_expense_booked": round(taxes_booked, 2),
                "canada_carbon_rebate": round(carbon_rebate, 2),
                "pretax_proxy_current_ytd": round(pretax_proxy, 2),
                "prior_ytd_pretax_proxy": round(
                    pl_compare["qb_profit_prior"]
                    + pl_compare["taxes_booked_prior"]
                    - pl_compare["canada_carbon_rebate_prior"],
                    2,
                ),
                "small_business_limit": round(historical["sbd_limit"], 2),
                "last_year_full_year_tuition": round(prior_full_year_tuition, 2),
                "last_year_total_tax": round(historical["prior_year"]["t2"]["total_tax_payable"], 2),
                "last_year_taxable_income": round(historical["prior_year"]["t2"]["taxable_income"], 2),
                "last_year_pretax_income": round(
                    historical["prior_year"]["review_fs"]["net_income_before_tax"], 2
                ),
                "installments": [],
                "installment_status_source": "No installment payment status is derived from the provided QuickBooks files.",
            },
            "shareholder": {
                "parent_2900_closing": sh_parent_all["closing_balance"],
                "ramzan_closing": sh_ramzan_cy["closing_balance"],
                "farah_closing": sh_farah_cy["closing_balance"],
                "net_current_balance": net_shareholder,
                "hajj_total_in_ramzan_account": shareholder_flags["hajj_total"],
                "journal_entries_without_memos": shareholder_flags["journal_entries_without_memos"],
                "possible_personal_card_items": shareholder_flags["possible_personal_card_items"],
            },
            "historical_context": {
                "three_year_totals": pl_3yr,
                "largest_account_moves": build_account_change_list(pl_compare["raw_accounts"]),
                "shareholder_all_dates": {
                    "parent_2900": sh_parent_all,
                    "ramzan": {
                        "opening_balance": sh_ramzan_all["opening_balance"],
                        "closing_balance": sh_ramzan_all["closing_balance"],
                    },
                    "farah": {
                        "opening_balance": sh_farah_all["opening_balance"],
                        "closing_balance": sh_farah_all["closing_balance"],
                    },
                },
            },
        },
        "briefing_order": CORE_REPORT_TOPICS,
        "codex_brief_prompts": {
            "marketing": [
                "Is marketing spend enough to meet the franchise obligation?",
                "Which channels look present, absent, or unusually changed?",
                "What should the owner do before year-end?",
            ],
            "tax": [
                "Does the revenue-profit relationship support a lower tax bill or not?",
                "Which estimate range is defensible from the evidence?",
                "What should be confirmed with the accountant from extra docs?",
            ],
            "deviation": [
                "Which changes are operational, and which are CRA-risk or classification issues?",
                "Which items need documentation or reclassification?",
                "Which issues are most urgent to explain cleanly?",
            ],
            "shareholder": [
                "What is the real net shareholder position after the parent account offset?",
                "Which transactions need documentation or cleanup?",
                "What should the owner do before year-end?",
            ],
        },
    }

    PACKET_PATH.write_text(json.dumps(packet, indent=2, ensure_ascii=False), encoding="utf-8")
    build_payload_template(packet)
    return packet


def build_payload_template(packet):
    template = {
        "meta": {
            "generated_for": packet["meta"]["company"],
            "report_date": date.today().isoformat(),
            "cutoff_date": packet["meta"]["cutoff_date"],
            "fiscal_year": packet["meta"]["fiscal_year"],
            "workflow": "Codex live-session payload template",
        },
        "reports": {},
    }

    for topic in CORE_REPORT_TOPICS:
        template["reports"][topic] = {
            "topic": topic,
            "title": "",
            "main_answer": "",
            "numbers_to_know": [],
            "watch_items": [],
            "sections": [
                {"heading": "What I See", "points": []},
                {"heading": "What This Means for You", "points": []},
                {"heading": "What To Do Next", "points": []},
            ],
            "questions_to_confirm": [],
            "bottom_line": "",
        }

    PAYLOAD_TEMPLATE_PATH.write_text(json.dumps(template, indent=2, ensure_ascii=False), encoding="utf-8")


def add_cell_text(cell, text, bold=False, size=10.5, color=None):
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_callout(doc, label, text, fill_hex, text_rgb):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)

    head = p.add_run(label + "  ")
    head.bold = True
    head.font.name = "Calibri"
    head.font.size = Pt(11)
    head.font.color.rgb = text_rgb

    body = p.add_run(text)
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = text_rgb

    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.space_after = Pt(8)


def make_report_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    return doc


def render_single_report(output_path: Path, payload, meta):
    doc = make_report_doc()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title = p.add_run(payload["title"])
    title.bold = True
    title.font.name = "Calibri"
    title.font.size = Pt(18)
    title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    p = doc.add_paragraph()
    subtitle = p.add_run(
        f"{meta['generated_for']} | Report date {meta['report_date']} | Data through {meta['cutoff_date']}"
    )
    subtitle.italic = True
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(9.5)
    subtitle.font.color.rgb = RGBColor(0x5C, 0x5C, 0x5C)
    p.paragraph_format.space_after = Pt(8)

    add_callout(doc, "Main Answer", payload["main_answer"], "E8F0FE", RGBColor(0x1F, 0x38, 0x96))

    if payload.get("numbers_to_know"):
        p = doc.add_paragraph()
        run = p.add_run("Numbers To Know")
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(12.5)

        table = doc.add_table(rows=1 + len(payload["numbers_to_know"]), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        add_cell_text(table.rows[0].cells[0], "Item", bold=True)
        add_cell_text(table.rows[0].cells[1], "Value", bold=True)
        shade_cell(table.rows[0].cells[0], "D9EAD3")
        shade_cell(table.rows[0].cells[1], "D9EAD3")
        for idx, row in enumerate(payload["numbers_to_know"], start=1):
            add_cell_text(table.rows[idx].cells[0], row.get("label", ""))
            add_cell_text(table.rows[idx].cells[1], row.get("value", ""))

    if payload.get("watch_items"):
        add_callout(
            doc,
            "Watch",
            " | ".join(payload["watch_items"]),
            "FDECEA",
            RGBColor(0x8B, 0x00, 0x00),
        )

    for section in payload.get("sections", []):
        p = doc.add_paragraph()
        run = p.add_run(section.get("heading", ""))
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

        for point in section.get("points", []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(point)
            run.font.name = "Calibri"
            run.font.size = Pt(10.8)

    if payload.get("questions_to_confirm"):
        p = doc.add_paragraph()
        run = p.add_run("Items To Confirm")
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
        for question in payload["questions_to_confirm"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(question)
            run.font.name = "Calibri"
            run.font.size = Pt(10.8)

    add_callout(doc, "Bottom Line", payload["bottom_line"], "EAF4EA", RGBColor(0x1A, 0x5C, 0x1A))

    p = doc.add_paragraph()
    note = p.add_run(
        "Prepared in a live Codex session. Python handled extraction and formatting; the report judgment was made during the chat review."
    )
    note.italic = True
    note.font.name = "Calibri"
    note.font.size = Pt(9)
    note.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)

    doc.save(str(output_path))


def render_report_bundle(payload_path: Path):
    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    meta = payload["meta"]
    outputs = []
    for topic in CORE_REPORT_TOPICS:
        report = payload["reports"].get(topic)
        if not report:
            raise ValueError(f"Missing report payload for topic: {topic}")
        if not report.get("title") or not report.get("main_answer") or not report.get("bottom_line"):
            raise ValueError(f"Report payload for '{topic}' is incomplete")
        output_path = BASE_DIR / "reports" / f"codex_live_report_{topic}_vau_{meta['report_date']}.docx"
        render_single_report(output_path, report, meta)
        outputs.append(str(output_path))
    return outputs

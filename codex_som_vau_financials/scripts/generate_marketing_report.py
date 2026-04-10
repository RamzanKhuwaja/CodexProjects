"""
Short-form marketing report for VAU.
"""

import datetime
import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import BS, CAL, bp, callout_blue, callout_red, ct, hr, make_doc, note, sdat, sh, shdr
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(BASE_DIR, "data", "extracted", "run_data.json"), encoding="utf-8") as f:
    D = json.load(f)

meta = D["meta"]
rev = D["revenue"]
mkt = D["marketing"]

cutoff_str = meta["ytd_cutoff_date"]
fy_label = meta["fiscal_year_label"]
fy_start = datetime.date.fromisoformat(meta["fiscal_year_start"])
fy_end = datetime.date.fromisoformat(meta["fiscal_year_end"])
ytd_cy = rev["ytd_tuition_current"]
ytd_py = rev["ytd_tuition_prior_year"]
proj_full_year = rev["projected_full_year"] or ytd_cy
yoy_pct = rev["yoy_growth_pct"] or 0.0

total_spent = mkt["total_ytd_current"]
total_spent_py = mkt["total_ytd_prior"]
obl_proj = mkt.get("obligation_projected", 0.0)
gap_proj = mkt.get("gap_projected", 0.0)
accounts = mkt["accounts"]

for fmt in ("%B %d, %Y", "%B %d %Y", "%b %d, %Y", "%Y-%m-%d"):
    try:
        cutoff_dt = datetime.datetime.strptime(cutoff_str, fmt).date()
        break
    except ValueError:
        pass
else:
    cutoff_dt = datetime.date.today()

months_elapsed = max(0.1, (cutoff_dt - fy_start).days / 30.44)
months_remaining = max(0.0, (fy_end - cutoff_dt).days / 30.44)
monthly_needed = gap_proj / months_remaining if months_remaining else gap_proj


def pct_change(current, prior):
    if prior == 0:
        return None
    return (current - prior) / abs(prior) * 100


def fmt_currency(value, decimals=0):
    if value is None:
        return "n/a"
    return f"${value:,.{decimals}f}"


def fmt_pct(value):
    if value is None:
        return "New"
    sign = "+" if value >= 0 else ""
    return f"{sign}{value:.1f}%"


def account_change_rows():
    rows = []
    for name, vals in accounts.items():
        current = vals.get("current", 0.0)
        prior = vals.get("prior", 0.0)
        delta = current - prior
        if current == 0 and prior == 0:
            continue
        rows.append(
            {
                "name": name,
                "current": current,
                "prior": prior,
                "delta": delta,
                "pct": pct_change(current, prior),
            }
        )
    return rows


changes = account_change_rows()
top_up = sorted([r for r in changes if r["delta"] > 0], key=lambda x: x["delta"], reverse=True)[:3]
top_down = sorted([r for r in changes if r["delta"] < 0], key=lambda x: x["delta"])[:3]

doc = make_doc()
file_date = datetime.date.today().strftime("%Y-%m-%d")
out_path = os.path.join(BASE_DIR, f"reports/claude_report_marketing_vau_{file_date}.docx")
today_str = datetime.date.today().strftime("%B %d, %Y")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Marketing Report")
r.bold = True
r.font.name = CAL
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("Spirit of Math Schools Vaughan  |  2236262 Ontario Inc.")
r.italic = True
r.font.name = CAL
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Report Date: {today_str}   |   Fiscal Year: {fy_label}")
r.italic = True
r.font.name = CAL
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"QuickBooks data through {cutoff_str}")
r.italic = True
r.font.name = CAL
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(6)
hr(doc)

sh(doc, "1. Key Numbers")
callout_blue(
    doc,
    (
        f"You have spent {fmt_currency(total_spent)} on marketing so far. "
        f"Your current full-year obligation is about {fmt_currency(obl_proj)}. "
        f"That leaves about {fmt_currency(gap_proj)} still to spend by {fy_end.strftime('%B %d, %Y')}."
    ),
)
callout_red(
    doc,
    (
        f"Calm action: spread the remaining {fmt_currency(gap_proj)} over about "
        f"{months_remaining:.1f} months. That is roughly {fmt_currency(monthly_needed)} per month."
    ),
)

summary_rows = [
    ("YTD tuition revenue", fmt_currency(ytd_cy, 2), "This year"),
    ("YTD tuition last year", fmt_currency(ytd_py, 2), f"{fmt_pct(yoy_pct)} vs last year"),
    ("Projected full-year revenue", fmt_currency(proj_full_year, 2), "Used for 3% test"),
    ("Projected marketing obligation", fmt_currency(obl_proj, 2), "3% of projected revenue"),
    ("Marketing spent so far", fmt_currency(total_spent, 2), "Actual YTD spend"),
    ("Remaining estimated gap", fmt_currency(gap_proj, 2), "Still to spend"),
]
table = doc.add_table(rows=1 + len(summary_rows), cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.columns[0].width = Inches(2.6)
table.columns[1].width = Inches(1.4)
table.columns[2].width = Inches(2.4)
ct(table.rows[0].cells[0], "Item", bold=True)
ct(table.rows[0].cells[1], "Amount", bold=True)
ct(table.rows[0].cells[2], "What it means", bold=True)
shdr(table.rows[0])
for i, row in enumerate(summary_rows, 1):
    ct(table.rows[i].cells[0], row[0])
    ct(table.rows[i].cells[1], row[1])
    ct(table.rows[i].cells[2], row[2])
    sdat(table.rows[i])

sh(doc, "2. Main Changes to Notice")
bp(
    doc,
    f"Marketing spend is {fmt_currency(total_spent)} this year versus {fmt_currency(total_spent_py)} last year. "
    "The lines below show what is unusually high and unusually low so you can keep balance.",
)

change_rows = []
for item in top_up + top_down:
    direction = "Higher" if item["delta"] >= 0 else "Lower"
    change_rows.append(
        (
            item["name"],
            direction,
            fmt_currency(abs(item["delta"])),
            fmt_pct(item["pct"]),
        )
    )

change_table = doc.add_table(rows=1 + len(change_rows), cols=4)
change_table.style = "Table Grid"
change_table.alignment = WD_TABLE_ALIGNMENT.LEFT
change_table.columns[0].width = Inches(2.8)
change_table.columns[1].width = Inches(1.0)
change_table.columns[2].width = Inches(1.0)
change_table.columns[3].width = Inches(1.2)
ct(change_table.rows[0].cells[0], "Account", bold=True)
ct(change_table.rows[0].cells[1], "Direction", bold=True)
ct(change_table.rows[0].cells[2], "$ Change", bold=True)
ct(change_table.rows[0].cells[3], "% Change", bold=True)
shdr(change_table.rows[0])
for i, row in enumerate(change_rows, 1):
    ct(change_table.rows[i].cells[0], row[0])
    ct(change_table.rows[i].cells[1], row[1])
    ct(change_table.rows[i].cells[2], row[2])
    ct(change_table.rows[i].cells[3], row[3])
    sdat(change_table.rows[i])

ftc = accounts.get("6201.2 FTC", {"current": 0.0, "prior": 0.0})
bp(
    doc,
    f"FTC is {fmt_currency(ftc['current'])} this year versus {fmt_currency(ftc['prior'])} last year. "
    "That lower FTC amount is one reason the marketing total looks behind.",
)

sh(doc, "3. What To Do")
for title, body in [
    (
        "Keep the tone calm and practical.",
        f"You do not need one big spend. Aim for about {fmt_currency(monthly_needed)} per month for the rest of the year.",
    ),
    (
        "Track every marketing invoice clearly.",
        "Keep receipts and a short note showing why the spend is marketing-related.",
    ),
    (
        "Confirm the FTC expectation.",
        "If head office will still charge FTC later, your gap will shrink. If not, you need to cover the gap with other marketing spend.",
    ),
]:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title + "  ")
    rb.bold = True
    rb.font.name = CAL
    rb.font.size = BS
    rn = p.add_run(body)
    rn.font.name = CAL
    rn.font.size = BS

sh(doc, "4. Bottom Line")
bp(
    doc,
    f"Revenue is up {fmt_pct(yoy_pct)}. The school is therefore expected to spend about {fmt_currency(obl_proj)} on marketing this year. "
    f"You have already spent {fmt_currency(total_spent)}, so the current shortfall is about {fmt_currency(gap_proj)}.",
)
bp(
    doc,
    "This is manageable. The key is to spread the remaining spend across the rest of the school year and keep simple records.",
)

note(
    doc,
    f"Based on QuickBooks data through {cutoff_str}. This is a management report, not professional advice.",
)

doc.save(out_path)
print("Saved:", out_path)

"""
Short-form deviation report for VAU.
"""

import datetime
import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import BS, CAL, bp, callout_blue, callout_red, ct, hr, make_doc, note, sdat, sh, shdr
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(BASE_DIR, "data", "extracted", "run_data.json"), encoding="utf-8") as f:
    D = json.load(f)

meta = D["meta"]
rev = D["revenue"]
exp = D["expenses"]
mkt = D["marketing"]

cutoff_str = meta["ytd_cutoff_date"]
fy_label = meta["fiscal_year_label"]
ytd_cy = rev["ytd_tuition_current"]
ytd_py = rev["ytd_tuition_prior_year"]
marketing_cy = mkt["total_ytd_current"]
marketing_py = mkt["total_ytd_prior"]


def e_cy(key):
    return exp.get(key, {}).get("current_ytd", 0.0)


def e_py(key):
    return exp.get(key, {}).get("prior_ytd", 0.0)


def pct_change(current, prior):
    if prior == 0:
        return None
    return (current - prior) / abs(prior) * 100


def fmt_currency(value, decimals=0):
    if value is None:
        return "n/a"
    return f"${value:,.{decimals}f}"


def fmt_pct(value):
    if value is None:
        return "New"
    sign = "+" if value >= 0 else ""
    return f"{sign}{value:.1f}%"


items = [
    ("Student Handouts", "5780", e_cy("5780"), e_py("5780"), "Higher"),
    ("Service Fee 5711", "5711", e_cy("5711"), e_py("5711"), "Higher"),
    ("Automobile costs", "6300_total", e_cy("6300_total"), e_py("6300_total"), "Higher"),
    ("Materials and Supplies", "5300", e_cy("5300"), e_py("5300"), "Lower"),
    ("IT Expense total", "6405_total", e_cy("6405_total") or e_cy("6405"), e_py("6405_total") or e_py("6405"), "Lower"),
    ("Office/Campus Expenses", "6401_total", e_cy("6401_total"), e_py("6401_total"), "Lower"),
]

rows = []
for label, key, current, prior, direction in items:
    delta = current - prior
    rows.append(
        {
            "label": label,
            "key": key,
            "current": current,
            "prior": prior,
            "delta": delta,
            "pct": pct_change(current, prior),
            "direction": direction,
        }
    )

doc = make_doc()
file_date = datetime.date.today().strftime("%Y-%m-%d")
out_path = os.path.join(BASE_DIR, f"reports/claude_report_deviation_vau_{file_date}.docx")
today_str = datetime.date.today().strftime("%B %d, %Y")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Deviation Report")
r.bold = True
r.font.name = CAL
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("Spirit of Math Schools Vaughan  |  2236262 Ontario Inc.")
r.italic = True
r.font.name = CAL
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Report Date: {today_str}   |   Fiscal Year: {fy_label}")
r.italic = True
r.font.name = CAL
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"QuickBooks data through {cutoff_str}")
r.italic = True
r.font.name = CAL
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(6)
hr(doc)

sh(doc, "1. Main Message")
callout_blue(
    doc,
    (
        f"The books are not alarming, but a few lines need explanation. "
        f"Revenue is {fmt_currency(ytd_cy)} this year versus {fmt_currency(ytd_py)} last year. "
        f"Marketing is {fmt_currency(marketing_cy)} this year versus {fmt_currency(marketing_py)} last year."
    ),
)
callout_red(
    doc,
    (
        "The main review items are Student Handouts, Service Fee 5711, and Automobile costs. "
        "The balancing low items are Materials and Supplies, IT Expense total, and Office/Campus Expenses."
    ),
)

sh(doc, "2. Items to Notice")
table = doc.add_table(rows=1 + len(rows), cols=6)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.columns[0].width = Inches(2.2)
table.columns[1].width = Inches(0.9)
table.columns[2].width = Inches(0.9)
table.columns[3].width = Inches(0.9)
table.columns[4].width = Inches(0.9)
table.columns[5].width = Inches(1.3)
ct(table.rows[0].cells[0], "Category", bold=True)
ct(table.rows[0].cells[1], "Now", bold=True)
ct(table.rows[0].cells[2], "Last Year", bold=True)
ct(table.rows[0].cells[3], "$ Change", bold=True)
ct(table.rows[0].cells[4], "% Change", bold=True)
ct(table.rows[0].cells[5], "Why it matters", bold=True)
shdr(table.rows[0])
for i, row in enumerate(rows, 1):
    note_text = "Needs explanation" if row["direction"] == "Higher" else "Gives balance"
    ct(table.rows[i].cells[0], row["label"])
    ct(table.rows[i].cells[1], fmt_currency(row["current"], 2))
    ct(table.rows[i].cells[2], fmt_currency(row["prior"], 2))
    ct(table.rows[i].cells[3], fmt_currency(abs(row["delta"]), 2))
    ct(table.rows[i].cells[4], fmt_pct(row["pct"]))
    ct(table.rows[i].cells[5], note_text)
    sdat(table.rows[i])

bp(
    doc,
    f"Student Handouts are {fmt_currency(e_cy('5780'))} now versus {fmt_currency(e_py('5780'))} last year. "
    f"Service Fee 5711 is {fmt_currency(e_cy('5711'))} this year and was {fmt_currency(e_py('5711'))} last year.",
)
bp(
    doc,
    "Service Fee 5711 has been identified as a new franchisor charge. It should still be confirmed with the accountant for correct accounting and tax treatment.",
)

sh(doc, "3. What To Do")
for title, body in [
    (
        "Keep a simple explanation for each higher item.",
        "If CRA or the accountant asks, you should be able to explain the business reason in one or two lines.",
    ),
    (
        "Do not only focus on high items.",
        "The lower categories matter too because they help show overall balance in the books.",
    ),
    (
        "Confirm new or unusual lines.",
        "Service Fee 5711 and any personal-looking items should be checked before year-end.",
    ),
]:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title + "  ")
    rb.bold = True
    rb.font.name = CAL
    rb.font.size = BS
    rn = p.add_run(body)
    rn.font.name = CAL
    rn.font.size = BS

sh(doc, "4. Bottom Line")
bp(
    doc,
    "This report is for explanation and cleanup, not panic. A few items are clearly higher, a few are clearly lower, and both should be understood before the year is closed.",
)

note(
    doc,
    f"Based on QuickBooks data through {cutoff_str}. Key validation numbers include tuition {fmt_currency(ytd_cy, 2)}, marketing {fmt_currency(marketing_cy, 2)}, student handouts {fmt_currency(e_cy('5780'), 2)}, and Service Fee 5711 {fmt_currency(e_cy('5711'), 2)}.",
)

doc.save(out_path)
print("Saved:", out_path)

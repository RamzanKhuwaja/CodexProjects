"""
Short-form tax report for VAU.
"""

import datetime
import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import BS, CAL, bp, callout_blue, callout_red, ct, hr, make_doc, note, sdat, sh, shdr
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(BASE_DIR, "data", "extracted", "run_data.json"), encoding="utf-8") as f:
    D = json.load(f)

meta = D["meta"]
rev = D["revenue"]
inc = D["income"]
mkt = D["marketing"]

cutoff_str = meta["ytd_cutoff_date"]
fy_label = meta["fiscal_year_label"]
fy_end = datetime.date.fromisoformat(meta["fiscal_year_end"])
qb_profit = inc["qb_profit"]
taxes_booked = inc["taxes_paid_8500"]
ccr = inc["canada_carbon_rebate"]
h1_pretax = inc["h1_pretax_proxy"]
prior_ytd_pretax = inc["prior_ytd_pretax_proxy"]
gap_proj = mkt.get("gap_projected") or 0.0

historical = D["tax"]["historical_reference"]
PY_TOTAL_TAX = historical["prior_total_tax"]
SBD_LIMIT = historical["sbd_limit"]
prior_taxable_income = historical["prior_taxable_income"]

completion_ratio = (prior_ytd_pretax / prior_taxable_income) if prior_taxable_income else None
estimated_taxable_income = (h1_pretax / completion_ratio) if completion_ratio else h1_pretax


def calc_tax(taxable_income):
    if taxable_income <= 0:
        return 0.0
    if taxable_income <= SBD_LIMIT:
        biz_tax = taxable_income * 0.122
    else:
        biz_tax = SBD_LIMIT * 0.122 + (taxable_income - SBD_LIMIT) * 0.265
    return biz_tax + 4_000


tax_low = calc_tax(estimated_taxable_income * 0.95)
tax_mid = calc_tax(estimated_taxable_income)
tax_high = calc_tax(estimated_taxable_income * 1.05)
sbd_excess = max(0.0, h1_pretax - SBD_LIMIT)
tax_saving_on_gap = gap_proj * (0.265 if estimated_taxable_income > SBD_LIMIT else 0.122)


def fmt_currency(value, decimals=0):
    if value is None:
        return "n/a"
    return f"${value:,.{decimals}f}"


doc = make_doc()
file_date = datetime.date.today().strftime("%Y-%m-%d")
out_path = os.path.join(BASE_DIR, f"reports/claude_report_tax_vau_{file_date}.docx")
today_str = datetime.date.today().strftime("%B %d, %Y")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Tax Report")
r.bold = True
r.font.name = CAL
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("Spirit of Math Schools Vaughan  |  2236262 Ontario Inc.")
r.italic = True
r.font.name = CAL
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Report Date: {today_str}   |   Fiscal Year: {fy_label}")
r.italic = True
r.font.name = CAL
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"QuickBooks data through {cutoff_str}")
r.italic = True
r.font.name = CAL
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(6)
hr(doc)

sh(doc, "1. Key Tax Answer")
callout_blue(
    doc,
    (
        f"If you pay no installments for the year ending {fy_end.strftime('%B %d, %Y')}, the current estimated final tax bill is about "
        f"{fmt_currency(tax_mid)}. A reasonable range is {fmt_currency(tax_low)} to {fmt_currency(tax_high)}."
    ),
)
callout_red(
    doc,
    (
        f"Calm action: plan cash for about {fmt_currency(tax_mid)}. "
        "This estimate is based on how the same cutoff point converted into full-year taxable income last year."
    ),
)

summary_rows = [
    ("QB profit", fmt_currency(qb_profit, 2), "Current QuickBooks profit"),
    ("Add back tax booked", fmt_currency(taxes_booked, 2), "Tax expense in books"),
    ("Less carbon rebate", fmt_currency(ccr, 2), "Non-taxable item removed"),
    ("H1 pre-tax proxy", fmt_currency(h1_pretax, 2), "Main YTD tax number"),
    ("Prior YTD pre-tax proxy", fmt_currency(prior_ytd_pretax, 2), "Same cutoff last year"),
    ("Estimated full-year taxable income", fmt_currency(estimated_taxable_income, 2), "Seasonality-based estimate"),
    ("Estimated final tax if no installments paid", fmt_currency(tax_mid, 2), "Best current estimate"),
]
table = doc.add_table(rows=1 + len(summary_rows), cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.columns[0].width = Inches(2.6)
table.columns[1].width = Inches(1.5)
table.columns[2].width = Inches(2.4)
ct(table.rows[0].cells[0], "Item", bold=True)
ct(table.rows[0].cells[1], "Amount", bold=True)
ct(table.rows[0].cells[2], "What it means", bold=True)
shdr(table.rows[0])
for i, row in enumerate(summary_rows, 1):
    ct(table.rows[i].cells[0], row[0])
    ct(table.rows[i].cells[1], row[1])
    ct(table.rows[i].cells[2], row[2])
    sdat(table.rows[i])

sh(doc, "2. What Matters Most")
bp(
    doc,
    f"Your H1 pre-tax proxy is {fmt_currency(h1_pretax)}. That is already {fmt_currency(sbd_excess)} above the small business limit of {fmt_currency(SBD_LIMIT)}."
    if sbd_excess > 0
    else f"Your H1 pre-tax proxy is {fmt_currency(h1_pretax)}. That is still within the small business limit of {fmt_currency(SBD_LIMIT)}.",
)
bp(
    doc,
    f"Last year's actual total tax was {fmt_currency(PY_TOTAL_TAX)}. "
    f"This year's current mid-case estimate is {fmt_currency(tax_mid)}.",
)
bp(
    doc,
    f"At the same cutoff last year, the pre-tax proxy was {fmt_currency(prior_ytd_pretax)} and full-year taxable income finished at {fmt_currency(prior_taxable_income)}.",
)
bp(
    doc,
    f"If you still spend the remaining marketing gap of {fmt_currency(gap_proj)}, it may reduce tax by about {fmt_currency(tax_saving_on_gap)}.",
)

sh(doc, "3. What To Do")
for title, body in [
    (
        "Plan cash for the final bill.",
        f"Set aside about {fmt_currency(tax_mid)} so the final payment does not feel like a surprise.",
    ),
    (
        "Do not assume installments are paid unless QuickBooks proves it.",
        "The current files do not prove installment payment status, so this report treats the full estimated tax as planning cash still needed.",
    ),
    (
        "Use the remaining marketing deduction wisely.",
        f"The remaining {fmt_currency(gap_proj)} of marketing spend can help reduce tax and also meet the franchise requirement.",
    ),
    (
        "Ask the accountant to confirm any large new expense.",
        "The new Service Fee 5711 should be checked for proper tax treatment.",
    ),
]:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title + "  ")
    rb.bold = True
    rb.font.name = CAL
    rb.font.size = BS
    rn = p.add_run(body)
    rn.font.name = CAL
    rn.font.size = BS

sh(doc, "4. Bottom Line")
bp(
    doc,
    f"The number to watch is {fmt_currency(tax_mid)}. That is the current estimate of what VAU may owe if no installments are paid for {fy_label}.",
)
bp(
    doc,
    "This is only an estimate. It is built from the current QuickBooks position plus last year's archived tax outcome, not from assumed installment payments.",
)

note(
    doc,
    f"Based on QuickBooks data through {cutoff_str}. FY2024-25 actual total tax used for comparison: {fmt_currency(PY_TOTAL_TAX, 2)}.",
)

doc.save(out_path)
print("Saved:", out_path)

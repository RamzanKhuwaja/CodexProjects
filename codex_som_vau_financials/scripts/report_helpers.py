"""
report_helpers.py — Shared python-docx styling helpers for all som_vau_financials report scripts.

Import pattern (add to each report script):
    import sys, os
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from report_helpers import (
        HB, CAL, BS, make_doc,
        sbg, sbd, shdr, sdat, ct, hr, bp, sh, note, sub_header,
        callout, callout_red, callout_green, callout_blue,
        callout_red_bullets, callout_blue_bullets,
    )
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Style constants ────────────────────────────────────────────────────────────
HB  = (0xBD, 0xD7, 0xEE)   # table header background (light blue)
CAL = "Calibri"
BS  = Pt(11)


# ── Document factory ───────────────────────────────────────────────────────────
def make_doc():
    """Create a Document with standard margins and return it."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
    return doc


# ── Cell helpers ───────────────────────────────────────────────────────────────
def sbg(cell, rgb):
    """Set cell background shading colour."""
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    h = "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2])
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), h)
    tcPr.append(shd)


def sbd(cell, sz=4):
    """Add thin grey borders to a cell."""
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement("w:" + side)
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "808080")
        tcB.append(b)
    tcPr.append(tcB)


def shdr(row):
    """Style a table header row (blue background, bold text)."""
    for cell in row.cells:
        sbg(cell, HB); sbd(cell)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True; run.font.name = CAL; run.font.size = BS


def sdat(row, bold=False, bg=None):
    """Style a table data row (borders only, optional background colour as RGB tuple)."""
    for cell in row.cells:
        sbd(cell)
        if bg:
            sbg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = bold; run.font.name = CAL; run.font.size = BS


def ct(cell, text, bold=False, color=None):
    """Write text into a cell with optional bold and font colour."""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold; run.font.name = CAL; run.font.size = BS
    if color is not None:
        run.font.color.rgb = color


# ── Document structure helpers ─────────────────────────────────────────────────
def hr(doc):
    """Horizontal rule divider."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "6")
    bt.set(qn("w:space"), "1"); bt.set(qn("w:color"), "999999")
    pBdr.append(bt); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def bp(doc, text="", bold=False, sz=None):
    """Body paragraph. sz overrides default font size if provided."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.font.name = CAL; run.font.size = sz or BS
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.space_before = Pt(0)
    return p


def sh(doc, text):
    """Blue section header (13pt bold)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.name = CAL; run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)


def sub_header(doc, text):
    """Dark-grey sub-header (11.5pt bold)."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = CAL; r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)


def note(doc, text):
    """Small italic grey note (9.5pt)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True; run.font.name = CAL; run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_after = Pt(4)


# ── Callout boxes ──────────────────────────────────────────────────────────────
def callout(doc, text, color_fill="FFF3CD", text_color=None, bg_rgb=None):
    """
    Parametric callout box (single string).
    Default: yellow background, brown text.
    Use color_fill (hex string) OR bg_rgb (RGB tuple) for background colour.
    Use text_color (RGBColor) to customise text colour.
    """
    if bg_rgb is not None:
        color_fill = "{:02X}{:02X}{:02X}".format(bg_rgb[0], bg_rgb[1], bg_rgb[2])
    if text_color is None:
        text_color = RGBColor(0x7D, 0x5A, 0x00)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_fill)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = CAL; run.font.size = BS
    run.font.color.rgb = text_color
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.1)


def callout_red(doc, text):
    """Red urgent-action callout (single string)."""
    callout(doc, text, color_fill="FDECEA", text_color=RGBColor(0x8B, 0x00, 0x00))


def callout_green(doc, text):
    """Green good-news callout (single string)."""
    callout(doc, text, color_fill="EAF4EA", text_color=RGBColor(0x1A, 0x5C, 0x1A))


def callout_blue(doc, text):
    """Blue context/key-facts callout (single string)."""
    callout(doc, text, color_fill="E8F0FE", text_color=RGBColor(0x1F, 0x38, 0x96))


def callout_blue_bullets(doc, bullets):
    """
    Blue callout box — renders each item in `bullets` as a separate bullet paragraph.
    Used by Task #1 (marketing) Quick Summary section.
    """
    for text in bullets:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "DDEEFF")
        pPr.append(shd)
        run = p.add_run("\u2022  " + text)
        run.font.name = CAL; run.font.size = BS
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.1)


def callout_red_bullets(doc, bullets):
    """
    Red urgent-action callout — renders each item in `bullets` as a separate bullet paragraph.
    Used by Task #1 (marketing) Quick Summary section.
    """
    for text in bullets:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "FFE0E0")
        pPr.append(shd)
        run = p.add_run("\u26a0  " + text)
        run.font.name = CAL; run.font.size = BS
        run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.1)

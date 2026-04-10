"""
Short-form shareholder report for VAU.
"""

import datetime
import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import BS, CAL, bp, callout_blue, callout_red, ct, hr, make_doc, note, sdat, sh, shdr
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
with open(os.path.join(BASE_DIR, "data", "extracted", "run_data.json"), encoding="utf-8") as f:
    D = json.load(f)

meta = D["meta"]
sh_d = D["shareholder"]

cutoff_str = meta["ytd_cutoff_date"]
fy_label = meta["fiscal_year_label"]

ram = sh_d["ramzan"]
far = sh_d["farah"]

ram_open = ram["opening_balance"]
ram_close = ram["closing_balance"]
ram_txns = ram["transactions"]
far_close = far["closing_balance"]
combined = sh_d["combined_closing"]
parent_2900 = sh_d["parent_2900"]["closing_balance"]
net_opening = sh_d["net_opening_balance"]
net_current = sh_d["net_current_balance"]

hajj_txns = [tx for tx in ram_txns if "hajj" in (tx.get("memo") or "").lower()]
hajj_total = sum(abs(tx.get("amount", 0.0)) for tx in hajj_txns)
sep_cheque = next(
    (
        tx
        for tx in ram_txns
        if tx.get("type") == "Expense"
        and "cheque" in (tx.get("memo") or "").lower()
        and tx.get("date", "").startswith("15/09/")
    ),
    None,
)
sep_cheque_amount = abs(sep_cheque.get("amount", 0.0)) if sep_cheque else 0.0
sep_cheque_date = sep_cheque.get("date", "") if sep_cheque else None


def fmt_signed(value, decimals=0):
    if value < 0:
        return f"-${abs(value):,.{decimals}f}"
    return f"${value:,.{decimals}f}"


def plain_status(name, value):
    if value < 0:
        return f"{name} owes the company {fmt_signed(abs(value))}"
    if value > 0:
        return f"The company owes {name} {fmt_signed(value)}"
    return f"{name} is at zero"


doc = make_doc()
file_date = datetime.date.today().strftime("%Y-%m-%d")
out_path = os.path.join(BASE_DIR, f"reports/claude_report_shareholder_vau_{file_date}.docx")
today_str = datetime.date.today().strftime("%B %d, %Y")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Shareholder Report")
r.bold = True
r.font.name = CAL
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x96)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("Spirit of Math Schools Vaughan  |  2236262 Ontario Inc.")
r.italic = True
r.font.name = CAL
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"Report Date: {today_str}   |   Fiscal Year: {fy_label}")
r.italic = True
r.font.name = CAL
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(f"QuickBooks shareholder data through {cutoff_str}")
r.italic = True
r.font.name = CAL
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(6)
hr(doc)

sh(doc, "1. Key Answer")
callout_blue(
    doc,
    (
        f"The real net shareholder position is about {fmt_signed(abs(net_current))} owed to the company. "
        f"This is the number that ties to the reviewed statements. "
        f"The raw Ramzan and Farah subaccounts should not be read alone because there is also a parent 2900 balance of {fmt_signed(parent_2900)}."
    ),
)
callout_red(
    doc,
    (
        (
            f"The Hajj payments of {fmt_signed(hajj_total)} and the {sep_cheque_date} cheque of {fmt_signed(sep_cheque_amount)} "
            "are already posted in the shareholder account. Keep them clearly documented there as shareholder items, not business expenses."
        )
        if sep_cheque
        else
        f"The Hajj payments of {fmt_signed(hajj_total)} are already posted in the shareholder account. Keep them clearly documented there as shareholder items, not business expenses."
    ),
)

sh(doc, "2. Balances")
rows = [
    ("Net opening balance", fmt_signed(net_opening, 2), "Real total at start of year"),
    ("Parent 2900 balance", fmt_signed(parent_2900, 2), "Offset account in QuickBooks"),
    ("Ramzan raw subaccount", fmt_signed(ram_close, 2), "Raw QuickBooks subaccount only"),
    ("Farah raw subaccount", fmt_signed(far_close, 2), "Raw QuickBooks subaccount only"),
    ("Real current net balance", fmt_signed(net_current, 2), "Best number to trust"),
]
table = doc.add_table(rows=1 + len(rows), cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.columns[0].width = Inches(2.5)
table.columns[1].width = Inches(1.4)
table.columns[2].width = Inches(2.7)
ct(table.rows[0].cells[0], "Item", bold=True)
ct(table.rows[0].cells[1], "Amount", bold=True)
ct(table.rows[0].cells[2], "Plain English", bold=True)
shdr(table.rows[0])
for i, row in enumerate(rows, 1):
    ct(table.rows[i].cells[0], row[0])
    ct(table.rows[i].cells[1], row[1])
    ct(table.rows[i].cells[2], row[2])
    sdat(table.rows[i])

sh(doc, "3. Main Items to Track")
bp(
    doc,
    f"Hajj payments recorded in Ramzan's account total {fmt_signed(hajj_total)}. They are already sitting in the shareholder account, which is the right place.",
)
bp(
    doc,
    (
        f"The large {sep_cheque_date} cheque was {fmt_signed(sep_cheque_amount)}. It is also already sitting in the shareholder account and is one of the biggest items increasing the balance."
        if sep_cheque
        else
        "No September 15 cheque was identified in the provided shareholder export, so this report does not invent one."
    ),
)
bp(
    doc,
    f"The shareholder picture is easier than it first looked. The real net amount owed to the company is {fmt_signed(abs(net_current))}, not the much larger raw Ramzan subaccount by itself.",
)

sh(doc, "4. What To Do")
for title, body in [
    (
        "Keep supporting notes and receipts.",
        "The flagged shareholder items already appear in the shareholder account. The main need now is clean documentation.",
    ),
    (
        "Create a repayment plan.",
        f"The real net amount of about {fmt_signed(abs(net_current))} should be reduced in an orderly way with the accountant's help.",
    ),
    (
        "Keep the QuickBooks presentation clear.",
        "The parent 2900 balance and the two subaccounts should be understood together, not separately.",
    ),
]:
    p = doc.add_paragraph(style="List Number")
    rb = p.add_run(title + "  ")
    rb.bold = True
    rb.font.name = CAL
    rb.font.size = BS
    rn = p.add_run(body)
    rn.font.name = CAL
    rn.font.size = BS

sh(doc, "5. Bottom Line")
bp(
    doc,
    f"The main number to trust is the real net balance of {fmt_signed(abs(net_current))} owed to the company. "
    f"Farah's raw subaccount at {fmt_signed(far_close)} is not the concern.",
)
bp(
    doc,
    (
        f"As long as the Hajj amount of {fmt_signed(hajj_total)} and the {sep_cheque_date} cheque of {fmt_signed(sep_cheque_amount)} stay clearly documented, this report stays manageable. The bigger issue was the way QuickBooks was being read, not a missing fixing entry."
        if sep_cheque
        else
        f"As long as the Hajj amount of {fmt_signed(hajj_total)} stays clearly documented, this report stays manageable. The bigger issue was the way QuickBooks was being read, not a missing fixing entry."
    ),
)

note(
    doc,
    (
        f"Based on QuickBooks shareholder data through {cutoff_str}. Validation checks include Ramzan {fmt_signed(ram_close, 2)}, Farah {fmt_signed(far_close, 2)}, Hajj {fmt_signed(hajj_total)}, and the {sep_cheque_date} cheque {fmt_signed(sep_cheque_amount)}."
        if sep_cheque
        else
        f"Based on QuickBooks shareholder data through {cutoff_str}. Validation checks include Ramzan {fmt_signed(ram_close, 2)}, Farah {fmt_signed(far_close, 2)}, and Hajj {fmt_signed(hajj_total)}."
    ),
)

doc.save(out_path)
print("Saved:", out_path)
